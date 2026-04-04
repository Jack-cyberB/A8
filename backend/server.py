from __future__ import annotations

import concurrent.futures
import csv
import datetime as dt
import io
import json
import math
import os
import re
import socket
import time
import uuid
from collections import Counter
from copy import deepcopy
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any, Generator
from urllib.error import HTTPError, URLError
from urllib.parse import parse_qs, quote, urlencode, urlparse
from urllib.request import Request, urlopen

from backend.mysql_support import MySQLClient, sql_literal

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
PROJECT_ROOT_FALLBACK = Path(r"D:/Project/2026/A8")


def resolve_project_path(*parts: str) -> Path:
    local_path = ROOT.joinpath(*parts)
    if local_path.exists():
        return local_path
    fallback_path = PROJECT_ROOT_FALLBACK.joinpath(*parts)
    return fallback_path


ENV_FILE = ROOT / ".env"
DEMO_DATA_FILE = ROOT / "data" / "energy_dataset.csv"
NORMALIZED_DATA_FILE = ROOT / "data" / "normalized" / "energy_normalized.csv"
METADATA_FILE = resolve_project_path("data", "raw", "bdg2", "data", "metadata", "metadata.csv")
RAW_ELECTRICITY_FILE = resolve_project_path("data", "raw", "bdg2", "data", "meters", "cleaned", "electricity_cleaned.csv")
WEATHER_FILE = resolve_project_path("data", "raw", "bdg2", "data", "weather", "weather.csv")
DICT_FILE = ROOT / "data" / "ai_dictionary.json"
KNOWLEDGE_FILE = ROOT / "data" / "normalized" / "knowledge_chunks.jsonl"
RUNTIME_DIR = ROOT / "data" / "runtime"
ACTION_LOG_FILE = RUNTIME_DIR / "anomaly_actions.jsonl"
AI_CALL_LOG_FILE = RUNTIME_DIR / "ai_calls.jsonl"
NOTE_LOG_FILE = RUNTIME_DIR / "anomaly_notes.jsonl"
REGRESSION_SUMMARY_FILE = RUNTIME_DIR / "regression_summary.json"
FRONTEND_DIR = ROOT / "frontend"
SCENE_KNOWLEDGE_DOC_DIR = resolve_project_path("docs", "ragflow", "sikong-kb-pack", "main-kb")
STANDARD_KNOWLEDGE_DOC_DIR = resolve_project_path("docs", "ragflow", "standard-kb-pack", "main-kb")

TIME_FMT = "%Y-%m-%d %H:%M:%S"
CARBON_FACTOR = 0.785
SEVERITY_SCORE = {"low": 1, "medium": 2, "high": 3}
SUPPORTED_ANALYSIS_METRICS = {"electricity"}
RAGFLOW_DEFAULT_BASE_URL = "http://127.0.0.1:8088/api/v1"
RAGFLOW_DEFAULT_DATASET_IDS = ""
RAGFLOW_DEFAULT_STANDARD_DATASET_IDS = ""
RAGFLOW_DEFAULT_CHAT_ID = ""
RAGFLOW_DEFAULT_TOP_K = 6
RAGFLOW_DEFAULT_SIMILARITY_THRESHOLD = 0.2
RAGFLOW_DEFAULT_VECTOR_SIMILARITY_WEIGHT = 0.45
RAGFLOW_DEFAULT_TIMEOUT_SEC = 6.0
RAGFLOW_DEFAULT_CHAT_TIMEOUT_SEC = 60.0
STORAGE_BACKEND_DEFAULT = "mysql"

SHOWCASE_BUILDINGS = {
    "Panther_education_Genevieve": {"display_category": "教学楼", "peer_category": "teaching_building"},
    "Panther_education_Jerome": {"display_category": "实验楼", "peer_category": "lab_building"},
    "Panther_office_Patti": {"display_category": "办公楼", "peer_category": "office_building"},
    "Panther_lodging_Marisol": {"display_category": "宿舍", "peer_category": "dormitory"},
    "Panther_assembly_Denice": {"display_category": "体育馆", "peer_category": "gymnasium"},
    "Fox_public_Martin": {"display_category": "图书馆", "peer_category": "library"},
    "Fox_food_Scott": {"display_category": "食堂", "peer_category": "canteen"},
}

PEER_CATEGORY_LABELS = {
    "teaching_building": "教学楼",
    "lab_building": "实验楼",
    "office_building": "办公楼",
    "dormitory": "宿舍",
    "library": "图书馆",
    "gymnasium": "体育馆",
    "canteen": "食堂",
}

STANDARD_QUERY_KEYWORDS = (
    "标准",
    "规范",
    "定额",
    "术语",
    "条文",
    "国标",
    "地标",
    "行标",
    "通用规范",
    "设计规范",
    "运行管理标准",
    "能耗定额",
    "依据",
    "要求",
    "是否符合",
    "是否合规",
)

SCENE_QUERY_KEYWORDS = (
    "排查",
    "运维",
    "怎么做",
    "怎么办",
    "建议",
    "原因",
    "为什么",
    "处理",
    "巡检",
    "故障",
    "异常",
    "诊断",
    "优化",
    "策略",
    "怎么排查",
    "优先检查",
)

MIXED_QUERY_HINT_KEYWORDS = (
    "因素",
    "有关",
    "原理",
    "机理",
    "要求",
    "定义",
    "区别",
    "影响",
)

KNOWLEDGE_DOMAIN_TERMS = (
    "教学楼",
    "教室",
    "实验楼",
    "办公楼",
    "宿舍",
    "空调",
    "通风",
    "新风",
    "排风",
    "风机",
    "热环境",
    "热舒适",
    "闷热",
    "温度",
    "湿度",
    "太阳辐射",
    "遮阳",
    "外窗",
    "过滤器",
    "盘管",
    "供回水",
    "负荷",
    "能耗",
    "节能",
    "照明",
    "排查",
    "检查",
    "运维",
)

STATUS_NEW = "new"
STATUS_ACK = "acknowledged"
STATUS_IGNORED = "ignored"
STATUS_RESOLVED = "resolved"
STATUS_VALUES = {STATUS_NEW, STATUS_ACK, STATUS_IGNORED, STATUS_RESOLVED}

ACTION_TO_STATUS = {
    "ack": STATUS_ACK,
    "ignore": STATUS_IGNORED,
    "resolve": STATUS_RESOLVED,
}

ALLOWED_TRANSITIONS = {
    STATUS_NEW: {"ack", "ignore"},
    STATUS_ACK: {"resolve", "ignore"},
    STATUS_IGNORED: set(),
    STATUS_RESOLVED: set(),
}

ANOMALY_RULE_META = {
    "anomaly_spike": {
        "rule_name": "瞬时突增",
        "time_scope_label": "单点瞬时",
        "summary_template": "当前负荷高于建筑自身突增阈值，属于短时异常抬升。",
    },
    "anomaly_sustained_high_load": {
        "rule_name": "持续高负荷",
        "time_scope_label": "连续高负荷窗口",
        "summary_template": "连续多个时点高于建筑自身高负荷阈值，说明存在持续运行压力。",
    },
    "anomaly_off_hours_load": {
        "rule_name": "非工作时段高负荷",
        "time_scope_label": "夜间/周末时段",
        "summary_template": "在夜间或周末等非工作时段，负荷仍显著高于非工作时段基线。",
    },
    "anomaly_workhour_offline": {
        "rule_name": "工作时段低负荷/疑似停运",
        "time_scope_label": "工作时段",
        "summary_template": "在正常工作时段负荷异常偏低，可能存在设备停运、联动未开启或采集异常。",
    },
    "anomaly_baseload_high": {
        "rule_name": "夜间基线偏高",
        "time_scope_label": "夜间基线窗口",
        "summary_template": "夜间平均负荷相对日均水平偏高，说明存在待机、常开或控制策略未收敛问题。",
    },
    "anomaly_schedule_shift": {
        "rule_name": "启停时段异常",
        "time_scope_label": "日程启停窗口",
        "summary_template": "启停时段相对典型工作日显著提前或延后，疑似排程设置异常。",
    },
}

ASSISTANT_REPORT_MODULE_META = {
    "saving": {"title": "节能建议报告", "filename": "节能建议报告"},
    "diagnosis": {"title": "异常诊断报告", "filename": "异常诊断报告"},
    "interpretation": {"title": "分析解读报告", "filename": "分析解读报告"},
}


def load_local_env() -> None:
    if not ENV_FILE.exists():
        return
    try:
        for raw_line in ENV_FILE.read_text(encoding="utf-8").splitlines():
            line = raw_line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, value = line.split("=", 1)
            key = key.strip().lstrip("\ufeff")
            value = value.strip().strip("'\"")
            if key and key not in os.environ:
                os.environ[key] = value
    except OSError:
        return


load_local_env()


def parse_time(value: str | None) -> dt.datetime | None:
    if not value:
        return None
    value = value.strip()
    for fmt in (TIME_FMT, "%Y-%m-%d"):
        try:
            return dt.datetime.strptime(value, fmt)
        except ValueError:
            continue
    return None


def to_iso(ts: dt.datetime) -> str:
    return ts.strftime(TIME_FMT)


def clean_text(value: Any) -> str:
    text = str(value or "").replace("\r", "\n")
    lines = [re.sub(r"\s{2,}", " ", line).strip() for line in text.split("\n")]
    return "\n".join(line for line in lines if line).strip()


def ensure_text_list(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    result: list[str] = []
    for item in value:
        text = clean_text(item)
        if text:
            result.append(text)
    return result


def decode_json_clone(value: Any) -> Any:
    try:
        return json.loads(json.dumps(value, ensure_ascii=False))
    except Exception:
        return deepcopy(value)


def showcase_display_name(building_id: str, peer_category: str | None = None) -> str:
    config = SHOWCASE_BUILDINGS.get(building_id)
    if config:
        return f"{building_id}（{config['display_category']}）"
    label = PEER_CATEGORY_LABELS.get(peer_category or "", "")
    return f"{building_id}（{label}）" if label else building_id


def infer_peer_category(primaryspaceusage: str | None, sub_primaryspaceusage: str | None) -> str | None:
    primary = str(primaryspaceusage or "").strip().lower()
    sub = str(sub_primaryspaceusage or "").strip().lower()
    if primary == "education" and "classroom" in sub:
        return "teaching_building"
    if primary == "education" and any(token in sub for token in ("laboratory", "research", "academic")):
        return "lab_building"
    if primary == "office" and (not sub or sub == "office"):
        return "office_building"
    if primary == "lodging/residential" and any(token in sub for token in ("dormitory", "residence hall")):
        return "dormitory"
    if primary == "public services" and "library" in sub:
        return "library"
    if primary == "entertainment/public assembly" and any(token in sub for token in ("gymnasium", "stadium", "fitness center")):
        return "gymnasium"
    if primary == "food sales and service":
        return "canteen"
    return None


def infer_ragflow_web_base_url(api_base_url: str) -> str:
    base = (api_base_url or "").rstrip("/")
    if base.endswith("/api/v1"):
        return base[:-7]
    if base.endswith("/api"):
        return base[:-4]
    return base


class DiagnoseProvider:
    name = "base_provider"

    def diagnose(self, repo: "EnergyRepository", payload: dict[str, Any]) -> dict[str, Any]:
        raise NotImplementedError


class TemplateDiagnoseProvider(DiagnoseProvider):
    name = "template_provider"

    def diagnose(self, repo: "EnergyRepository", payload: dict[str, Any]) -> dict[str, Any]:
        return repo._diagnose_by_template(payload)


class LLMDiagnoseProvider(DiagnoseProvider):
    name = "llm_provider"

    def _extract_json_object(self, text: str) -> dict[str, Any]:
        raw = (text or "").strip()
        if not raw:
            raise ValueError("empty llm content")

        # 1) direct JSON parse
        try:
            obj = json.loads(raw)
            if isinstance(obj, dict):
                return obj
        except json.JSONDecodeError:
            pass

        # 2) fenced code block parse
        fence_match = re.search(r"```(?:json)?\s*([\s\S]*?)```", raw, re.IGNORECASE)
        if fence_match:
            candidate = fence_match.group(1).strip()
            try:
                obj = json.loads(candidate)
                if isinstance(obj, dict):
                    return obj
            except json.JSONDecodeError:
                pass

        # 3) first {...} parse
        start = raw.find("{")
        end = raw.rfind("}")
        if start >= 0 and end > start:
            candidate = raw[start : end + 1]
            try:
                obj = json.loads(candidate)
                if isinstance(obj, dict):
                    return obj
            except json.JSONDecodeError:
                pass

        raise ValueError("llm content is not valid json object")

    def _coerce_list_of_str(self, value: Any) -> list[str]:
        if isinstance(value, list):
            out = []
            for item in value:
                if item is None:
                    continue
                s = str(item).strip()
                if s:
                    out.append(s)
            return out
        if value is None:
            return []
        single = str(value).strip()
        return [single] if single else []

    def _sanitize_llm_result(self, llm_obj: dict[str, Any]) -> dict[str, Any]:
        confidence_raw = llm_obj.get("confidence", 0.6)
        try:
            confidence = float(confidence_raw)
        except (TypeError, ValueError):
            confidence = 0.6
        confidence = max(0.0, min(1.0, confidence))

        risk_level = str(llm_obj.get("risk_level", "")).strip().lower()
        if risk_level not in {"low", "medium", "high"}:
            risk_level = "high" if confidence >= 0.8 else "medium" if confidence >= 0.6 else "low"

        evidence_val = llm_obj.get("evidence", [])
        evidence: list[dict[str, str]] = []
        if isinstance(evidence_val, list):
            for idx, item in enumerate(evidence_val, start=1):
                if isinstance(item, dict):
                    evidence.append(
                        {
                            "chunk_id": str(item.get("chunk_id", f"llm-{idx}")),
                            "title": str(item.get("title", "LLM")),
                            "section": str(item.get("section", "")),
                            "excerpt": str(item.get("excerpt", "")).strip(),
                            "source_type": str(item.get("source_type", "llm_generated")),
                        }
                    )
                else:
                    txt = str(item).strip()
                    if txt:
                        evidence.append(
                            {
                                "chunk_id": f"llm-{idx}",
                                "title": "LLM",
                                "section": "",
                                "excerpt": txt,
                                "source_type": "llm_generated",
                            }
                        )

        return {
            "conclusion": str(llm_obj.get("conclusion", "")).strip(),
            "causes": self._coerce_list_of_str(llm_obj.get("causes", llm_obj.get("possible_causes"))),
            "steps": self._coerce_list_of_str(llm_obj.get("steps")),
            "prevention": self._coerce_list_of_str(llm_obj.get("prevention")),
            "recommended_actions": self._coerce_list_of_str(llm_obj.get("recommended_actions")),
            "evidence": evidence,
            "confidence": round(confidence, 2),
            "risk_level": risk_level,
        }

    def _call_chat_completion(
        self,
        *,
        base_url: str,
        api_key: str,
        model: str,
        timeout_sec: float,
        messages: list[dict[str, str]],
        max_tokens: int = 900,
        response_format: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        endpoint = f"{base_url.rstrip('/')}/chat/completions"
        payload: dict[str, Any] = {
            "model": model,
            "messages": messages,
            "temperature": 0.2,
            "max_tokens": max(128, int(max_tokens)),
        }
        if response_format:
            payload["response_format"] = response_format
        body = json.dumps(payload).encode("utf-8")
        req = Request(
            endpoint,
            data=body,
            method="POST",
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            },
        )
        with urlopen(req, timeout=timeout_sec) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        return data

    def _call_chat_completion_stream(
        self,
        *,
        base_url: str,
        api_key: str,
        model: str,
        timeout_sec: float,
        messages: list[dict[str, str]],
    ) -> Generator[str, None, None]:
        """Yield text tokens from a streaming chat completion (SSE)."""
        endpoint = f"{base_url.rstrip('/')}/chat/completions"
        body = json.dumps(
            {
                "model": model,
                "messages": messages,
                "temperature": 0.3,
                "stream": True,
            }
        ).encode("utf-8")
        req = Request(
            endpoint,
            data=body,
            method="POST",
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            },
        )
        with urlopen(req, timeout=timeout_sec) as resp:
            for raw_line in resp:
                line = raw_line.decode("utf-8").strip()
                if not line.startswith("data:"):
                    continue
                data_str = line[5:].strip()
                if data_str == "[DONE]":
                    break
                try:
                    chunk = json.loads(data_str)
                    delta = chunk.get("choices", [{}])[0].get("delta", {})
                    text = delta.get("content", "")
                    if text:
                        yield text
                except (json.JSONDecodeError, IndexError, KeyError):
                    continue

    def _build_diagnosis_prompt_context(
        self,
        context: dict[str, Any],
        diag_template: dict[str, Any],
    ) -> str:
        window_context = context.get("window_context") or {}
        peer_context = context.get("peer_context") or {}
        weather_context = context.get("weather_context") or {}
        systems = "、".join(context.get("likely_systems") or []) or "未识别"
        tags = "、".join(context.get("phenomenon_tags") or []) or "无"
        lines = [
            f"建筑：{context.get('building_name', '')}",
            f"建筑类型：{context.get('building_type', '')}",
            f"异常类型：{diag_template.get('anomaly_name', '')}",
            f"异常时间：{context.get('timestamp', '')}",
            f"当前负荷：{context.get('value_kwh', '')} kWh",
            f"偏差比例：{context.get('deviation_pct', '')}%",
            f"24h基线：{window_context.get('baseline_24h_avg_kwh', '')} kWh",
            f"同小时均值：{window_context.get('same_hour_avg_kwh', '')} kWh",
            f"异常前24h均值：{window_context.get('before_24h_avg_kwh', '')} kWh",
            f"异常后24h均值：{window_context.get('after_24h_avg_kwh', '')} kWh",
            f"同类偏差：{peer_context.get('gap_pct', '')}%",
            f"同类百分位：{peer_context.get('peer_percentile', '')}",
            f"异常时温度：{weather_context.get('temperature_c', '')}°C",
            f"温度区间：{weather_context.get('temperature_band', '')}",
            f"优先怀疑系统：{systems}",
            f"现象标签：{tags}",
        ]
        return "\n".join(line for line in lines if line and not line.endswith("："))

    def diagnose(self, repo: "EnergyRepository", payload: dict[str, Any]) -> dict[str, Any]:
        if payload.get("simulate_llm_failure"):
            raise RuntimeError("Simulated llm failure")

        api_key = os.getenv("OPENAI_API_KEY", "").strip()
        if not api_key:
            raise RuntimeError("LLM provider not configured")

        base_url = os.getenv("OPENAI_BASE_URL", "https://api.deepseek.com").strip() or "https://api.deepseek.com"
        model = os.getenv("OPENAI_MODEL", "deepseek-chat").strip() or "deepseek-chat"
        base_timeout_sec = float(os.getenv("OPENAI_TIMEOUT_SEC", "12"))
        timeout_sec = float(os.getenv("OPENAI_DIAGNOSE_TIMEOUT_SEC", str(max(base_timeout_sec, 45.0))))
        max_retries = int(os.getenv("OPENAI_MAX_RETRIES", "2"))

        template_result = repo._diagnose_by_template(payload)
        diag_template = template_result.get("diagnosis", {})
        context = template_result.get("context", {})

        anomaly_name = str(diag_template.get("anomaly_name", ""))
        evidence = diag_template.get("evidence", [])
        data_evidence = diag_template.get("data_evidence", [])
        evidence_text = "\n".join(
            [
                f"- {str(x.get('title', ''))}: {str(x.get('excerpt', ''))}"
                for x in evidence[:3]
                if isinstance(x, dict)
            ]
        ).strip()
        data_evidence_text = "\n".join(
            [
                f"- {str(x.get('title', ''))}: {str(x.get('detail', ''))}"
                for x in data_evidence[:6]
                if isinstance(x, dict)
            ]
        ).strip()
        prompt_context_text = self._build_diagnosis_prompt_context(context, diag_template)
        prompt_message = str(payload.get("message", "")).strip()
        if not prompt_message:
            prompt_message = "请基于异常上下文给出诊断建议。"

        system_prompt = (
            "你是建筑能源运维诊断助手，面向真实运维排障场景。必须只输出一个JSON对象，不要输出其他文本。"
            "JSON字段必须包含：conclusion, causes, steps, prevention, recommended_actions, evidence, confidence, risk_level。"
            "其中causes/steps/prevention/recommended_actions为字符串数组，每个数组至少包含3条内容，不允许输出空数组。"
            "evidence为数组。risk_level只能是low/medium/high。confidence为0.0到1.0之间的数字。"
            "conclusion必须引用建筑名称、异常发生时间段、当前负荷、24h基线和具体偏差数值（如'较基线偏高X%'）。"
            "如提供了知识证据，请在conclusion中用[1][2]等标注引用来源。"
            "causes禁止写成'设备异常启动''临时任务''传感器偏差'这类泛词，必须写出具体系统、场景和导致偏高/偏低的机制。"
            "steps必须是按顺序执行的排查动作，明确先查什么数据、再查什么设备、最后怎么交叉验证。"
            "recommended_actions必须是本班次就能执行的即时动作，prevention必须是后续制度、排程或阈值治理动作。"
            "请使用中文，不要复述空话，不要只摆数据，要像资深运维工程师给现场班组下达诊断意见。"
        )
        user_prompt = (
            f"异常摘要:\n{prompt_context_text}\n"
            f"数据证据:\n{data_evidence_text}\n"
            f"知识证据:\n{evidence_text}\n"
            f"用户问题: {prompt_message}\n"
            "输出要求：\n"
            "1. conclusion 用 1 段话讲清这次异常意味着什么，先点判断，再点理由。\n"
            "2. causes 至少 3 条，每条都要写成“系统/对象 + 异常机制 + 为什么符合当前数据”。\n"
            "3. steps 至少 3 条，必须是现场可执行顺序，优先写分项回路、BMS排程、设备状态、现场记录。\n"
            "4. recommended_actions 至少 3 条，必须是立即动作，不要和 prevention 重复。\n"
            "5. prevention 至少 3 条，必须是防复发治理动作。\n"
            "6. 如果知识证据和数据证据不足以支持某个判断，就不要编造。\n"
            "请输出严格JSON。"
        )

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ]

        last_err: Exception | None = None
        for attempt in range(max_retries + 1):
            try:
                response = self._call_chat_completion(
                    base_url=base_url,
                    api_key=api_key,
                    model=model,
                    timeout_sec=timeout_sec,
                    messages=messages,
                    response_format={"type": "json_object"},
                )
                choices = response.get("choices", [])
                if not choices:
                    raise RuntimeError("llm response missing choices")
                content = choices[0].get("message", {}).get("content", "")
                if isinstance(content, list):
                    content = "".join(
                        str(part.get("text", "")) if isinstance(part, dict) else str(part)
                        for part in content
                    )
                llm_obj = self._extract_json_object(str(content))
                llm_diag = self._sanitize_llm_result(llm_obj)

                merged = dict(diag_template)
                llm_causes = llm_diag.get("causes", llm_diag.get("possible_causes"))
                merged["conclusion"] = str(llm_diag.get("conclusion", "")).strip() or diag_template.get("conclusion", "")
                merged["causes"] = repo._clean_text_list(llm_causes, max_items=4)
                merged["steps"] = repo._clean_text_list(llm_diag.get("steps"), max_items=4)
                merged["prevention"] = repo._clean_text_list(llm_diag.get("prevention"), max_items=4)
                merged["recommended_actions"] = repo._clean_text_list(llm_diag.get("recommended_actions"), max_items=3)
                if llm_diag.get("confidence") not in (None, ""):
                    merged["confidence"] = llm_diag["confidence"]
                if llm_diag.get("risk_level"):
                    merged["risk_level"] = llm_diag["risk_level"]
                merged["possible_causes"] = merged.get("causes", [])
                merged["data_evidence"] = diag_template.get("data_evidence", [])
                merged["evidence"] = diag_template.get("evidence", [])
                if not repo._required_diag_fields_complete(merged):
                    raise RuntimeError("llm returned incomplete diagnosis fields")
                return {"diagnosis": merged, "context": context}
            except HTTPError as exc:
                retryable = exc.code == 429 or 500 <= exc.code < 600
                last_err = RuntimeError(f"llm http status {exc.code}")
                if retryable and attempt < max_retries:
                    time.sleep(0.6 * (attempt + 1))
                    continue
                break
            except (URLError, TimeoutError, socket.timeout, ConnectionResetError) as exc:
                last_err = RuntimeError(f"llm network error: {type(exc).__name__}")
                if attempt < max_retries:
                    time.sleep(0.6 * (attempt + 1))
                    continue
                break
            except Exception as exc:
                last_err = RuntimeError(str(exc) if str(exc) else f"llm parse error: {type(exc).__name__}")
                break

        raise RuntimeError(str(last_err or "llm unknown error"))


class EnergyRepository:
    def __init__(
        self,
        demo_data_file: Path,
        normalized_data_file: Path,
        metadata_file: Path,
        weather_file: Path,
        dict_file: Path,
        knowledge_file: Path,
        action_log_file: Path,
        ai_call_log_file: Path,
        note_log_file: Path,
        regression_summary_file: Path,
    ) -> None:
        self.demo_data_file = demo_data_file
        self.normalized_data_file = normalized_data_file
        self.metadata_file = metadata_file
        self.raw_electricity_file = RAW_ELECTRICITY_FILE
        self.weather_file = weather_file
        self.dict_file = dict_file
        self.knowledge_file = knowledge_file
        self.action_log_file = action_log_file
        self.ai_call_log_file = ai_call_log_file
        self.note_log_file = note_log_file
        self.regression_summary_file = regression_summary_file

        self.bdq2_metadata = self._load_bdg2_metadata()
        self.showcase_config = SHOWCASE_BUILDINGS
        self.peer_category_to_buildings = self._build_peer_category_to_buildings()
        self.raw_electricity_headers = self._load_raw_electricity_headers()
        self.compare_pool_cache: dict[tuple[str, str | None, str | None], dict[str, float]] = {}
        self.rows = self._load_rows()
        self.building_site_map = self._load_building_site_map()
        self.weather_by_site = self._load_weather_by_site()
        self.by_building: dict[str, list[dict[str, Any]]] = {}
        self.stats: dict[str, dict[str, float]] = {}
        self.anomalies: list[dict[str, Any]] = []
        self.buildings_meta: dict[str, dict[str, Any]] = {}
        self.dict_data = self._load_dictionary()
        self.knowledge_chunks = self._load_knowledge_chunks()
        self.providers: dict[str, DiagnoseProvider] = {
            "template": TemplateDiagnoseProvider(),
            "llm": LLMDiagnoseProvider(),
        }

        self.action_events: list[dict[str, Any]] = []
        self.action_index: dict[int, dict[str, Any]] = {}
        self.ai_events: list[dict[str, Any]] = []
        self.note_events: list[dict[str, Any]] = []
        self.note_index: dict[int, dict[str, Any]] = {}

        self._prepare_indexes()
        self._ensure_runtime_storage()
        self._load_actions()
        self._load_ai_events()
        self._load_notes()

    def _ensure_runtime_storage(self) -> None:
        self.action_log_file.parent.mkdir(parents=True, exist_ok=True)
        if not self.action_log_file.exists():
            self.action_log_file.write_text("", encoding="utf-8")
        if not self.ai_call_log_file.exists():
            self.ai_call_log_file.write_text("", encoding="utf-8")
        if not self.note_log_file.exists():
            self.note_log_file.write_text("", encoding="utf-8")
        if not self.regression_summary_file.exists():
            self.regression_summary_file.write_text(
                json.dumps(
                    {
                        "updated_at": None,
                        "all_ok": False,
                        "status": "unknown",
                        "steps": [],
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )

    def _load_rows(self) -> list[dict[str, Any]]:
        if self.normalized_data_file.exists():
            return self._load_rows_normalized(self.normalized_data_file)
        return self._load_rows_demo(self.demo_data_file)

    def _load_rows_normalized(self, file_path: Path) -> list[dict[str, Any]]:
        rows: list[dict[str, Any]] = []
        idx = 1
        with file_path.open("r", encoding="utf-8", newline="") as f:
            reader = csv.DictReader(f)
            for raw in reader:
                ts = parse_time(raw.get("timestamp", ""))
                if not ts:
                    continue
                try:
                    kwh = float(raw.get("electricity_kwh", ""))
                except ValueError:
                    continue
                rows.append(
                    {
                        "record_id": idx,
                        "building_id": raw.get("building_id", "UNKNOWN"),
                        "building_name": raw.get("building_name", "UNKNOWN"),
                        "building_type": raw.get("building_type", "unknown"),
                        "timestamp": ts,
                        "hour": ts.hour,
                        "electricity_kwh": kwh,
                        "source": raw.get("source", "normalized"),
                    }
                )
                idx += 1
        return rows

    def _load_rows_demo(self, file_path: Path) -> list[dict[str, Any]]:
        rows: list[dict[str, Any]] = []
        with file_path.open("r", encoding="utf-8", newline="") as f:
            reader = csv.DictReader(f)
            for raw in reader:
                ts = parse_time(raw.get("timestamp", ""))
                if not ts:
                    continue
                rows.append(
                    {
                        "record_id": int(raw.get("record_id", len(rows) + 1)),
                        "building_id": raw.get("building_id", "UNKNOWN"),
                        "building_name": raw.get("building_name", "UNKNOWN"),
                        "building_type": raw.get("building_type", "unknown"),
                        "timestamp": ts,
                        "hour": int(raw.get("hour", ts.hour)),
                        "electricity_kwh": float(raw.get("electricity_kwh", 0.0)),
                        "source": "demo",
                    }
                )
        return rows

    def _load_dictionary(self) -> dict[str, Any]:
        with self.dict_file.open("r", encoding="utf-8-sig") as f:
            return json.load(f).get("anomaly_type_dict", {})

    def _anomaly_dictionary(self, anomaly_type: str) -> dict[str, Any]:
        return self.dict_data.get(anomaly_type, {})

    def _anomaly_name(self, anomaly_type: str) -> str:
        return str(self._anomaly_dictionary(anomaly_type).get("name", "")).strip() or anomaly_type

    def _rule_meta(self, anomaly_type: str) -> dict[str, str]:
        meta = ANOMALY_RULE_META.get(anomaly_type, {})
        return {
            "rule_code": anomaly_type,
            "rule_name": str(meta.get("rule_name", "")).strip() or self._anomaly_name(anomaly_type),
            "time_scope_label": str(meta.get("time_scope_label", "")).strip() or "异常时段",
            "summary_template": str(meta.get("summary_template", "")).strip() or "当前点位命中异常规则。",
        }

    def _load_bdg2_metadata(self) -> dict[str, dict[str, Any]]:
        if not self.metadata_file.exists():
            return {}
        metadata: dict[str, dict[str, Any]] = {}
        with self.metadata_file.open("r", encoding="utf-8", newline="") as f:
            reader = csv.DictReader(f)
            for raw in reader:
                building_id = str(raw.get("building_id", "")).strip()
                if not building_id:
                    continue
                peer_category = infer_peer_category(raw.get("primaryspaceusage"), raw.get("sub_primaryspaceusage"))
                metadata[building_id] = {
                    "building_id": building_id,
                    "site_id": str(raw.get("site_id", "")).strip(),
                    "primaryspaceusage": str(raw.get("primaryspaceusage", "")).strip(),
                    "sub_primaryspaceusage": str(raw.get("sub_primaryspaceusage", "")).strip(),
                    "peer_category": peer_category,
                    "display_category": SHOWCASE_BUILDINGS.get(building_id, {}).get("display_category") or PEER_CATEGORY_LABELS.get(peer_category or "", ""),
                    "display_name": showcase_display_name(building_id, peer_category),
                }
        return metadata

    def _build_peer_category_to_buildings(self) -> dict[str, list[str]]:
        mapping: dict[str, list[str]] = {}
        for building_id, meta in self.bdq2_metadata.items():
            peer_category = str(meta.get("peer_category") or "").strip()
            if not peer_category:
                continue
            mapping.setdefault(peer_category, []).append(building_id)
        return mapping

    def _load_raw_electricity_headers(self) -> set[str]:
        if not self.raw_electricity_file.exists():
            return set()
        with self.raw_electricity_file.open("r", encoding="utf-8", newline="") as f:
            reader = csv.reader(f)
            headers = next(reader, [])
        return {str(item).strip() for item in headers[1:] if str(item).strip()}

    def _load_building_site_map(self) -> dict[str, str]:
        mapping: dict[str, str] = {}
        for building_id, meta in self.bdq2_metadata.items():
            site_id = str(meta.get("site_id", "")).strip()
            if site_id:
                mapping[building_id] = site_id
        return mapping

    def _load_weather_by_site(self) -> dict[str, dict[dt.datetime, dict[str, float]]]:
        if not self.weather_file.exists():
            return {}
        weather_by_site: dict[str, dict[dt.datetime, dict[str, float]]] = {}
        with self.weather_file.open("r", encoding="utf-8", newline="") as f:
            reader = csv.DictReader(f)
            for raw in reader:
                site_id = str(raw.get("site_id", "")).strip()
                timestamp = parse_time(raw.get("timestamp", ""))
                if not site_id or not timestamp:
                    continue
                try:
                    temperature = float(raw.get("airTemperature", ""))
                except (TypeError, ValueError):
                    continue
                try:
                    wind_speed = float(raw.get("windSpeed", ""))
                except (TypeError, ValueError):
                    wind_speed = 0.0
                weather_by_site.setdefault(site_id, {})[timestamp] = {
                    "temperature_c": round(temperature, 2),
                    "wind_speed": round(wind_speed, 2),
                }
        return weather_by_site

    def _load_knowledge_chunks(self) -> list[dict[str, Any]]:
        if not self.knowledge_file.exists():
            return []
        chunks: list[dict[str, Any]] = []
        with self.knowledge_file.open("r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    chunks.append(json.loads(line))
                except json.JSONDecodeError:
                    continue
        return chunks

    def _ragflow_settings(self) -> dict[str, Any]:
        base_url = (os.getenv("RAGFLOW_BASE_URL", RAGFLOW_DEFAULT_BASE_URL).strip() or RAGFLOW_DEFAULT_BASE_URL).rstrip("/")
        dataset_ids_raw = os.getenv("RAGFLOW_DATASET_IDS", RAGFLOW_DEFAULT_DATASET_IDS).strip() or RAGFLOW_DEFAULT_DATASET_IDS
        dataset_ids = [item.strip() for item in dataset_ids_raw.split(",") if item.strip()]
        standard_dataset_ids_raw = os.getenv("RAGFLOW_STANDARD_DATASET_IDS", RAGFLOW_DEFAULT_STANDARD_DATASET_IDS).strip() or RAGFLOW_DEFAULT_STANDARD_DATASET_IDS
        standard_dataset_ids = [item.strip() for item in standard_dataset_ids_raw.split(",") if item.strip()]
        api_key = os.getenv("RAGFLOW_API_KEY", "").strip()
        web_base_url = (os.getenv("RAGFLOW_WEB_BASE_URL", "").strip() or infer_ragflow_web_base_url(base_url)).rstrip("/")
        chat_id = (
            os.getenv("RAGFLOW_CHAT_ID", "").strip()
            or os.getenv("RAGFLOW_SHARED_DIALOG_ID", "").strip()
            or RAGFLOW_DEFAULT_CHAT_ID
        )
        try:
            top_k = max(1, int(os.getenv("RAGFLOW_TOP_K", str(RAGFLOW_DEFAULT_TOP_K))))
        except ValueError:
            top_k = RAGFLOW_DEFAULT_TOP_K
        try:
            similarity_threshold = float(os.getenv("RAGFLOW_SIMILARITY_THRESHOLD", str(RAGFLOW_DEFAULT_SIMILARITY_THRESHOLD)))
        except ValueError:
            similarity_threshold = RAGFLOW_DEFAULT_SIMILARITY_THRESHOLD
        try:
            vector_similarity_weight = float(os.getenv("RAGFLOW_VECTOR_SIMILARITY_WEIGHT", str(RAGFLOW_DEFAULT_VECTOR_SIMILARITY_WEIGHT)))
        except ValueError:
            vector_similarity_weight = RAGFLOW_DEFAULT_VECTOR_SIMILARITY_WEIGHT
        try:
            timeout_sec = float(os.getenv("RAGFLOW_TIMEOUT_SEC", str(RAGFLOW_DEFAULT_TIMEOUT_SEC)))
        except ValueError:
            timeout_sec = RAGFLOW_DEFAULT_TIMEOUT_SEC
        try:
            chat_timeout_sec = float(os.getenv("RAGFLOW_CHAT_TIMEOUT_SEC", str(RAGFLOW_DEFAULT_CHAT_TIMEOUT_SEC)))
        except ValueError:
            chat_timeout_sec = RAGFLOW_DEFAULT_CHAT_TIMEOUT_SEC

        all_dataset_ids = dataset_ids + [item for item in standard_dataset_ids if item not in dataset_ids]
        enabled = bool(base_url and all_dataset_ids)
        configured = bool(enabled and api_key)
        chat_ready = bool(configured and all_dataset_ids and chat_id)
        return {
            "base_url": base_url,
            "web_base_url": web_base_url,
            "dataset_ids": dataset_ids,
            "scene_dataset_ids": dataset_ids,
            "standard_dataset_ids": standard_dataset_ids,
            "dataset_count": len(dataset_ids),
            "standard_dataset_count": len(standard_dataset_ids),
            "all_dataset_count": len(all_dataset_ids),
            "api_key": api_key,
            "top_k": top_k,
            "similarity_threshold": similarity_threshold,
            "vector_similarity_weight": vector_similarity_weight,
            "timeout_sec": timeout_sec,
            "chat_timeout_sec": max(timeout_sec, chat_timeout_sec),
            "enabled": enabled,
            "configured": configured,
            "chat_id": chat_id,
            "chat_ready": chat_ready,
            "assistant_ready": chat_ready,
            "standard_enabled": bool(base_url and standard_dataset_ids),
            "standard_configured": bool(api_key and standard_dataset_ids),
        }

    def _sanitize_excerpt(self, value: Any, limit: int = 220) -> str:
        text = str(value or "")
        text = re.sub(r"</?em>", "", text, flags=re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)
        text = text.replace("\r", " ").replace("\n", " ")
        text = re.sub(r"\s{2,}", " ", text).strip()
        if len(text) <= limit:
            return text
        return f"{text[:limit].rstrip()}..."

    def _sanitize_multiline_reference_text(self, value: Any) -> str:
        text = str(value or "")
        text = re.sub(r"</?p>", "\n\n", text, flags=re.IGNORECASE)
        text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
        text = re.sub(r"</?em>", "", text, flags=re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)
        text = text.replace("\r", "\n")
        lines = [re.sub(r"[ \t]{2,}", " ", line).strip() for line in text.split("\n")]
        text = "\n".join(lines)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        return text

    def _strip_knowledge_document_suffix(self, value: Any) -> str:
        text = str(value or "").strip()
        text = re.sub(r"\.(md|markdown|txt|pdf)$", "", text, flags=re.IGNORECASE)
        text = re.sub(r"(?i)([-_\s]?part\d+)$", "", text)
        return text.strip(" -_")

    def _normalize_knowledge_document_lookup_key(self, value: Any) -> str:
        text = self._strip_knowledge_document_suffix(value).lower()
        return re.sub(r"[^0-9a-z\u4e00-\u9fff]+", "", text)

    def _ragflow_document_key(self, chunk: dict[str, Any], fallback: str = "") -> str:
        raw_title = (
            chunk.get("document_name")
            or chunk.get("document_keyword")
            or chunk.get("docnm_kwd")
            or chunk.get("doc_title")
            or chunk.get("title")
            or chunk.get("document_id")
            or fallback
        )
        return self._strip_knowledge_document_suffix(raw_title)

    def _extract_ragflow_snippet(self, chunk: dict[str, Any]) -> str:
        raw = (
            chunk.get("content_with_weight")
            or chunk.get("content")
            or chunk.get("highlight")
            or chunk.get("text")
            or ""
        )
        text = self._sanitize_multiline_reference_text(raw)
        if not text:
            return ""

        question_match = re.search(r"Question:\s*(.+?)(?:\nAnswer:|\n|$)", text, flags=re.IGNORECASE | re.DOTALL)
        answer_match = re.search(r"Answer:\s*(.+)$", text, flags=re.IGNORECASE | re.DOTALL)
        if question_match or answer_match:
            question = re.sub(r"\s+", " ", question_match.group(1)).strip() if question_match else ""
            answer = re.sub(r"\s+", " ", answer_match.group(1)).strip() if answer_match else ""
            if answer and re.fullmatch(r"共\s*\d+\s*条问答。?", answer):
                return ""
            parts: list[str] = []
            if question:
                parts.append(f"问题：{question}")
            if answer:
                parts.append(f"答案：{answer}")
            return "\n\n".join(parts).strip()
        return text

    def _normalize_ragflow_title(self, chunk: dict[str, Any], fallback: str = "RAGFlow 片段") -> str:
        raw_title = (
            chunk.get("document_name")
            or chunk.get("document_keyword")
            or chunk.get("docnm_kwd")
            or chunk.get("doc_title")
            or chunk.get("title")
            or chunk.get("document_id")
            or fallback
        )
        title = str(raw_title or "").strip()
        title = re.sub(r"\.(md|markdown|txt|pdf)$", "", title, flags=re.IGNORECASE)
        return title or fallback

    def _extract_ragflow_excerpt(self, chunk: dict[str, Any]) -> str:
        snippet = self._extract_ragflow_snippet(chunk)
        if not snippet:
            return ""
        return self._sanitize_excerpt(snippet)

    def _is_noisy_ragflow_chunk(self, chunk: dict[str, Any], excerpt: str) -> bool:
        text = str(excerpt or "")
        if not text:
            return True
        if re.search(r"共\s*\d+\s*条问答", text):
            return True
        raw = str(
            chunk.get("content_with_weight")
            or chunk.get("content")
            or chunk.get("highlight")
            or chunk.get("text")
            or ""
        )
        if raw.count("[fQTQT") >= 2:
            return True
        if len(re.findall(r"/G[0-9A-F]{2,3}", raw, flags=re.IGNORECASE)) >= 6:
            return True
        chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", text))
        weird_chars = len(re.findall(r"[/\\\\][A-Z0-9]{2,4}", raw))
        if chinese_chars < 6 and weird_chars >= 6:
            return True
        return False

    def _clean_ragflow_answer_text(self, value: Any) -> str:
        text = str(value or "")
        text = re.sub(r"\[ID:\s*\d+\]", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\[\s*\d+\]", "", text)
        text = re.sub(r"\[\s*ID:\s*$", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\[\s*\d+\s*$", "", text)
        text = re.sub(r"ID:\s*\d+\s*\]", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\bID:\s*", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\bD:\s*", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\bD:\s*\d+\]", "", text, flags=re.IGNORECASE)
        text = text.replace("**", "")
        text = text.replace("[", "").replace("]", "")
        text = text.replace("\r", "")
        text = re.sub(r"[ \t]+\n", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]{2,}", " ", text)
        text = re.sub(r"\s+([，。！？；：、])", r"\1", text)
        return text.strip()

    def _postprocess_knowledge_answer(self, value: Any) -> str:
        text = self._clean_ragflow_answer_text(value)
        if not text:
            return ""
        text = text.replace("\r", "\n")
        text = re.sub(r"(?m)^\s*(结论|依据与分析|标准依据|优先检查|运维建议|执行提示|关键要求|说明)\s*[:：]?\s*", "", text)
        text = re.sub(r"[ \t]+\n", "\n", text)
        text = re.sub(r"\n[ \t]+", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _merge_ragflow_stream_text(self, full_text: str, incoming_text: Any) -> tuple[str, str]:
        piece = str(incoming_text or "")
        if not piece:
            return full_text, ""
        if piece.startswith(full_text):
            delta = piece[len(full_text) :]
            return piece, delta
        if full_text.startswith(piece):
            return full_text, ""
        if piece in full_text:
            return full_text, ""
        merged = f"{full_text}{piece}"
        return merged, piece

    def _normalize_evidence_item(
        self,
        *,
        chunk_id: Any,
        title: Any,
        section: Any,
        excerpt: Any,
        source_type: str,
        similarity: float | None = None,
        snippet_text: Any = "",
        document_key: Any = "",
    ) -> dict[str, Any]:
        return {
            "chunk_id": str(chunk_id or "").strip() or f"{source_type}-chunk",
            "title": str(title or "").strip() or ("RAGFlow" if source_type == "ragflow" else "本地知识"),
            "section": str(section or "").strip(),
            "excerpt": self._sanitize_excerpt(excerpt),
            "snippet_text": self._sanitize_multiline_reference_text(snippet_text or excerpt),
            "document_key": self._strip_knowledge_document_suffix(document_key or title),
            "source_type": source_type,
            "similarity": round(float(similarity), 4) if similarity is not None else None,
        }

    def _knowledge_document_root_for_source(self, source_type: str) -> Path:
        return STANDARD_KNOWLEDGE_DOC_DIR if str(source_type or "").strip() == "standard" else SCENE_KNOWLEDGE_DOC_DIR

    def _read_text_file(self, path: Path) -> str:
        last_error: Exception | None = None
        for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk"):
            try:
                return path.read_text(encoding=encoding)
            except Exception as exc:  # pragma: no cover - fallback branch depends on file encoding
                last_error = exc
        raise RuntimeError(f"无法读取原文文件：{path}") from last_error

    def _locate_knowledge_document(self, source_type: str, title: Any, document_key: Any) -> Path:
        root = self._knowledge_document_root_for_source(source_type)
        if not root.exists():
            raise LookupError(f"原文目录不存在：{root}")

        candidate_keys = [
            self._normalize_knowledge_document_lookup_key(document_key),
            self._normalize_knowledge_document_lookup_key(title),
        ]
        candidate_keys = [item for item in candidate_keys if item]
        if not candidate_keys:
            raise LookupError("引用缺少可定位的原文标识")

        files = [item for item in root.iterdir() if item.is_file()]
        for candidate in candidate_keys:
            for file_path in files:
                if self._normalize_knowledge_document_lookup_key(file_path.stem) == candidate:
                    return file_path
        for candidate in candidate_keys:
            for file_path in files:
                normalized_name = self._normalize_knowledge_document_lookup_key(file_path.stem)
                if normalized_name.endswith(candidate) or candidate.endswith(normalized_name):
                    return file_path
        raise LookupError("未找到对应原文")

    def ask_ragflow_reference_document(self, payload: dict[str, Any]) -> dict[str, Any]:
        source_type = str(payload.get("source_type", "")).strip() or "ragflow"
        title = str(payload.get("title", "")).strip()
        document_key = str(payload.get("document_key", "")).strip()
        if not title and not document_key:
            raise ValueError("title or document_key is required")
        file_path = self._locate_knowledge_document(source_type, title, document_key)
        return {
            "title": file_path.stem,
            "source_type": source_type,
            "document_key": self._strip_knowledge_document_suffix(document_key or file_path.stem),
            "content": self._read_text_file(file_path),
            "format": "markdown",
        }

    def _knowledge_route_for_question(self, question: str) -> str:
        text = str(question or "").strip().lower()
        if not text:
            return "scene"
        standard_hits = sum(1 for token in STANDARD_QUERY_KEYWORDS if token.lower() in text)
        scene_hits = sum(1 for token in SCENE_QUERY_KEYWORDS if token.lower() in text)
        if standard_hits and scene_hits:
            return "mixed"
        if standard_hits:
            return "standard"
        return "scene"

    def _should_use_mixed_knowledge_route(self, question: str, settings: dict[str, Any]) -> bool:
        if not settings.get("scene_dataset_ids") or not settings.get("standard_dataset_ids"):
            return False
        text = str(question or "").strip().lower()
        if not text:
            return False
        return any(token.lower() in text for token in MIXED_QUERY_HINT_KEYWORDS)

    def _ragflow_dataset_route(self, question: str) -> tuple[str, list[str]]:
        settings = self._ragflow_settings()
        scene_ids = list(settings.get("scene_dataset_ids") or [])
        standard_ids = list(settings.get("standard_dataset_ids") or [])
        route = self._knowledge_route_for_question(question)
        if route == "scene" and self._should_use_mixed_knowledge_route(question, settings):
            route = "mixed"
        if route == "standard":
            return ("standard", standard_ids or scene_ids)
        if route == "mixed":
            merged = scene_ids + [item for item in standard_ids if item not in scene_ids]
            return ("mixed", merged or scene_ids)
        return ("scene", scene_ids)

    def _ragflow_source_type_for_dataset(self, dataset_id: str, settings: dict[str, Any]) -> str:
        if dataset_id and dataset_id in set(settings.get("standard_dataset_ids") or []):
            return "standard"
        return "ragflow"

    def _knowledge_query_terms(self, text: str) -> list[str]:
        raw = str(text or "").strip().lower()
        if not raw:
            return []
        terms = [term for term in KNOWLEDGE_DOMAIN_TERMS if term.lower() in raw]
        if terms:
            return terms
        parts = re.split(r"[，。！？、；：\s“”‘’（）()]+", raw)
        return [part for part in parts if len(part) >= 2][:8]

    def _knowledge_relevance_score(self, question: str, title: str, excerpt: str, similarity: float | None) -> float:
        haystack = f"{title}\n{excerpt}".lower()
        score = float(similarity or 0.0) * 10.0
        overlap_count = 0
        for term in self._knowledge_query_terms(question):
            term_lower = term.lower()
            if term_lower and term_lower in haystack:
                overlap_count += 1
                score += max(1.0, min(len(term_lower), 4))
        if "问题：" in excerpt and "答案：" in excerpt:
            score += 0.6
        if overlap_count == 0 and float(similarity or 0.0) < 0.65:
            score -= 2.5
        return score

    def _build_knowledge_retrieval_query(self, question: str, route: str) -> str:
        raw = str(question or "").strip()
        if not raw:
            return ""
        pieces = [raw]
        terms = self._knowledge_query_terms(raw)[:6]
        if terms:
            pieces.append("关键词：" + "、".join(terms))
        if route in {"standard", "mixed"}:
            pieces.append("关注：标准要求、设计规范、运行管理")
        if route in {"scene", "mixed"} and any(token in raw for token in ("排查", "检查", "优化", "闷热", "通风", "空调", "负荷", "能耗")):
            pieces.append("场景：运维排查、运行优化")
        return "\n".join(piece for piece in pieces if piece)

    def _merge_knowledge_retrieval_results(self, question: str, *results: dict[str, Any], limit: int = 6) -> dict[str, Any]:
        scored: list[tuple[float, dict[str, Any]]] = []
        seen: set[tuple[str, str, str]] = set()
        retrieval_error_types: list[str] = []
        for result in results:
            if not isinstance(result, dict):
                continue
            error_type = str(result.get("retrieval_error_type", "")).strip()
            if error_type:
                retrieval_error_types.append(error_type)
            for item in list(result.get("items") or []):
                if not isinstance(item, dict):
                    continue
                key = (
                    str(item.get("chunk_id", "")).strip(),
                    str(item.get("title", "")).strip(),
                    str(item.get("excerpt", "")).strip(),
                )
                if key in seen:
                    continue
                seen.add(key)
                similarity = item.get("similarity")
                try:
                    similarity_value = float(similarity) if similarity is not None else None
                except (TypeError, ValueError):
                    similarity_value = None
                score = self._knowledge_relevance_score(
                    question,
                    str(item.get("title", "")),
                    str(item.get("excerpt", "")),
                    similarity_value,
                )
                scored.append((score, item))

        sorted_scored = sorted(scored, key=lambda pair: pair[0], reverse=True)
        items: list[dict[str, Any]] = []
        if sorted_scored:
            by_source: dict[str, list[dict[str, Any]]] = {}
            for _, item in sorted_scored:
                by_source.setdefault(str(item.get("source_type", "")), []).append(item)
            if "standard" in by_source and "ragflow" in by_source:
                for source_type in ("standard", "ragflow"):
                    for item in by_source.get(source_type, [])[: max(1, min(2, limit // 2 or 1))]:
                        if item not in items and len(items) < limit:
                            items.append(item)
            for _, item in sorted_scored:
                if len(items) >= limit:
                    break
                if item not in items:
                    items.append(item)
        if items:
            source_types = {str(item.get("source_type", "")) for item in items}
            if source_types == {"standard"}:
                knowledge_source = "standard"
            elif "standard" in source_types and "ragflow" in source_types:
                knowledge_source = "mixed"
            else:
                knowledge_source = "ragflow"
            return {
                "items": items,
                "knowledge_source": knowledge_source,
                "retrieval_hit_count": len(items),
                "retrieval_error_type": "",
                "ragflow_session_id": "",
            }

        return {
            "items": [],
            "knowledge_source": "none",
            "retrieval_hit_count": 0,
            "retrieval_error_type": retrieval_error_types[0] if retrieval_error_types else "empty_result",
            "ragflow_session_id": "",
        }

    def _llm_settings(self) -> dict[str, Any]:
        api_key = os.getenv("OPENAI_API_KEY", "").strip()
        base_url = os.getenv("OPENAI_BASE_URL", "https://api.deepseek.com").strip() or "https://api.deepseek.com"
        model = os.getenv("OPENAI_MODEL", "deepseek-chat").strip() or "deepseek-chat"
        try:
            timeout_sec = float(os.getenv("OPENAI_TIMEOUT_SEC", "20"))
        except ValueError:
            timeout_sec = 20.0
        return {
            "configured": bool(api_key),
            "api_key": api_key,
            "base_url": base_url,
            "model": model,
            "timeout_sec": timeout_sec,
        }

    def _knowledge_prompt_messages(self, question: str, references: list[dict[str, Any]], route: str) -> list[dict[str, str]]:
        evidence_lines: list[str] = []
        for idx, item in enumerate(references, start=1):
            title = str(item.get("title", "")).strip() or f"参考资料{idx}"
            source_type = "标准规范" if str(item.get("source_type", "")) == "standard" else "场景知识"
            section = str(item.get("section", "")).strip()
            excerpt = str(item.get("excerpt", "")).strip()
            evidence_lines.append(f"[{idx}] 来源类型：{source_type}")
            evidence_lines.append(f"[{idx}] 文档：{title}")
            if section:
                evidence_lines.append(f"[{idx}] 位置：{section}")
            evidence_lines.append(f"[{idx}] 摘录：{excerpt}")
        scope_hint = {
            "standard": "请直接输出自然中文答案，优先用2到4个完整短段落说明标准要点，不要写“结论”“依据与分析”等标题，不要编号。",
            "scene": "请直接输出自然中文答案，先概括核心判断，再分两到三段说明重点检查项和优化方向，不要写标题，不要数字编号。",
            "mixed": "请直接输出自然中文答案，先说明标准或原理，再补充运维做法，整体用2到4个完整短段落表达，不要写标题，不要数字编号。",
        }.get(route, "请结合召回资料直接回答。")
        system_prompt = (
            "你是建筑能源与运维知识助手。请严格基于给定资料作答，不要编造未提供的标准条文。"
            "如果资料不足，请明确说明“当前检索到的资料不足以确认”。"
            "引用资料时用[1][2]这种编号内联引用，正文后不要额外输出参考文献列表。"
            "回答使用中文，优先给出完整判断，再补充必要依据。"
            "不要输出 Markdown 表格，不要输出数字列表，不要输出“结论”“依据与分析”“标准依据”等小标题。"
            "每句话都要完整，必须使用自然中文标点，不要把多条内容挤在同一行。"
            "不要照抄零碎短句，不要输出残缺片段。若资料中出现数字范围、单位、条款编号或尺寸，请原样保留，不要改写或拆开。"
        )
        user_prompt = (
            f"用户问题：{question}\n"
            f"{scope_hint}\n\n"
            "已检索资料如下：\n"
            f"{chr(10).join(evidence_lines)}\n\n"
            "请输出可直接展示的中文答案。若资料支持多个判断点，请自然分段，不要使用1. 2. 3.编号。"
        )
        return [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ]

    def _parse_ragflow_stream_payload(self, raw_text: str) -> list[dict[str, Any]]:
        events: list[dict[str, Any]] = []
        for block in re.split(r"\r?\n\r?\n", str(raw_text or "")):
            if not block.strip():
                continue
            for line in block.splitlines():
                stripped = line.strip()
                if not stripped.startswith("data:"):
                    continue
                payload = stripped[5:].strip()
                if not payload:
                    continue
                try:
                    event = json.loads(payload)
                except json.JSONDecodeError:
                    continue
                if isinstance(event, dict):
                    events.append(event)
        return events

    def _ragflow_headers(self) -> dict[str, str]:
        settings = self._ragflow_settings()
        headers = {"Content-Type": "application/json"}
        if settings["api_key"]:
            headers["Authorization"] = f"Bearer {settings['api_key']}"
        return headers

    def _parse_ragflow_json_or_stream(self, raw_text: str) -> dict[str, Any]:
        text = str(raw_text or "").strip()
        if not text:
            return {}
        try:
            payload = json.loads(text)
        except json.JSONDecodeError:
            payload = None
        if isinstance(payload, dict):
            return payload

        events = self._parse_ragflow_stream_payload(text)
        if not events:
            raise ValueError("RAGFlow returned empty payload")

        for event in reversed(events):
            if isinstance(event.get("data"), dict):
                return event
        return events[-1]

    def _ragflow_request(
        self,
        path: str,
        *,
        method: str = "POST",
        payload: dict[str, Any] | None = None,
        timeout_sec: float | None = None,
    ) -> dict[str, Any]:
        settings = self._ragflow_settings()
        req = Request(
            f"{settings['base_url']}{path}",
            data=json.dumps(payload).encode("utf-8") if payload is not None else None,
            method=method,
            headers=self._ragflow_headers(),
        )
        with urlopen(req, timeout=timeout_sec or settings["timeout_sec"]) as resp:
            raw = resp.read().decode("utf-8")
        return self._parse_ragflow_json_or_stream(raw)

    def _ensure_ragflow_session(self, session_id: str | None = None) -> str:
        if session_id:
            return session_id
        settings = self._ragflow_settings()
        if not settings["chat_ready"]:
            raise RuntimeError("RAGFlow chat not configured")
        payload = {
            "name": f"A8-{dt.datetime.now().strftime('%m%d-%H%M%S')}",
        }
        result = self._ragflow_request(
            f"/chats/{settings['chat_id']}/sessions",
            method="POST",
            payload=payload,
        )
        data = result.get("data") if isinstance(result, dict) else None
        if not isinstance(data, dict) or not str(data.get("id", "")).strip():
            raise RuntimeError("RAGFlow session create failed")
        return str(data.get("id", "")).strip()

    def _normalize_ragflow_reference(self, reference: dict[str, Any] | None, limit: int = 6) -> list[dict[str, str]]:
        chunks = reference.get("chunks", []) if isinstance(reference, dict) else []
        settings = self._ragflow_settings()
        items: list[dict[str, str]] = []
        for idx, chunk in enumerate(chunks[:limit], start=1):
            if not isinstance(chunk, dict):
                continue
            similarity = chunk.get("similarity")
            similarity_value: float | None
            try:
                similarity_value = float(similarity) if similarity is not None else None
            except (TypeError, ValueError):
                similarity_value = None
            section = ""
            if similarity_value is not None:
                section = f"相似度 {similarity_value:.2f}"
            title = self._normalize_ragflow_title(chunk, "RAGFlow 文档")
            snippet_text = self._extract_ragflow_snippet(chunk)
            items.append(
                self._normalize_evidence_item(
                    chunk_id=chunk.get("id") or f"ragflow-chat-{idx}",
                    title=title,
                    section=section,
                    excerpt=snippet_text or chunk.get("content") or chunk.get("highlight") or "",
                    snippet_text=snippet_text,
                    document_key=self._ragflow_document_key(chunk, title),
                    source_type=self._ragflow_source_type_for_dataset(str(chunk.get("dataset_id") or ""), settings),
                    similarity=similarity_value,
                )
            )
        return items

    def _knowledge_source_from_references(self, references: list[dict[str, Any]] | None) -> str:
        source_types = {
            str(item.get("source_type", "")).strip()
            for item in (references or [])
            if isinstance(item, dict) and str(item.get("source_type", "")).strip()
        }
        if source_types == {"standard"}:
            return "standard"
        if "standard" in source_types and "ragflow" in source_types:
            return "mixed"
        return "ragflow"

    def _ragflow_reference_limit_for_answer(self, answer: Any, default: int = 6, cap: int = 12) -> int:
        text = str(answer or "")
        citation_ids = [
            int(match.group(1))
            for match in re.finditer(r"\[ID:\s*(\d+)\]", text, flags=re.IGNORECASE)
        ]
        if not citation_ids:
            return default
        return min(max(default, max(citation_ids) + 1), cap)

    def _ragflow_chat_completion(self, question: str, session_id: str | None = None) -> dict[str, Any]:
        settings = self._ragflow_settings()
        if not settings["chat_ready"]:
            raise RuntimeError("RAGFlow chat not configured")

        ensured_session_id = self._ensure_ragflow_session(session_id)
        started = time.perf_counter()
        body: dict[str, Any] = {
            "question": question,
            "stream": False,
            "session_id": ensured_session_id,
        }
        event = self._ragflow_request(
            f"/chats/{settings['chat_id']}/completions",
            method="POST",
            payload=body,
            timeout_sec=settings["chat_timeout_sec"],
        )
        latest = event.get("data") if isinstance(event.get("data"), dict) else event
        if not isinstance(latest, dict):
            raise RuntimeError("RAGFlow chat returned no payload")
        return {
            "answer": str(latest.get("answer", "") or ""),
            "reference": latest.get("reference") if isinstance(latest.get("reference"), dict) else {},
            "session_id": str(latest.get("session_id", "")).strip() or ensured_session_id,
            "message_id": str(latest.get("id", "")).strip(),
            "latency_ms": int((time.perf_counter() - started) * 1000),
        }

    def _ragflow_chat_stream_completion(self, question: str, session_id: str | None = None) -> dict[str, Any]:
        settings = self._ragflow_settings()
        if not settings["chat_ready"]:
            raise RuntimeError("RAGFlow chat not configured")

        ensured_session_id = self._ensure_ragflow_session(session_id)
        body: dict[str, Any] = {
            "question": question,
            "stream": True,
            "session_id": ensured_session_id,
        }
        req = Request(
            f"{settings['base_url']}/chats/{settings['chat_id']}/completions",
            data=json.dumps(body).encode("utf-8"),
            method="POST",
            headers=self._ragflow_headers(),
        )

        def event_iter():
            with urlopen(req, timeout=settings["chat_timeout_sec"]) as resp:
                for raw_line in resp:
                    line = raw_line.decode("utf-8", errors="ignore").strip()
                    if not line.startswith("data:"):
                        continue
                    payload = line[5:].strip()
                    if not payload or payload == "[DONE]":
                        continue
                    try:
                        event = json.loads(payload)
                    except json.JSONDecodeError:
                        continue
                    if not isinstance(event, dict):
                        continue
                    data = event.get("data")
                    if isinstance(data, dict):
                        yield data

        return {
            "session_id": ensured_session_id,
            "event_iter": event_iter(),
        }

    def _retrieve_ragflow_knowledge(
        self,
        query_text: str,
        limit: int = 3,
        *,
        dataset_ids: list[str] | None = None,
    ) -> dict[str, Any]:
        settings = self._ragflow_settings()
        active_dataset_ids = list(dataset_ids or settings["scene_dataset_ids"] or settings["dataset_ids"] or [])
        if not settings["configured"] or not active_dataset_ids:
            return {
                "items": [],
                "knowledge_source": "none",
                "retrieval_hit_count": 0,
                "retrieval_error_type": "not_configured",
            }

        body = json.dumps(
            {
                "dataset_ids": active_dataset_ids,
                "question": query_text,
                "top_k": max(limit * 4, settings["top_k"], 12),
                "similarity_threshold": settings["similarity_threshold"],
                "vector_similarity_weight": settings["vector_similarity_weight"],
                "highlight": False,
            }
        ).encode("utf-8")
        req = Request(
            f"{settings['base_url']}/retrieval",
            data=body,
            method="POST",
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {settings['api_key']}",
            },
        )
        try:
            with urlopen(req, timeout=settings["timeout_sec"]) as resp:
                payload = json.loads(resp.read().decode("utf-8"))
        except HTTPError as exc:
            payload = None
            retrieval_error_type = "http_error"
        except (URLError, TimeoutError, socket.timeout):
            payload = None
            retrieval_error_type = "timeout"
        except json.JSONDecodeError:
            payload = None
            retrieval_error_type = "parse_error"
        else:
            retrieval_error_type = ""

        if payload is not None and (not isinstance(payload, dict) or payload.get("code") not in (0, None)):
            retrieval_error_type = "http_error"
            payload = None

        scored_items: list[tuple[float, dict[str, Any]]] = []
        if payload is not None:
            data = payload.get("data") or {}
            chunks = data.get("chunks", []) if isinstance(data, dict) else []
            for idx, chunk in enumerate(chunks, start=1):
                if not isinstance(chunk, dict):
                    continue
                similarity = chunk.get("similarity")
                section = chunk.get("section") or chunk.get("important_keywords") or ""
                if not section and similarity is not None:
                    try:
                        section = f"相似度 {float(similarity):.2f}"
                    except (TypeError, ValueError):
                        section = ""
                excerpt = self._extract_ragflow_excerpt(chunk)
                if self._is_noisy_ragflow_chunk(chunk, excerpt):
                    continue
                title = self._normalize_ragflow_title(chunk)
                item = self._normalize_evidence_item(
                    chunk_id=chunk.get("id") or chunk.get("chunk_id") or f"ragflow-{idx}",
                    title=title,
                    section=section or chunk.get("dataset_id") or "",
                    excerpt=excerpt,
                    source_type=self._ragflow_source_type_for_dataset(str(chunk.get("dataset_id") or ""), settings),
                    similarity=float(similarity) if similarity is not None else None,
                )
                scored_items.append(
                    (
                        self._knowledge_relevance_score(query_text, title, excerpt, float(similarity) if similarity is not None else None),
                        item,
                    )
                )

        items = [item for _, item in sorted(scored_items, key=lambda pair: pair[0], reverse=True)[:limit]]

        if items:
            source_types = {str(item.get("source_type", "")) for item in items}
            if source_types == {"standard"}:
                knowledge_source = "standard"
            elif "standard" in source_types and "ragflow" in source_types:
                knowledge_source = "mixed"
            else:
                knowledge_source = "ragflow"
            return {
                "items": items,
                "knowledge_source": knowledge_source,
                "retrieval_hit_count": len(items),
                "retrieval_error_type": "",
                "ragflow_session_id": "",
            }

        return {
            "items": [],
            "knowledge_source": "none",
            "retrieval_hit_count": 0,
            "retrieval_error_type": retrieval_error_type or "empty_result",
            "ragflow_session_id": "",
        }

    def _answer_knowledge_question(
        self,
        question: str,
        *,
        limit: int = 6,
        stream: bool = False,
        session_id: str | None = None,
    ):
        llm_settings = self._llm_settings()
        if not llm_settings["configured"]:
            raise RuntimeError("DeepSeek not configured for knowledge answer generation")

        settings = self._ragflow_settings()
        route, dataset_ids = self._ragflow_dataset_route(question)
        retrieval_query = self._build_knowledge_retrieval_query(question, route)
        if route == "mixed" and settings.get("scene_dataset_ids") and settings.get("standard_dataset_ids"):
            retrieval = self._merge_knowledge_retrieval_results(
                retrieval_query,
                self._retrieve_ragflow_knowledge(
                    retrieval_query,
                    limit=max(limit, 4),
                    dataset_ids=list(settings.get("scene_dataset_ids") or []),
                ),
                self._retrieve_ragflow_knowledge(
                    retrieval_query,
                    limit=max(limit, 4),
                    dataset_ids=list(settings.get("standard_dataset_ids") or []),
                ),
                limit=limit,
            )
        else:
            retrieval = self._retrieve_ragflow_knowledge(retrieval_query, limit=limit, dataset_ids=dataset_ids)
        references = list(retrieval.get("items") or [])
        if not references:
            raise RuntimeError("未从知识库检索到可用资料")

        messages = self._knowledge_prompt_messages(question, references, route)
        llm_provider = self.providers.get("llm")
        if not isinstance(llm_provider, LLMDiagnoseProvider):
            raise RuntimeError("LLM provider unavailable")

        session_id = str(session_id or "").strip() or f"kb-{uuid.uuid4().hex[:12]}"
        started = time.perf_counter()
        if stream:
            token_iter = llm_provider._call_chat_completion_stream(
                base_url=llm_settings["base_url"],
                api_key=llm_settings["api_key"],
                model=llm_settings["model"],
                timeout_sec=llm_settings["timeout_sec"],
                messages=messages,
            )
            return {
                "mode": "stream",
                "route": route,
                "references": references,
                "session_id": session_id,
                "latency_ms": lambda: int((time.perf_counter() - started) * 1000),
                "token_iter": token_iter,
                "knowledge_source": retrieval.get("knowledge_source", "ragflow"),
            }

        response = llm_provider._call_chat_completion(
            base_url=llm_settings["base_url"],
            api_key=llm_settings["api_key"],
            model=llm_settings["model"],
            timeout_sec=llm_settings["timeout_sec"],
            messages=messages,
            max_tokens=900,
            response_format=None,
        )
        choices = response.get("choices", [])
        if not choices:
            raise RuntimeError("knowledge answer missing choices")
        content = choices[0].get("message", {}).get("content", "")
        if isinstance(content, list):
            content = "".join(
                str(part.get("text", "")) if isinstance(part, dict) else str(part)
                for part in content
            )
        return {
            "mode": "sync",
            "route": route,
            "answer": self._postprocess_knowledge_answer(content),
            "references": references,
            "session_id": session_id,
            "latency_ms": int((time.perf_counter() - started) * 1000),
            "knowledge_source": retrieval.get("knowledge_source", "ragflow"),
        }

    def ask_ragflow_chat(self, payload: dict[str, Any]) -> dict[str, Any]:
        question = str(payload.get("question", "")).strip()
        if not question:
            raise ValueError("question is required")

        _sid = payload.get("session_id")
        session_id = str(_sid).strip() if _sid and str(_sid).strip() not in ("None", "null") else None
        result = self._ragflow_chat_completion(question, session_id=session_id)
        references = self._normalize_ragflow_reference(
            result.get("reference"),
            limit=self._ragflow_reference_limit_for_answer(result.get("answer", "")),
        )
        return {
            "answer": str(result.get("answer", "")),
            "session_id": str(result.get("session_id", "")).strip(),
            "message_id": str(result.get("message_id", "")).strip(),
            "chat_id": self._ragflow_settings().get("chat_id", ""),
            "references": references,
            "knowledge_source": self._knowledge_source_from_references(references),
            "provider": "ragflow_chat",
            "latency_ms": result.get("latency_ms", 0),
        }

    def ragflow_chat_stream_events(self, payload: dict[str, Any]):
        """Yield (event_name, data_dict) tuples streaming RAGFlow chat completion."""
        question = str(payload.get("question", "")).strip()
        if not question:
            yield "error", {"message": "question is required"}
            return
        try:
            _sid = payload.get("session_id")
            session_id = str(_sid).strip() if _sid and str(_sid).strip() not in ("None", "null") else None
            stream_result = self._ragflow_chat_stream_completion(question, session_id=session_id)
        except Exception as exc:
            yield "error", {"message": str(exc)}
            return
        yield "start", {"session_id": stream_result["session_id"]}
        full_answer = ""
        latest_reference: dict[str, Any] = {}
        final_session_id = str(stream_result["session_id"]).strip()
        message_id = ""
        try:
            for event_data in stream_result["event_iter"]:
                final_session_id = str(event_data.get("session_id", "")).strip() or final_session_id
                message_id = str(event_data.get("id", "")).strip() or message_id
                reference = event_data.get("reference")
                if isinstance(reference, dict) and reference:
                    latest_reference = reference
                full_answer, delta = self._merge_ragflow_stream_text(full_answer, event_data.get("answer", ""))
                if delta:
                    yield "token", {"text": delta}
                if bool(event_data.get("final")):
                    references = self._normalize_ragflow_reference(
                        latest_reference,
                        limit=self._ragflow_reference_limit_for_answer(full_answer),
                    )
                    yield "done", {
                        "answer": full_answer,
                        "session_id": final_session_id,
                        "message_id": message_id,
                        "chat_id": self._ragflow_settings().get("chat_id", ""),
                        "references": references,
                        "knowledge_source": self._knowledge_source_from_references(references),
                    }
                    return
        except (URLError, TimeoutError, socket.timeout) as exc:
            yield "error", {"message": f"RAGFlow timeout: {exc}"}
            return
        except Exception as exc:
            yield "error", {"message": str(exc)}
            return
        references = self._normalize_ragflow_reference(
            latest_reference,
            limit=self._ragflow_reference_limit_for_answer(full_answer),
        )
        yield "done", {
            "answer": full_answer,
            "session_id": final_session_id,
            "message_id": message_id,
            "chat_id": self._ragflow_settings().get("chat_id", ""),
            "references": references,
            "knowledge_source": self._knowledge_source_from_references(references),
        }

    def _build_diagnose_knowledge_query(
        self,
        anomaly_name: str,
        message: str,
        context: dict[str, Any] | None,
    ) -> str:
        context = context or {}
        tags = "、".join([str(item).strip() for item in (context.get("phenomenon_tags") or [])[:4] if str(item).strip()])
        systems = "、".join([str(item).strip() for item in (context.get("likely_systems") or [])[:4] if str(item).strip()])
        parts = [
            f"异常类型：{anomaly_name}",
            f"建筑类型：{context.get('building_type') or ''}",
            f"建筑名称：{context.get('building_name') or ''}",
            f"发生时间：{context.get('timestamp') or ''}",
            f"偏差比例：{context.get('deviation_pct') or ''}",
            f"关键现象：{tags}",
            f"可能设备系统：{systems}",
            f"用户问题：{message or '请给出诊断与运维建议'}",
        ]
        return "\n".join([part for part in parts if str(part).strip()])

    def _build_analysis_knowledge_query(self, payload: dict[str, Any], context: dict[str, Any], insights: dict[str, Any]) -> str:
        summary = context.get("summary") or {}
        peer_group = (context.get("compare") or {}).get("peer_group") or {}
        insight_lines = []
        for item in (insights.get("trend_findings") or [])[:2]:
            if isinstance(item, dict):
                insight_lines.append(f"{item.get('title', '')}：{item.get('detail', '')}")
        for item in (insights.get("weather_findings") or [])[:2]:
            if isinstance(item, dict):
                insight_lines.append(f"{item.get('title', '')}：{item.get('detail', '')}")
        user_question = str(payload.get("message", "")).strip() or "请围绕当前建筑能耗分析给出节能与运维建议。"
        parts = [
            f"建筑名称：{context.get('building_name') or ''}",
            f"建筑类型：{context.get('building_type') or ''}",
            f"分析指标：{context.get('metric_type') or 'electricity'}",
            f"时间范围：{context.get('start_time') or ''} ~ {context.get('end_time') or ''}",
            f"均值：{summary.get('avg_value', 0)} {summary.get('unit', 'kWh')}",
            f"峰值：{summary.get('peak_value', 0)} {summary.get('unit', 'kWh')}",
            f"波动率：{summary.get('volatility_pct', 0)}%",
            f"同类差距：{peer_group.get('gap_pct', 0)}%",
            f"关键发现：{'；'.join(insight_lines)}",
            f"用户问题：{user_question}",
        ]
        return "\n".join([part for part in parts if str(part).strip()])

    def _prepare_indexes(self) -> None:
        for row in self.rows:
            self.by_building.setdefault(row["building_id"], []).append(row)

        for building_id, items in self.by_building.items():
            items.sort(key=lambda r: r["timestamp"])
            values = [r["electricity_kwh"] for r in items]
            mean_val = sum(values) / len(values)
            variance = sum((v - mean_val) ** 2 for v in values) / len(values)
            std_val = math.sqrt(variance)
            self.stats[building_id] = {
                "mean": mean_val,
                "std": std_val,
                "spike_threshold": mean_val + 2 * std_val,
                "high_load_threshold": mean_val * 1.5,
            }
            raw_meta = self.bdq2_metadata.get(building_id, {})
            peer_category = (
                self.showcase_config.get(building_id, {}).get("peer_category")
                or raw_meta.get("peer_category")
                or ""
            )
            display_category = (
                self.showcase_config.get(building_id, {}).get("display_category")
                or raw_meta.get("display_category")
                or items[0]["building_type"]
            )
            display_name = (
                raw_meta.get("display_name")
                or showcase_display_name(building_id, peer_category)
            )
            self.buildings_meta[building_id] = {
                "building_id": building_id,
                "building_name": display_name,
                "building_type": display_category,
                "display_name": display_name,
                "display_category": display_category,
                "peer_category": peer_category,
                "raw_building_name": items[0]["building_name"],
                "primaryspaceusage": raw_meta.get("primaryspaceusage", ""),
                "sub_primaryspaceusage": raw_meta.get("sub_primaryspaceusage", ""),
                "site_id": self.building_site_map.get(building_id),
                "record_count": len(items),
                "start_time": items[0]["timestamp"],
                "end_time": items[-1]["timestamp"],
            }

        self.anomalies = self._detect_anomalies()

    def _off_hours_row(self, row: dict[str, Any]) -> bool:
        hour = int(row.get("hour", 0))
        weekday = row["timestamp"].weekday()
        return weekday >= 5 or hour < 7 or hour >= 21

    def _work_hours_row(self, row: dict[str, Any]) -> bool:
        hour = int(row.get("hour", 0))
        weekday = row["timestamp"].weekday()
        return weekday < 5 and 8 <= hour <= 18

    def _median_value(self, values: list[float], fallback: float = 0.0) -> float:
        if not values:
            return fallback
        ordered = sorted(float(v) for v in values)
        middle = len(ordered) // 2
        if len(ordered) % 2 == 1:
            return ordered[middle]
        return (ordered[middle - 1] + ordered[middle]) / 2

    def _severity_for_anomaly(
        self,
        anomaly_type: str,
        row_value: float,
        baseline_value: float,
        threshold_value: float,
        trigger_count: int,
        extra: dict[str, Any] | None = None,
    ) -> str:
        extra = extra or {}
        baseline_gap_pct = ((row_value - baseline_value) / baseline_value * 100) if baseline_value else 0.0
        threshold_gap_pct = ((row_value - threshold_value) / threshold_value * 100) if threshold_value else 0.0
        threshold_drop_pct = ((threshold_value - row_value) / threshold_value * 100) if threshold_value else 0.0
        abs_gap_pct = abs(baseline_gap_pct)

        if anomaly_type == "anomaly_spike":
            if threshold_gap_pct >= 18 or abs_gap_pct >= 55:
                return "high"
            if threshold_gap_pct >= 8 or abs_gap_pct >= 30:
                return "medium"
            return "low"

        if anomaly_type == "anomaly_sustained_high_load":
            if trigger_count >= 8 or threshold_gap_pct >= 18:
                return "high"
            if trigger_count >= 5 or threshold_gap_pct >= 8:
                return "medium"
            return "low"

        if anomaly_type == "anomaly_off_hours_load":
            if trigger_count >= 6 or threshold_gap_pct >= 22:
                return "high"
            if trigger_count >= 3 or threshold_gap_pct >= 10:
                return "medium"
            return "low"

        if anomaly_type == "anomaly_workhour_offline":
            if trigger_count >= 4 or threshold_drop_pct >= 45:
                return "high"
            if trigger_count >= 2 or threshold_drop_pct >= 22:
                return "medium"
            return "low"

        if anomaly_type == "anomaly_baseload_high":
            night_ratio_pct = float(extra.get("night_ratio_pct", 0.0) or 0.0)
            if night_ratio_pct >= 85:
                return "high"
            if night_ratio_pct >= 65:
                return "medium"
            return "low"

        if anomaly_type == "anomaly_schedule_shift":
            shift_hours = abs(float(extra.get("shift_hours", 0.0) or 0.0))
            offhour_ratio = float(extra.get("offhour_ratio", 0.0) or 0.0)
            if shift_hours >= 4 and offhour_ratio >= 1.35:
                return "high"
            if shift_hours >= 2 and offhour_ratio >= 1.15:
                return "medium"
            return "low"

        if abs_gap_pct >= 60:
            return "high"
        if abs_gap_pct >= 30:
            return "medium"
        return "low"

    def _window_label(self, rows: list[dict[str, Any]]) -> str:
        if not rows:
            return ""
        if len(rows) == 1:
            return to_iso(rows[0]["timestamp"])
        return f"{to_iso(rows[0]['timestamp'])} ~ {to_iso(rows[-1]['timestamp'])}"

    def _append_anomaly(
        self,
        anomalies: list[dict[str, Any]],
        seen: set[tuple[int, str]],
        anomaly_id: int,
        row: dict[str, Any],
        anomaly_type: str,
        threshold: float,
        stat_mean: float,
        *,
        baseline_value: float | None = None,
        trigger_rows: list[dict[str, Any]] | None = None,
        trigger_count: int | None = None,
        time_scope_label: str | None = None,
        rule_summary: str | None = None,
        extra: dict[str, Any] | None = None,
    ) -> int:
        key = (int(row["record_id"]), anomaly_type)
        if key in seen:
            return anomaly_id
        seen.add(key)
        anomalies.append(
            self._to_anomaly(
                anomaly_id,
                row,
                anomaly_type,
                threshold,
                stat_mean,
                baseline_value=baseline_value,
                trigger_rows=trigger_rows,
                trigger_count=trigger_count,
                time_scope_label=time_scope_label,
                rule_summary=rule_summary,
                extra=extra,
            )
        )
        return anomaly_id + 1

    def _detect_anomalies(self) -> list[dict[str, Any]]:
        anomalies: list[dict[str, Any]] = []
        a_id = 1
        seen: set[tuple[int, str]] = set()

        for building_id, items in self.by_building.items():
            stat = self.stats[building_id]
            spike_threshold = stat["spike_threshold"]
            high_threshold = stat["high_load_threshold"]
            off_hour_values = [float(row["electricity_kwh"]) for row in items if self._off_hours_row(row)]
            work_hour_values = [float(row["electricity_kwh"]) for row in items if self._work_hours_row(row)]
            night_values = [float(row["electricity_kwh"]) for row in items if int(row.get("hour", 0)) < 6]
            off_hour_avg = sum(off_hour_values) / len(off_hour_values) if off_hour_values else stat["mean"]
            work_hour_avg = sum(work_hour_values) / len(work_hour_values) if work_hour_values else stat["mean"]
            night_avg = sum(night_values) / len(night_values) if night_values else off_hour_avg
            night_ratio = (night_avg / stat["mean"] * 100) if stat["mean"] else 0.0
            workhour_offline_threshold = min(max(work_hour_avg * 0.18, 0.25), max(work_hour_avg * 0.4, 0.5))
            off_hours_threshold = max(off_hour_avg * 1.45, stat["mean"] * 0.95)
            night_base_threshold = max(night_avg * 1.2, stat["mean"] * 0.52)
            schedule_active_threshold = max(work_hour_avg * 0.92, off_hour_avg * 1.22, stat["mean"] * 0.98)

            for index, row in enumerate(items):
                value = float(row["electricity_kwh"])
                prev_value = float(items[index - 1]["electricity_kwh"]) if index > 0 else value
                next_value = float(items[index + 1]["electricity_kwh"]) if index + 1 < len(items) else value
                if value > spike_threshold and value >= prev_value and value >= next_value:
                    a_id = self._append_anomaly(
                        anomalies,
                        seen,
                        a_id,
                        row,
                        "anomaly_spike",
                        spike_threshold,
                        stat["mean"],
                        baseline_value=stat["mean"],
                        trigger_rows=[row],
                        trigger_count=1,
                        time_scope_label="单点瞬时",
                        rule_summary=f"当前点负荷 {round(value, 2)} kWh，高于建筑瞬时突增阈值 {round(spike_threshold, 2)} kWh。",
                    )

            consec_rows: list[dict[str, Any]] = []
            for row in items:
                if float(row["electricity_kwh"]) > high_threshold:
                    consec_rows.append(row)
                else:
                    if len(consec_rows) >= 4:
                        peak_row = max(consec_rows, key=lambda current: float(current["electricity_kwh"]))
                        a_id = self._append_anomaly(
                            anomalies,
                            seen,
                            a_id,
                            peak_row,
                            "anomaly_sustained_high_load",
                            high_threshold,
                            stat["mean"],
                            baseline_value=stat["mean"],
                            trigger_rows=list(consec_rows),
                            trigger_count=len(consec_rows),
                            time_scope_label="连续高负荷窗口",
                            rule_summary=f"连续 {len(consec_rows)} 个时点高于高负荷阈值 {round(high_threshold, 2)} kWh。",
                        )
                    consec_rows = []
            if len(consec_rows) >= 4:
                peak_row = max(consec_rows, key=lambda current: float(current["electricity_kwh"]))
                a_id = self._append_anomaly(
                    anomalies,
                    seen,
                    a_id,
                    peak_row,
                    "anomaly_sustained_high_load",
                    high_threshold,
                    stat["mean"],
                    baseline_value=stat["mean"],
                    trigger_rows=list(consec_rows),
                    trigger_count=len(consec_rows),
                    time_scope_label="连续高负荷窗口",
                    rule_summary=f"连续 {len(consec_rows)} 个时点高于高负荷阈值 {round(high_threshold, 2)} kWh。",
                )

            off_hour_run: list[dict[str, Any]] = []
            workhour_offline_run: list[dict[str, Any]] = []
            for row in items:
                value = float(row["electricity_kwh"])
                if self._off_hours_row(row) and value > off_hours_threshold:
                    off_hour_run.append(row)
                else:
                    if len(off_hour_run) >= 2:
                        peak_row = max(off_hour_run, key=lambda current: float(current["electricity_kwh"]))
                        a_id = self._append_anomaly(
                            anomalies,
                            seen,
                            a_id,
                            peak_row,
                            "anomaly_off_hours_load",
                            off_hours_threshold,
                            stat["mean"],
                            baseline_value=off_hour_avg,
                            trigger_rows=list(off_hour_run),
                            trigger_count=len(off_hour_run),
                            time_scope_label="夜间/周末时段",
                            rule_summary=f"非工作时段连续 {len(off_hour_run)} 个时点高于基线阈值 {round(off_hours_threshold, 2)} kWh。",
                        )
                    off_hour_run = []

                if self._work_hours_row(row) and value < workhour_offline_threshold:
                    workhour_offline_run.append(row)
                else:
                    if len(workhour_offline_run) >= 2:
                        low_row = min(workhour_offline_run, key=lambda current: float(current["electricity_kwh"]))
                        a_id = self._append_anomaly(
                            anomalies,
                            seen,
                            a_id,
                            low_row,
                            "anomaly_workhour_offline",
                            workhour_offline_threshold,
                            stat["mean"],
                            baseline_value=work_hour_avg,
                            trigger_rows=list(workhour_offline_run),
                            trigger_count=len(workhour_offline_run),
                            time_scope_label="工作时段",
                            rule_summary=f"工作时段连续 {len(workhour_offline_run)} 个时点低于最低运行阈值 {round(workhour_offline_threshold, 2)} kWh。",
                        )
                    workhour_offline_run = []

            if len(off_hour_run) >= 2:
                peak_row = max(off_hour_run, key=lambda current: float(current["electricity_kwh"]))
                a_id = self._append_anomaly(
                    anomalies,
                    seen,
                    a_id,
                    peak_row,
                    "anomaly_off_hours_load",
                    off_hours_threshold,
                    stat["mean"],
                    baseline_value=off_hour_avg,
                    trigger_rows=list(off_hour_run),
                    trigger_count=len(off_hour_run),
                    time_scope_label="夜间/周末时段",
                    rule_summary=f"非工作时段连续 {len(off_hour_run)} 个时点高于基线阈值 {round(off_hours_threshold, 2)} kWh。",
                )
            if len(workhour_offline_run) >= 2:
                low_row = min(workhour_offline_run, key=lambda current: float(current["electricity_kwh"]))
                a_id = self._append_anomaly(
                    anomalies,
                    seen,
                    a_id,
                    low_row,
                    "anomaly_workhour_offline",
                    workhour_offline_threshold,
                    stat["mean"],
                    baseline_value=work_hour_avg,
                    trigger_rows=list(workhour_offline_run),
                    trigger_count=len(workhour_offline_run),
                    time_scope_label="工作时段",
                    rule_summary=f"工作时段连续 {len(workhour_offline_run)} 个时点低于最低运行阈值 {round(workhour_offline_threshold, 2)} kWh。",
                )

            night_rows_by_day: dict[str, list[dict[str, Any]]] = {}
            day_rows: dict[str, list[dict[str, Any]]] = {}
            for row in items:
                date_key = row["timestamp"].strftime("%Y-%m-%d")
                day_rows.setdefault(date_key, []).append(row)
                if int(row.get("hour", 0)) < 6:
                    night_rows_by_day.setdefault(date_key, []).append(row)
            for date_key, rows_in_day in night_rows_by_day.items():
                values = [float(row["electricity_kwh"]) for row in rows_in_day]
                if not values:
                    continue
                day_avg = sum(float(row["electricity_kwh"]) for row in day_rows.get(date_key, rows_in_day)) / max(len(day_rows.get(date_key, rows_in_day)), 1)
                local_night_avg = sum(values) / len(values)
                ratio_pct = (local_night_avg / day_avg * 100) if day_avg else 0.0
                if ratio_pct > 58 and local_night_avg > night_base_threshold:
                    anchor_row = max(rows_in_day, key=lambda current: float(current["electricity_kwh"]))
                    a_id = self._append_anomaly(
                        anomalies,
                        seen,
                        a_id,
                        anchor_row,
                        "anomaly_baseload_high",
                        night_base_threshold,
                        stat["mean"],
                        baseline_value=day_avg,
                        trigger_rows=list(rows_in_day),
                        trigger_count=len(rows_in_day),
                        time_scope_label="夜间基线窗口",
                        rule_summary=f"{date_key} 夜间平均负荷为日均的 {round(ratio_pct, 1)}%，高于夜间基线判定阈值。",
                        extra={"night_ratio_pct": round(ratio_pct, 2)},
                    )

            rows_by_day: dict[str, list[dict[str, Any]]] = {}
            for row in items:
                if row["timestamp"].weekday() < 5:
                    rows_by_day.setdefault(row["timestamp"].strftime("%Y-%m-%d"), []).append(row)
            first_active_hours: list[float] = []
            last_active_hours: list[float] = []
            daily_active_rows: dict[str, list[dict[str, Any]]] = {}
            for date_key, rows_in_day in rows_by_day.items():
                active_rows = [row for row in rows_in_day if float(row["electricity_kwh"]) >= schedule_active_threshold]
                if len(active_rows) < 2:
                    continue
                daily_active_rows[date_key] = active_rows
                first_active = min(active_rows, key=lambda current: current["timestamp"])
                last_active = max(active_rows, key=lambda current: current["timestamp"])
                first_active_hours.append(float(first_active.get("hour", 0)))
                last_active_hours.append(float(last_active.get("hour", 0)))

            typical_start_hour = int(round(self._median_value(first_active_hours, 8.0)))
            typical_end_hour = int(round(self._median_value(last_active_hours, 18.0)))

            for date_key, active_rows in daily_active_rows.items():
                first_active = min(active_rows, key=lambda current: current["timestamp"])
                last_active = max(active_rows, key=lambda current: current["timestamp"])
                first_hour = int(first_active.get("hour", 0))
                last_hour = int(last_active.get("hour", 0))
                first_value = float(first_active["electricity_kwh"])
                last_value = float(last_active["electricity_kwh"])
                early_shift_hours = typical_start_hour - first_hour
                late_shift_hours = last_hour - typical_end_hour
                early_ratio = (first_value / off_hour_avg) if off_hour_avg else 0.0
                late_ratio = (last_value / off_hour_avg) if off_hour_avg else 0.0

                if early_shift_hours >= 2 and first_value >= schedule_active_threshold * 1.03 and early_ratio >= 1.12:
                    a_id = self._append_anomaly(
                        anomalies,
                        seen,
                        a_id,
                        first_active,
                        "anomaly_schedule_shift",
                        schedule_active_threshold,
                        stat["mean"],
                        baseline_value=work_hour_avg,
                        trigger_rows=[first_active],
                        trigger_count=1,
                        time_scope_label="日程启停窗口",
                        rule_summary=f"{date_key} 首个高负荷时点出现在 {first_hour:02d}:00，较典型启用时段 {typical_start_hour:02d}:00 提前了 {early_shift_hours} 小时。",
                        extra={
                            "schedule_direction": "early_start",
                            "shift_hours": early_shift_hours,
                            "typical_start_hour": typical_start_hour,
                            "offhour_ratio": round(early_ratio, 4),
                        },
                    )
                if late_shift_hours >= 2 and last_value >= schedule_active_threshold * 1.03 and late_ratio >= 1.12:
                    a_id = self._append_anomaly(
                        anomalies,
                        seen,
                        a_id,
                        last_active,
                        "anomaly_schedule_shift",
                        schedule_active_threshold,
                        stat["mean"],
                        baseline_value=work_hour_avg,
                        trigger_rows=[last_active],
                        trigger_count=1,
                        time_scope_label="日程启停窗口",
                        rule_summary=f"{date_key} 最后一个高负荷时点出现在 {last_hour:02d}:00，较典型停用时段 {typical_end_hour:02d}:00 延后了 {late_shift_hours} 小时。",
                        extra={
                            "schedule_direction": "late_stop",
                            "shift_hours": late_shift_hours,
                            "typical_end_hour": typical_end_hour,
                            "offhour_ratio": round(late_ratio, 4),
                        },
                    )

        anomalies.sort(key=lambda x: x["timestamp"], reverse=True)
        return anomalies

    def _to_anomaly(
        self,
        anomaly_id: int,
        row: dict[str, Any],
        anomaly_type: str,
        threshold: float,
        stat_mean: float,
        *,
        baseline_value: float | None = None,
        trigger_rows: list[dict[str, Any]] | None = None,
        trigger_count: int | None = None,
        time_scope_label: str | None = None,
        rule_summary: str | None = None,
        extra: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        reference_value = float(baseline_value if baseline_value is not None else stat_mean)
        diff_pct = ((row["electricity_kwh"] - reference_value) / reference_value * 100) if reference_value else 0.0
        trigger_rows = trigger_rows or [row]
        rule_meta = self._rule_meta(anomaly_type)
        merged_extra = dict(extra or {})
        severity = self._severity_for_anomaly(
            anomaly_type,
            float(row["electricity_kwh"]),
            reference_value,
            float(threshold),
            int(trigger_count or len(trigger_rows) or 1),
            merged_extra,
        )
        return {
            "anomaly_id": anomaly_id,
            "record_id": row["record_id"],
            "building_id": row["building_id"],
            "building_name": row["building_name"],
            "building_type": row["building_type"],
            "timestamp": row["timestamp"],
            "electricity_kwh": round(row["electricity_kwh"], 4),
            "mean_kwh": round(stat_mean, 4),
            "baseline_value": round(reference_value, 4),
            "current_value": round(float(row["electricity_kwh"]), 4),
            "deviation_pct": round(diff_pct, 2),
            "threshold": round(threshold, 4),
            "threshold_value": round(threshold, 4),
            "anomaly_type": anomaly_type,
            "anomaly_name": self._anomaly_name(anomaly_type),
            "display_name": self._anomaly_name(anomaly_type),
            "rule_code": rule_meta["rule_code"],
            "rule_name": rule_meta["rule_name"],
            "rule_summary": str(rule_summary or rule_meta["summary_template"]).strip(),
            "trigger_window": self._window_label(trigger_rows),
            "trigger_count": int(trigger_count or len(trigger_rows) or 1),
            "time_scope_label": str(time_scope_label or rule_meta["time_scope_label"]).strip(),
            "severity": severity,
            **merged_extra,
        }

    def _load_actions(self) -> None:
        self.action_events = []
        with self.action_log_file.open("r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    event = json.loads(line)
                except json.JSONDecodeError:
                    continue
                if not isinstance(event, dict):
                    continue
                try:
                    event["anomaly_id"] = int(event.get("anomaly_id"))
                except (ValueError, TypeError):
                    continue
                event["created_at"] = str(event.get("created_at", ""))
                self.action_events.append(event)
        self._rebuild_action_index()

    def _rebuild_action_index(self) -> None:
        self.action_index = {}
        for e in sorted(self.action_events, key=lambda x: x.get("created_at", "")):
            aid = int(e["anomaly_id"])
            bucket = self.action_index.setdefault(
                aid,
                {
                    "status": STATUS_NEW,
                    "assignee": "",
                    "note": "",
                    "last_action_at": "",
                    "history": [],
                },
            )
            bucket["status"] = str(e.get("status", STATUS_NEW))
            bucket["assignee"] = str(e.get("assignee", ""))
            bucket["note"] = str(e.get("note", ""))
            bucket["last_action_at"] = str(e.get("created_at", ""))
            bucket["history"].append(
                {
                    "action": str(e.get("action", "")),
                    "status": str(e.get("status", STATUS_NEW)),
                    "assignee": str(e.get("assignee", "")),
                    "note": str(e.get("note", "")),
                    "created_at": str(e.get("created_at", "")),
                }
            )

    def _append_action_event(self, event: dict[str, Any]) -> None:
        with self.action_log_file.open("a", encoding="utf-8") as f:
            f.write(json.dumps(event, ensure_ascii=False) + "\n")
        self.action_events.append(event)
        self._rebuild_action_index()

    def _load_notes(self) -> None:
        self.note_events = []
        with self.note_log_file.open("r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    event = json.loads(line)
                except json.JSONDecodeError:
                    continue
                if not isinstance(event, dict):
                    continue
                try:
                    event["anomaly_id"] = int(event.get("anomaly_id"))
                except (ValueError, TypeError):
                    continue
                event["updated_at"] = str(event.get("updated_at", ""))
                self.note_events.append(event)
        self._rebuild_note_index()

    def _rebuild_note_index(self) -> None:
        self.note_index = {}
        for e in sorted(self.note_events, key=lambda x: x.get("updated_at", "")):
            self.note_index[int(e["anomaly_id"])] = {
                "anomaly_id": int(e["anomaly_id"]),
                "cause_confirmed": str(e.get("cause_confirmed", "")),
                "action_taken": str(e.get("action_taken", "")),
                "result_summary": str(e.get("result_summary", "")),
                "recurrence_risk": str(e.get("recurrence_risk", "")),
                "reviewer": str(e.get("reviewer", "")),
                "updated_at": str(e.get("updated_at", "")),
            }

    def _append_note_event(self, event: dict[str, Any]) -> None:
        with self.note_log_file.open("a", encoding="utf-8") as f:
            f.write(json.dumps(event, ensure_ascii=False) + "\n")
        self.note_events.append(event)
        self._rebuild_note_index()

    def _load_ai_events(self) -> None:
        self.ai_events = []
        with self.ai_call_log_file.open("r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    event = json.loads(line)
                except json.JSONDecodeError:
                    continue
                if isinstance(event, dict):
                    self.ai_events.append(event)

    def _append_ai_event(self, event: dict[str, Any]) -> None:
        with self.ai_call_log_file.open("a", encoding="utf-8") as f:
            f.write(json.dumps(event, ensure_ascii=False) + "\n")
        self.ai_events.append(event)

    def query_ai_stats(self, hours: int = 24) -> dict[str, Any]:
        safe_hours = min(max(hours, 1), 168)
        now = dt.datetime.now()
        start = now - dt.timedelta(hours=safe_hours)
        window_events = []
        for ev in self.ai_events:
            ts = parse_time(str(ev.get("timestamp", "")))
            if ts and ts >= start:
                window_events.append(ev)

        total = len(window_events)
        llm_calls = sum(1 for ev in window_events if str(ev.get("requested_provider", "")).lower() in {"llm", "auto"})
        fallback_calls = sum(1 for ev in window_events if bool(ev.get("fallback_used", False)))
        latencies = [int(ev.get("latency_ms", 0)) for ev in window_events if str(ev.get("latency_ms", "")).isdigit()]
        avg_latency = round(sum(latencies) / len(latencies), 2) if latencies else 0.0

        by_provider: dict[str, int] = {}
        error_types: dict[str, int] = {}
        knowledge_sources: dict[str, int] = {}
        for ev in window_events:
            provider = str(ev.get("provider", "unknown"))
            by_provider[provider] = by_provider.get(provider, 0) + 1
            err = str(ev.get("error_type", "")).strip()
            if err:
                error_types[err] = error_types.get(err, 0) + 1
            knowledge = str(ev.get("knowledge_source", "")).strip() or "none"
            knowledge_sources[knowledge] = knowledge_sources.get(knowledge, 0) + 1

        fallback_rate = round((fallback_calls / total) * 100, 2) if total else 0.0
        return {
            "window_hours": safe_hours,
            "total_calls": total,
            "llm_calls": llm_calls,
            "fallback_calls": fallback_calls,
            "fallback_rate_pct": fallback_rate,
            "avg_latency_ms": avg_latency,
            "by_provider": by_provider,
            "knowledge_sources": knowledge_sources,
            "error_types": error_types,
            "updated_at": to_iso(now),
        }

    def _required_diag_fields_complete(self, diagnosis: dict[str, Any]) -> bool:
        required = ["conclusion", "causes", "steps", "prevention", "recommended_actions", "evidence", "confidence", "risk_level"]
        for key in required:
            if key not in diagnosis:
                return False
            value = diagnosis.get(key)
            if key in {"causes", "steps", "prevention", "recommended_actions", "evidence"}:
                if not isinstance(value, list) or not value:
                    return False
            elif value in (None, ""):
                return False
        return True

    def _required_analysis_fields_complete(self, analysis: dict[str, Any]) -> bool:
        required = [
            "summary",
            "findings",
            "possible_causes",
            "energy_saving_suggestions",
            "operations_suggestions",
            "evidence",
        ]
        for key in required:
            if key not in analysis:
                return False
            value = analysis.get(key)
            if key == "summary":
                if value in (None, ""):
                    return False
            elif not isinstance(value, list) or not value:
                return False
        return True

    def upsert_anomaly_note(self, payload: dict[str, Any]) -> dict[str, Any]:
        anomaly_id_raw = payload.get("anomaly_id")
        if anomaly_id_raw is None:
            raise ValueError("anomaly_id required")
        try:
            anomaly_id = int(anomaly_id_raw)
        except (TypeError, ValueError):
            raise ValueError("invalid anomaly_id")

        if not any(a["anomaly_id"] == anomaly_id for a in self.anomalies):
            raise LookupError("anomaly not found")

        cause_confirmed = str(payload.get("cause_confirmed", "")).strip()
        action_taken = str(payload.get("action_taken", "")).strip()
        result_summary = str(payload.get("result_summary", "")).strip()
        recurrence_risk = str(payload.get("recurrence_risk", "")).strip().lower()
        reviewer = str(payload.get("reviewer", "")).strip()

        if not cause_confirmed or not action_taken or not result_summary:
            raise ValueError("cause_confirmed/action_taken/result_summary required")
        if recurrence_risk not in {"low", "medium", "high"}:
            recurrence_risk = "medium"

        event = {
            "anomaly_id": anomaly_id,
            "cause_confirmed": cause_confirmed,
            "action_taken": action_taken,
            "result_summary": result_summary,
            "recurrence_risk": recurrence_risk,
            "reviewer": reviewer,
            "updated_at": to_iso(dt.datetime.now()),
        }
        self._append_note_event(event)
        return dict(self.note_index.get(anomaly_id, event))

    def query_anomaly_note(self, anomaly_id: int) -> dict[str, Any]:
        note = self.note_index.get(anomaly_id)
        if note:
            return dict(note)
        return {
            "anomaly_id": anomaly_id,
            "cause_confirmed": "",
            "action_taken": "",
            "result_summary": "",
            "recurrence_risk": "medium",
            "reviewer": "",
            "updated_at": "",
        }

    def _compute_processing_duration_hours(self, anomaly_id: int) -> float:
        state = self.action_index.get(anomaly_id)
        if not state or not state.get("history"):
            return 0.0
        history = state["history"]
        first_ts = parse_time(str(history[0].get("created_at", "")))
        resolved_ts = None
        for h in history:
            if str(h.get("status")) == STATUS_RESOLVED:
                resolved_ts = parse_time(str(h.get("created_at", "")))
                break
        end_ts = resolved_ts or parse_time(str(history[-1].get("created_at", "")))
        if not first_ts or not end_ts:
            return 0.0
        return round(max((end_ts - first_ts).total_seconds(), 0.0) / 3600.0, 2)

    def export_anomalies_csv(
        self,
        building_id: str | None,
        start_time: dt.datetime | None,
        end_time: dt.datetime | None,
        anomaly_type: str | None,
        severity: str | None,
        status: str | None,
        sort: str = "timestamp_desc",
    ) -> str:
        rows = self._sort_anomalies(self._filter_anomalies(building_id, start_time, end_time, anomaly_type, severity, status), sort)
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(
            [
                "anomaly_id",
                "building_id",
                "building_name",
                "anomaly_type",
                "severity",
                "status",
                "assignee",
                "timestamp",
                "deviation_pct",
                "estimated_loss_kwh",
                "processing_duration_hours",
                "final_conclusion",
            ]
        )
        for item in rows:
            aid = int(item["anomaly_id"])
            action_state = self._action_state(aid)
            estimated_loss = round(max(float(item.get("electricity_kwh", 0)) - float(item.get("mean_kwh", 0)), 0.0), 4)
            note = self.note_index.get(aid, {})
            conclusion = str(note.get("result_summary", "")).strip() or str(item.get("anomaly_type", ""))
            writer.writerow(
                [
                    aid,
                    item["building_id"],
                    item["building_name"],
                    item["anomaly_type"],
                    item["severity"],
                    action_state["status"],
                    action_state["assignee"],
                    to_iso(item["timestamp"]),
                    item["deviation_pct"],
                    estimated_loss,
                    self._compute_processing_duration_hours(aid),
                    conclusion,
                ]
            )
        return output.getvalue()

    def _assistant_report_meta(self, module: str) -> dict[str, str]:
        meta = ASSISTANT_REPORT_MODULE_META.get(module, {})
        return {
            "title": str(meta.get("title", "")).strip() or "智能助手报告",
            "filename": str(meta.get("filename", "")).strip() or "智能助手报告",
        }

    def _normalize_report_context_meta(self, payload: dict[str, Any]) -> dict[str, str]:
        meta = payload.get("context_meta", {})
        if not isinstance(meta, dict):
            meta = {}
        building_id = clean_text(meta.get("building_id")) or "ALL"
        building_name = clean_text(meta.get("building_name")) or showcase_display_name(building_id)
        start_time = clean_text(meta.get("start_time")) or "-"
        end_time = clean_text(meta.get("end_time")) or "-"
        return {
            "building_id": building_id,
            "building_name": building_name,
            "start_time": start_time,
            "end_time": end_time,
            "scope_text": clean_text(meta.get("scope_text")) or f"{building_name} | {start_time} ~ {end_time}",
            "anomaly_id": clean_text(meta.get("anomaly_id")) or "-",
            "anomaly_name": clean_text(meta.get("anomaly_name")) or "-",
            "timestamp": clean_text(meta.get("timestamp")) or "-",
            "metric_label": clean_text(meta.get("metric_label")) or "电力",
            "building_type": clean_text(meta.get("building_type")) or "-",
        }

    def _normalize_report_messages(self, payload: dict[str, Any], module: str) -> list[dict[str, Any]]:
        raw_messages = payload.get("session_messages", [])
        if not isinstance(raw_messages, list):
            raw_messages = []
        module_type = "analysis" if module == "interpretation" else module
        messages: list[dict[str, Any]] = []
        for raw in raw_messages:
            if not isinstance(raw, dict):
                continue
            role = str(raw.get("role", "")).strip()
            if role == "user":
                content = clean_text(raw.get("content"))
                if content:
                    messages.append({"role": "user", "content": content})
                continue
            if role != "assistant":
                continue
            msg_type = str(raw.get("type", "")).strip()
            if msg_type != module_type or bool(raw.get("pending")):
                continue
            data = decode_json_clone(raw.get("data")) if isinstance(raw.get("data"), dict) else {}
            content = clean_text(raw.get("content"))
            if not data and not content:
                continue
            messages.append({"role": "assistant", "type": msg_type, "data": data, "content": content})
        return messages

    def _coerce_report_latest_result(self, payload: dict[str, Any], messages: list[dict[str, Any]]) -> dict[str, Any]:
        latest = payload.get("latest_result", {})
        if isinstance(latest, dict) and latest:
            return decode_json_clone(latest)
        for item in reversed(messages):
            if item.get("role") == "assistant" and isinstance(item.get("data"), dict) and item["data"]:
                return decode_json_clone(item["data"])
        return {}

    def _assistant_report_question_summary(self, messages: list[dict[str, Any]]) -> list[str]:
        questions = [clean_text(item.get("content")) for item in messages if item.get("role") == "user"]
        compact: list[str] = []
        for item in questions:
            if item and item not in compact:
                compact.append(item)
        return compact

    def _assistant_report_brief_history(self, messages: list[dict[str, Any]], module: str) -> list[str]:
        history: list[str] = []
        for item in messages:
            if item.get("role") != "assistant":
                continue
            data = item.get("data") if isinstance(item.get("data"), dict) else {}
            if module == "diagnosis":
                conclusion = clean_text(data.get("conclusion"))
            else:
                conclusion = clean_text(data.get("summary"))
            if conclusion:
                history.append(conclusion)
        return history[:-1] if len(history) > 1 else []

    def _assistant_report_sections(self, module: str, latest_result: dict[str, Any], history: list[str]) -> list[dict[str, Any]]:
        sections: list[dict[str, Any]] = []
        if module == "saving":
            sections.append({"title": "1. 节能结论摘要", "type": "paragraph", "content": clean_text(latest_result.get("summary")) or "当前会话未返回节能结论。"})
            sections.append({"title": "2. 优先动作", "type": "list", "items": ensure_text_list(latest_result.get("energy_saving_suggestions"))})
            sections.append({"title": "3. 预估收益与影响", "type": "list", "items": ensure_text_list(latest_result.get("report_impacts", latest_result.get("saving_impacts")))})
            sections.append({"title": "4. 实施配合要点", "type": "list", "items": ensure_text_list(latest_result.get("operations_suggestions"))})
        elif module == "diagnosis":
            sections.append({"title": "1. 诊断结论", "type": "paragraph", "content": clean_text(latest_result.get("conclusion")) or "当前会话未返回诊断结论。"})
            sections.append({"title": "2. 可能原因", "type": "list", "items": ensure_text_list(latest_result.get("causes", latest_result.get("possible_causes")))})
            sections.append({"title": "3. 排查步骤", "type": "list", "items": ensure_text_list(latest_result.get("steps"))})
            sections.append({"title": "4. 立即动作", "type": "list", "items": ensure_text_list(latest_result.get("recommended_actions"))})
            sections.append({"title": "5. 预防建议", "type": "list", "items": ensure_text_list(latest_result.get("prevention"))})
        else:
            sections.append({"title": "1. 分析结论", "type": "paragraph", "content": clean_text(latest_result.get("summary")) or "当前会话未返回分析结论。"})
            sections.append({"title": "2. 主要发现", "type": "list", "items": ensure_text_list(latest_result.get("findings"))})
            sections.append({"title": "3. 可能原因", "type": "list", "items": ensure_text_list(latest_result.get("possible_causes"))})
            sections.append({"title": "4. 运维建议", "type": "list", "items": ensure_text_list(latest_result.get("operations_suggestions"))})

        if history:
            sections.append({"title": "会话补充结论", "type": "list", "items": history})
        return sections

    def _build_operator_form_rows(self, module: str, latest_result: dict[str, Any]) -> list[dict[str, Any]]:
        if module == "diagnosis":
            action_items = ensure_text_list(latest_result.get("recommended_actions"))[:2]
            review_items = ensure_text_list(latest_result.get("prevention"))[:2]
            review_summary = clean_text(latest_result.get("conclusion"))
        elif module == "saving":
            action_items = ensure_text_list(latest_result.get("energy_saving_suggestions"))[:2]
            review_items = ensure_text_list(latest_result.get("operations_suggestions"))[:2]
            review_summary = clean_text(latest_result.get("summary"))
        else:
            action_items = ensure_text_list(latest_result.get("operations_suggestions"))[:2]
            review_items = ensure_text_list(latest_result.get("possible_causes"))[:2]
            review_summary = clean_text(latest_result.get("summary"))

        action_text = "；".join(action_items) if action_items else "结合本次报告建议填写执行动作。"
        review_text = "；".join(review_items) if review_items else "结合本次报告结论填写复核意见。"
        if review_summary:
            review_text = f"{review_summary}；{review_text}"

        return [
            {"label": "处理人", "value": "张三 / 李四 / 王五（任选其一）", "row_height_lines": 1},
            {"label": "处理时间", "value": "____年__月__日 __:__", "row_height_lines": 1},
            {"label": "执行结果", "value": action_text, "row_height_lines": 2},
            {"label": "复核结论", "value": review_text, "row_height_lines": 2},
        ]

    def _override_operator_form_rows(
        self,
        default_rows: list[dict[str, Any]],
        payload: dict[str, Any],
    ) -> list[dict[str, Any]]:
        raw = payload.get("operator_form", {})
        if not isinstance(raw, dict):
            return default_rows
        label_to_row = {str(item.get("label", "")).strip(): dict(item) for item in default_rows}
        mapping = {
            "assignee": "处理人",
            "processing_time": "处理时间",
            "execution_result": "执行结果",
            "review_conclusion": "复核结论",
        }
        for field_key, label in mapping.items():
            if label not in label_to_row:
                continue
            value = clean_text(raw.get(field_key))
            if value:
                label_to_row[label]["value"] = value
        return [label_to_row.get(item.get("label", ""), item) for item in default_rows]

    def _build_assistant_report(self, payload: dict[str, Any]) -> dict[str, Any]:
        module = str(payload.get("module", "")).strip().lower()
        if module not in ASSISTANT_REPORT_MODULE_META:
            raise ValueError("module must be saving/diagnosis/interpretation")
        format_name = str(payload.get("format", "")).strip().lower()
        if format_name not in {"docx", "pdf"}:
            raise ValueError("format must be docx/pdf")
        messages = self._normalize_report_messages(payload, module)
        if not messages:
            raise ValueError("current module session is empty")
        latest_result = self._coerce_report_latest_result(payload, messages)
        if not latest_result:
            raise ValueError("latest_result required")
        meta = self._normalize_report_context_meta(payload)
        questions = self._assistant_report_question_summary(messages)
        history = self._assistant_report_brief_history(messages, module)
        module_meta = self._assistant_report_meta(module)
        meta_items = [
            ("生成时间", to_iso(dt.datetime.now())),
            ("报告模块", module_meta["title"].replace("报告", "")),
            ("建筑名称", meta["building_name"]),
            ("建筑编号", meta["building_id"]),
            ("建筑类型", meta["building_type"]),
            ("分析指标", meta["metric_label"]),
            ("时间范围", f"{meta['start_time']} ~ {meta['end_time']}"),
        ]
        if module == "diagnosis":
            meta_items.extend(
                [
                    ("异常类型", meta["anomaly_name"]),
                    ("异常时间", meta["timestamp"]),
                    ("异常编号", meta["anomaly_id"]),
                ]
            )
        operator_form_rows = self._override_operator_form_rows(
            self._build_operator_form_rows(module, latest_result),
            payload,
        )
        return {
            "module": module,
            "format": format_name,
            "report_title": module_meta["title"],
            "filename_prefix": module_meta["filename"],
            "meta_items": meta_items,
            "question_summary": questions,
            "sections": self._assistant_report_sections(module, latest_result, history),
            "operator_form_fields": operator_form_rows,
            "latest_result": latest_result,
        }

    def _set_doc_cell_text(self, cell: Any, text: str, bold: bool = False) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.bold = bold
        self._set_doc_run_font(run, 10.5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _set_doc_run_font(self, run: Any, size: float, bold: bool | None = None) -> None:
        if bold is not None:
            run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Microsoft YaHei"
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    def _docx_add_title(self, document: Document, text: str, level: int = 0) -> None:
        if level == 0:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            self._set_doc_run_font(run, 16, bold=True)
            return
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        self._set_doc_run_font(run, 12 if level == 1 else 10.5, bold=True)

    def _docx_add_paragraph(self, document: Document, text: str, indent: bool = False) -> None:
        paragraph = document.add_paragraph()
        if indent:
            paragraph.paragraph_format.first_line_indent = Cm(0.74)
        run = paragraph.add_run(text)
        self._set_doc_run_font(run, 10.5)

    def _docx_add_list(self, document: Document, items: list[str]) -> None:
        if not items:
            self._docx_add_paragraph(document, "无")
            return
        for idx, item in enumerate(items, start=1):
            self._docx_add_paragraph(document, f"{idx}. {item}")

    def _apply_doc_table_border(self, table: Any) -> None:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        borders = tbl_pr.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = borders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "B8C0CC")

    def _build_docx_report(self, report: dict[str, Any]) -> bytes:
        document = Document()
        section = document.sections[0]
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        self._docx_add_title(document, report["report_title"])
        self._docx_add_paragraph(document, f"导出时间：{to_iso(dt.datetime.now())}")

        meta_table = document.add_table(rows=0, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_table.autofit = True
        for label, value in report["meta_items"]:
            row = meta_table.add_row().cells
            self._set_doc_cell_text(row[0], label, bold=True)
            self._set_doc_cell_text(row[1], clean_text(value) or "-")
        self._apply_doc_table_border(meta_table)

        self._docx_add_title(document, "会话问题摘要", level=1)
        self._docx_add_list(document, report["question_summary"])

        for section_item in report["sections"]:
            self._docx_add_title(document, section_item["title"], level=1)
            if section_item["type"] == "paragraph":
                self._docx_add_paragraph(document, section_item.get("content") or "无", indent=True)
            else:
                self._docx_add_list(document, ensure_text_list(section_item.get("items")))

        self._docx_add_title(document, "操作员处理栏", level=1)
        operator_table = document.add_table(rows=4, cols=2)
        operator_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        rows = report["operator_form_fields"]
        for idx, row in enumerate(rows):
            self._set_doc_cell_text(operator_table.cell(idx, 0), str(row.get("label", "")), bold=True)
            self._set_doc_cell_text(operator_table.cell(idx, 1), str(row.get("value", "")))
        self._apply_doc_table_border(operator_table)

        stream = io.BytesIO()
        document.save(stream)
        return stream.getvalue()

    def _build_pdf_report(self, report: dict[str, Any]) -> bytes:
        try:
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        except ModuleNotFoundError as exc:
            raise RuntimeError("reportlab not installed") from exc

        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("ReportTitle", parent=styles["Title"], fontName="STSong-Light", fontSize=16, leading=22, alignment=TA_CENTER)
        heading_style = ParagraphStyle("ReportHeading", parent=styles["Heading2"], fontName="STSong-Light", fontSize=12, leading=18, spaceBefore=8, spaceAfter=6)
        body_style = ParagraphStyle("ReportBody", parent=styles["BodyText"], fontName="STSong-Light", fontSize=10.5, leading=16)
        story: list[Any] = [Paragraph(report["report_title"], title_style), Spacer(1, 5 * mm)]

        meta_rows = [[Paragraph("<b>字段</b>", body_style), Paragraph("<b>内容</b>", body_style)]]
        for label, value in report["meta_items"]:
            meta_rows.append([Paragraph(label, body_style), Paragraph(clean_text(value) or "-", body_style)])
        meta_table = Table(meta_rows, colWidths=[34 * mm, 138 * mm])
        meta_table.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B8C0CC")),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F4F6F8")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        story.extend([meta_table, Spacer(1, 4 * mm), Paragraph("会话问题摘要", heading_style)])
        questions = report["question_summary"] or ["无"]
        for idx, item in enumerate(questions, start=1):
            story.append(Paragraph(f"{idx}. {item}", body_style))
        story.append(Spacer(1, 3 * mm))

        for section_item in report["sections"]:
            story.append(Paragraph(section_item["title"], heading_style))
            if section_item["type"] == "paragraph":
                story.append(Paragraph(section_item.get("content") or "无", body_style))
            else:
                items = ensure_text_list(section_item.get("items")) or ["无"]
                for idx, item in enumerate(items, start=1):
                    story.append(Paragraph(f"{idx}. {item}", body_style))
            story.append(Spacer(1, 2 * mm))

        story.append(Paragraph("操作员处理栏", heading_style))
        operator_rows = [[Paragraph("<b>字段</b>", body_style), Paragraph("<b>填写内容</b>", body_style)]]
        row_heights = [None]
        for row in report["operator_form_fields"]:
            operator_rows.append([Paragraph(str(row.get("label", "")), body_style), Paragraph(clean_text(row.get("value")) or "-", body_style)])
            row_heights.append((10 if int(row.get("row_height_lines", 1) or 1) <= 1 else 22) * mm)
        operator_table = Table(operator_rows, colWidths=[34 * mm, 138 * mm], rowHeights=row_heights)
        operator_table.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B8C0CC")),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F4F6F8")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        story.append(operator_table)

        stream = io.BytesIO()
        doc = SimpleDocTemplate(stream, pagesize=A4, leftMargin=16 * mm, rightMargin=16 * mm, topMargin=14 * mm, bottomMargin=14 * mm)
        doc.build(story)
        return stream.getvalue()

    def export_assistant_report(self, payload: dict[str, Any]) -> tuple[bytes, str, str]:
        report = self._build_assistant_report(payload)
        timestamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
        latest_result = report.get("latest_result", {})
        building_id = clean_text(next((value for label, value in report["meta_items"] if label == "建筑编号"), "ALL")) or "ALL"
        filename = f"A8_{report['filename_prefix']}_{building_id}_{timestamp}.{report['format']}"
        if report["format"] == "docx":
            content = self._build_docx_report(report)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            content = self._build_pdf_report(report)
            content_type = "application/pdf"
        if not content:
            raise RuntimeError("report content is empty")
        return content, filename, content_type

    def query_ai_evaluate(self, hours: int = 24) -> dict[str, Any]:
        safe_hours = min(max(hours, 1), 168)
        now = dt.datetime.now()
        start = now - dt.timedelta(hours=safe_hours)
        events = []
        for ev in self.ai_events:
            ts = parse_time(str(ev.get("timestamp", "")))
            if ts and ts >= start:
                events.append(ev)

        def provider_metrics(requested: str) -> dict[str, Any]:
            subset = [ev for ev in events if str(ev.get("requested_provider", "")).lower() == requested]
            total = len(subset)
            if total == 0:
                return {"total": 0, "success_rate_pct": 0.0, "avg_latency_ms": 0.0, "fallback_rate_pct": 0.0, "field_completeness_pct": 0.0}
            success = sum(1 for ev in subset if not str(ev.get("error_type", "")).strip())
            fallback = sum(1 for ev in subset if bool(ev.get("fallback_used", False)))
            complete = sum(1 for ev in subset if bool(ev.get("field_complete", False)))
            latencies = [int(ev.get("latency_ms", 0)) for ev in subset if str(ev.get("latency_ms", "")).isdigit()]
            avg_latency = round(sum(latencies) / len(latencies), 2) if latencies else 0.0
            return {
                "total": total,
                "success_rate_pct": round((success / total) * 100, 2),
                "avg_latency_ms": avg_latency,
                "fallback_rate_pct": round((fallback / total) * 100, 2),
                "field_completeness_pct": round((complete / total) * 100, 2),
            }

        template_metrics = provider_metrics("template")
        llm_metrics = provider_metrics("llm")
        auto_metrics = provider_metrics("auto")
        feedback_subset = [ev for ev in events if str(ev.get("feedback_label", "")).strip()]
        useful = sum(1 for ev in feedback_subset if str(ev.get("feedback_label")) == "useful")

        return {
            "window_hours": safe_hours,
            "template": template_metrics,
            "llm": llm_metrics,
            "auto": auto_metrics,
            "feedback": {
                "total_labeled": len(feedback_subset),
                "useful_rate_pct": round((useful / len(feedback_subset)) * 100, 2) if feedback_subset else 0.0,
            },
            "updated_at": to_iso(now),
        }

    def save_ai_feedback(self, payload: dict[str, Any]) -> dict[str, Any]:
        trace_id = str(payload.get("trace_id", "")).strip()
        label = str(payload.get("label", "")).strip().lower()
        if not trace_id:
            raise ValueError("trace_id required")
        if label not in {"useful", "not_useful"}:
            raise ValueError("label must be useful/not_useful")
        for ev in reversed(self.ai_events):
            if str(ev.get("trace_id", "")) == trace_id:
                ev["feedback_label"] = label
                with self.ai_call_log_file.open("w", encoding="utf-8") as f:
                    for row in self.ai_events:
                        f.write(json.dumps(row, ensure_ascii=False) + "\n")
                return {"trace_id": trace_id, "label": label}
        raise LookupError("trace_id not found")

    def query_system_health(self) -> dict[str, Any]:
        regression = {
            "status": "unknown",
            "updated_at": None,
            "all_ok": False,
            "steps": [],
        }
        if self.regression_summary_file.exists():
            try:
                regression = json.loads(self.regression_summary_file.read_text(encoding="utf-8"))
            except json.JSONDecodeError:
                pass

        data_source = {
            "normalized_energy": self.normalized_data_file.exists(),
            "dictionary": self.dict_file.exists(),
            "knowledge": self.knowledge_file.exists(),
        }
        ragflow = self._ragflow_settings()
        ai_status = {
            "configured": bool(os.getenv("OPENAI_API_KEY", "").strip()),
            "base_url": bool(os.getenv("OPENAI_BASE_URL", "").strip()),
            "model": bool(os.getenv("OPENAI_MODEL", "").strip()),
        }
        return {
            "status": "ok" if all(data_source.values()) else "degraded",
            "data_source": data_source,
            "storage": {
                "backend": getattr(self, "storage_backend", "file"),
            },
            "ai_provider": ai_status,
            "ragflow": {
                "configured": ragflow["configured"],
                "base_url": ragflow["base_url"],
                "web_base_url": ragflow["web_base_url"],
                "dataset_count": ragflow["dataset_count"],
                "scene_dataset_count": ragflow["dataset_count"],
                "standard_dataset_count": ragflow["standard_dataset_count"],
                "standard_configured": ragflow["standard_configured"],
                "standard_enabled": ragflow["standard_enabled"],
                "enabled": ragflow["enabled"],
                "chat_ready": ragflow["chat_ready"],
                "chat_id": ragflow["chat_id"],
                "assistant_ready": ragflow["assistant_ready"],
            },
            "recent_regression": regression,
            "updated_at": to_iso(dt.datetime.now()),
        }

    def _action_state(self, anomaly_id: int) -> dict[str, Any]:
        state = self.action_index.get(anomaly_id)
        if not state:
            return {
                "status": STATUS_NEW,
                "assignee": "",
                "last_note": "",
                "last_action_at": "",
                "history_count": 0,
            }
        return {
            "status": state["status"],
            "assignee": state["assignee"],
            "last_note": state["note"],
            "last_action_at": state["last_action_at"],
            "history_count": len(state["history"]),
        }

    def query_anomaly_history(self, anomaly_id: int) -> dict[str, Any] | None:
        if not any(a["anomaly_id"] == anomaly_id for a in self.anomalies):
            return None
        state = self.action_index.get(anomaly_id)
        history = []
        if state:
            history = list(reversed(state["history"]))
        return {"anomaly_id": anomaly_id, "count": len(history), "items": history}

    def apply_anomaly_action(self, payload: dict[str, Any]) -> dict[str, Any]:
        anomaly_id_raw = payload.get("anomaly_id")
        action = str(payload.get("action", "")).strip().lower()
        assignee = str(payload.get("assignee", "")).strip()
        note = str(payload.get("note", "")).strip()

        if anomaly_id_raw is None:
            raise ValueError("anomaly_id required")
        try:
            anomaly_id = int(anomaly_id_raw)
        except (ValueError, TypeError):
            raise ValueError("invalid anomaly_id")

        if not any(a["anomaly_id"] == anomaly_id for a in self.anomalies):
            raise LookupError("anomaly not found")

        if action not in ACTION_TO_STATUS:
            raise ValueError("invalid action")

        current_status = self._action_state(anomaly_id)["status"]
        if action not in ALLOWED_TRANSITIONS.get(current_status, set()):
            raise ValueError(f"invalid transition: {current_status} -> {action}")

        target_status = ACTION_TO_STATUS[action]
        now = to_iso(dt.datetime.now())

        event = {
            "anomaly_id": anomaly_id,
            "action": action,
            "status_before": current_status,
            "status": target_status,
            "assignee": assignee,
            "note": note,
            "created_at": now,
        }
        self._append_action_event(event)

        state = self._action_state(anomaly_id)
        return {
            "anomaly_id": anomaly_id,
            "status": state["status"],
            "assignee": state["assignee"],
            "last_note": state["last_note"],
            "last_action_at": state["last_action_at"],
            "history_count": state["history_count"],
        }

    def _filter_rows(self, building_id: str | None, start_time: dt.datetime | None, end_time: dt.datetime | None) -> list[dict[str, Any]]:
        rows = list(self.by_building.get(building_id, [])) if building_id and building_id in self.by_building else list(self.rows)
        if start_time:
            rows = [r for r in rows if r["timestamp"] >= start_time]
        if end_time:
            rows = [r for r in rows if r["timestamp"] <= end_time]
        return rows

    def _require_metric_type(self, metric_type: str | None) -> str:
        chosen = str(metric_type or "electricity").strip().lower() or "electricity"
        if chosen not in SUPPORTED_ANALYSIS_METRICS:
            raise ValueError("暂未接入该分析类型")
        return chosen

    def _summarize_values(self, values: list[float]) -> dict[str, float]:
        if not values:
            return {
                "total": 0.0,
                "avg": 0.0,
                "peak": 0.0,
                "min": 0.0,
                "std": 0.0,
                "volatility_pct": 0.0,
            }
        total = sum(values)
        avg = total / len(values)
        variance = sum((v - avg) ** 2 for v in values) / len(values)
        std = math.sqrt(variance)
        return {
            "total": round(total, 4),
            "avg": round(avg, 4),
            "peak": round(max(values), 4),
            "min": round(min(values), 4),
            "std": round(std, 4),
            "volatility_pct": round((std / avg * 100) if avg else 0.0, 2),
        }

    def _series_with_weather(self, rows: list[dict[str, Any]], building_id: str | None) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
        ordered = sorted(rows, key=lambda x: x["timestamp"])
        if building_id:
            site_id = self.building_site_map.get(building_id)
        else:
            site_id = None

        weather_lookup = self.weather_by_site.get(site_id or "", {})
        series: list[dict[str, Any]] = []
        weather_series: list[dict[str, Any]] = []

        if not building_id:
            buckets: dict[dt.datetime, float] = {}
            for row in ordered:
                buckets.setdefault(row["timestamp"], 0.0)
                buckets[row["timestamp"]] += row["electricity_kwh"]
            for timestamp, value in sorted(buckets.items(), key=lambda x: x[0]):
                weather = weather_lookup.get(timestamp, {})
                point = {
                    "timestamp": to_iso(timestamp),
                    "value": round(value, 4),
                    "temperature_c": weather.get("temperature_c"),
                }
                series.append(point)
                if weather.get("temperature_c") is not None:
                    weather_series.append({"timestamp": to_iso(timestamp), "value": weather["temperature_c"]})
            return series, weather_series

        for row in ordered:
            weather = weather_lookup.get(row["timestamp"], {})
            point = {
                "timestamp": to_iso(row["timestamp"]),
                "value": round(row["electricity_kwh"], 4),
                "temperature_c": weather.get("temperature_c"),
            }
            series.append(point)
            if weather.get("temperature_c") is not None:
                weather_series.append({"timestamp": to_iso(row["timestamp"]), "value": weather["temperature_c"]})
        return series, weather_series

    def _temperature_correlation(self, series: list[dict[str, Any]]) -> float:
        pairs = [(float(item["value"]), float(item["temperature_c"])) for item in series if item.get("temperature_c") is not None]
        if len(pairs) < 2:
            return 0.0
        xs = [x for x, _ in pairs]
        ys = [y for _, y in pairs]
        x_avg = sum(xs) / len(xs)
        y_avg = sum(ys) / len(ys)
        numerator = sum((x - x_avg) * (y - y_avg) for x, y in pairs)
        denominator = math.sqrt(sum((x - x_avg) ** 2 for x in xs) * sum((y - y_avg) ** 2 for y in ys))
        if not denominator:
            return 0.0
        return round(numerator / denominator, 4)

    def _granularity_label(self, rows: list[dict[str, Any]]) -> str:
        ordered = sorted(rows, key=lambda item: item["timestamp"])
        if len(ordered) < 2:
            return "single-point"
        deltas = [
            max(1.0, round((ordered[index + 1]["timestamp"] - ordered[index]["timestamp"]).total_seconds() / 3600, 2))
            for index in range(min(len(ordered) - 1, 48))
        ]
        avg_delta = sum(deltas) / len(deltas)
        if avg_delta <= 1.2:
            return "hourly"
        if avg_delta <= 6.5:
            return "6-hour"
        if avg_delta <= 12.5:
            return "12-hour"
        return "daily"

    def _comparison_series(self, rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
        ordered = sorted(rows, key=lambda item: item["timestamp"])
        if not ordered:
            return []
        by_hour: dict[int, list[float]] = {hour: [] for hour in range(24)}
        for row in ordered:
            by_hour[int(row["hour"])].append(float(row["electricity_kwh"]))
        baseline_by_hour = {
            hour: (sum(values) / len(values)) if values else 0.0
            for hour, values in by_hour.items()
        }
        return [
            {
                "timestamp": to_iso(row["timestamp"]),
                "value": round(baseline_by_hour[int(row["hour"])], 4),
            }
            for row in ordered
        ]

    def _weather_relation_stats(self, series: list[dict[str, Any]]) -> dict[str, float]:
        pairs = [
            (float(item["temperature_c"]), float(item["value"]))
            for item in series
            if item.get("temperature_c") is not None
        ]
        if len(pairs) < 6:
            return {
                "hot_avg_load": 0.0,
                "cold_avg_load": 0.0,
                "hot_avg_temp": 0.0,
                "cold_avg_temp": 0.0,
                "hot_cold_gap_pct": 0.0,
            }

        sorted_pairs = sorted(pairs, key=lambda item: item[0])
        bucket_size = max(2, len(sorted_pairs) // 3)
        cold_bucket = sorted_pairs[:bucket_size]
        hot_bucket = sorted_pairs[-bucket_size:]
        cold_avg_load = sum(item[1] for item in cold_bucket) / len(cold_bucket)
        hot_avg_load = sum(item[1] for item in hot_bucket) / len(hot_bucket)
        cold_avg_temp = sum(item[0] for item in cold_bucket) / len(cold_bucket)
        hot_avg_temp = sum(item[0] for item in hot_bucket) / len(hot_bucket)
        gap_pct = ((hot_avg_load - cold_avg_load) / cold_avg_load * 100) if cold_avg_load else 0.0
        return {
            "hot_avg_load": round(hot_avg_load, 4),
            "cold_avg_load": round(cold_avg_load, 4),
            "hot_avg_temp": round(hot_avg_temp, 2),
            "cold_avg_temp": round(cold_avg_temp, 2),
            "hot_cold_gap_pct": round(gap_pct, 2),
        }

    def _trend_markers(self, anomalies: list[dict[str, Any]]) -> list[dict[str, Any]]:
        ranked = sorted(
            anomalies,
            key=lambda item: (SEVERITY_SCORE.get(str(item.get("severity")), 0), float(item.get("deviation_pct", 0.0))),
            reverse=True,
        )[:12]
        ranked.sort(key=lambda item: item["timestamp"])
        return [
            {
                "anomaly_id": int(item["anomaly_id"]),
                "timestamp": to_iso(item["timestamp"]),
                "anomaly_name": self.dict_data.get(str(item.get("anomaly_type", "")), {}).get("name", item.get("anomaly_type", "异常")),
                "severity": item["severity"],
                "deviation_pct": round(float(item["deviation_pct"]), 2),
                "value": round(float(item["electricity_kwh"]), 4),
                "estimated_loss_kwh": round(max(float(item["electricity_kwh"]) - float(item["mean_kwh"]), 0.0), 4),
            }
            for item in ranked
        ]

    def _insight_item(self, title: str, detail: str, severity: str = "info") -> dict[str, str]:
        return {"title": title, "detail": detail, "severity": severity}

    def _opportunity_item(self, title: str, detail: str, priority: str, estimated_kwh: float = 0.0) -> dict[str, Any]:
        estimated_value = round(float(estimated_kwh), 4)
        return {
            "title": title,
            "detail": detail,
            "priority": priority,
            "estimated_kwh": estimated_value,
            "estimated_loss_kwh": estimated_value,
        }

    def _analysis_target_building(self, building_id: str | None, rows: list[dict[str, Any]]) -> tuple[str | None, list[dict[str, Any]]]:
        if building_id and building_id in self.by_building:
            return building_id, self._filter_rows(building_id, None, None)
        if rows:
            fallback_id = rows[0]["building_id"]
            return fallback_id, list(self.by_building.get(fallback_id, []))
        if self.buildings_meta:
            fallback_id = sorted(self.buildings_meta.keys())[0]
            return fallback_id, list(self.by_building.get(fallback_id, []))
        return None, []

    def _peer_compare_pool(
        self,
        peer_category: str,
        start_time: dt.datetime | None,
        end_time: dt.datetime | None,
    ) -> dict[str, float]:
        cache_key = (
            peer_category,
            to_iso(start_time) if start_time else None,
            to_iso(end_time) if end_time else None,
        )
        cached = self.compare_pool_cache.get(cache_key)
        if cached is not None:
            return dict(cached)

        candidate_ids = [
            bid
            for bid in self.peer_category_to_buildings.get(peer_category, [])
            if bid in self.raw_electricity_headers
        ]
        if not candidate_ids or not self.raw_electricity_file.exists():
            self.compare_pool_cache[cache_key] = {}
            return {}

        sums = {bid: 0.0 for bid in candidate_ids}
        counts = {bid: 0 for bid in candidate_ids}
        with self.raw_electricity_file.open("r", encoding="utf-8", newline="") as f:
            reader = csv.DictReader(f)
            for row in reader:
                ts = parse_time(row.get("timestamp", ""))
                if not ts:
                    continue
                if start_time and ts < start_time:
                    continue
                if end_time and ts > end_time:
                    continue
                for bid in candidate_ids:
                    raw_val = row.get(bid, "")
                    if raw_val is None:
                        continue
                    raw_val = str(raw_val).strip()
                    if not raw_val:
                        continue
                    try:
                        value = float(raw_val)
                    except ValueError:
                        continue
                    if value < 0:
                        continue
                    sums[bid] += value
                    counts[bid] += 1
        result = {
            bid: (sums[bid] / counts[bid])
            for bid in candidate_ids
            if counts[bid] > 0
        }
        self.compare_pool_cache[cache_key] = dict(result)
        return result

    def _sort_anomalies(self, rows: list[dict[str, Any]], sort: str) -> list[dict[str, Any]]:
        if sort == "timestamp_asc":
            return sorted(rows, key=lambda x: x["timestamp"])
        if sort == "deviation_desc":
            return sorted(rows, key=lambda x: abs(float(x.get("deviation_pct", 0))), reverse=True)
        if sort == "severity_desc":
            return sorted(rows, key=lambda x: (SEVERITY_SCORE.get(x.get("severity", "low"), 0), x["timestamp"]), reverse=True)
        return sorted(rows, key=lambda x: x["timestamp"], reverse=True)

    def _filter_anomalies(
        self,
        building_id: str | None,
        start_time: dt.datetime | None,
        end_time: dt.datetime | None,
        anomaly_type: str | None = None,
        severity: str | None = None,
        status: str | None = None,
    ) -> list[dict[str, Any]]:
        result = list(self.anomalies)
        if building_id:
            result = [x for x in result if x["building_id"] == building_id]
        if start_time:
            result = [x for x in result if x["timestamp"] >= start_time]
        if end_time:
            result = [x for x in result if x["timestamp"] <= end_time]
        if anomaly_type:
            result = [x for x in result if x["anomaly_type"] == anomaly_type]
        if severity:
            result = [x for x in result if x["severity"] == severity]
        if status:
            result = [x for x in result if self._action_state(int(x["anomaly_id"]))["status"] == status]
        return result

    def query_buildings(self) -> dict[str, Any]:
        items = []
        for building_id in self.showcase_config:
            meta = self.buildings_meta.get(building_id)
            if not meta:
                continue
            items.append(
                {
                    "building_id": meta["building_id"],
                    "building_name": meta["building_name"],
                    "building_type": meta["building_type"],
                    "display_name": meta.get("display_name", meta["building_name"]),
                    "display_category": meta.get("display_category", meta["building_type"]),
                    "peer_category": meta.get("peer_category", ""),
                    "record_count": meta["record_count"],
                    "start_time": to_iso(meta["start_time"]),
                    "end_time": to_iso(meta["end_time"]),
                }
            )
        global_start = min((m["start_time"] for m in self.buildings_meta.values()), default=None)
        global_end = max((m["end_time"] for m in self.buildings_meta.values()), default=None)
        return {
            "count": len(items),
            "items": items,
            "global_range": {
                "start_time": to_iso(global_start) if global_start else None,
                "end_time": to_iso(global_end) if global_end else None,
            },
        }

    def query_analysis_summary(
        self,
        building_id: str | None,
        start_time: dt.datetime | None,
        end_time: dt.datetime | None,
        metric_type: str | None = "electricity",
    ) -> dict[str, Any]:
        metric = self._require_metric_type(metric_type)
        rows = self._filter_rows(building_id, start_time, end_time)
        values = [r["electricity_kwh"] for r in rows]
        summary = self._summarize_values(values)
        anomalies = self._filter_anomalies(building_id, start_time, end_time)
        working_hours = [r["electricity_kwh"] for r in rows if 8 <= r["hour"] <= 20]
        off_hours = [r["electricity_kwh"] for r in rows if r["hour"] < 8 or r["hour"] > 20]

        return {
            "metric_type": metric,
            "metric_label": "电力",
            "unit": "kWh",
            "building_id": building_id or "ALL",
            "point_count": len(rows),
            "total_value": summary["total"],
            "avg_value": summary["avg"],
            "peak_value": summary["peak"],
            "min_value": summary["min"],
            "std_value": summary["std"],
            "volatility_pct": summary["volatility_pct"],
            "anomaly_count": len(anomalies),
            "working_hour_avg": round(sum(working_hours) / len(working_hours), 4) if working_hours else 0.0,
            "off_hour_avg": round(sum(off_hours) / len(off_hours), 4) if off_hours else 0.0,
            "supported_metric_types": sorted(SUPPORTED_ANALYSIS_METRICS),
            "method_note": "normalized electricity + anomaly context + weather ready",
        }

    def query_analysis_trend(
        self,
        building_id: str | None,
        start_time: dt.datetime | None,
        end_time: dt.datetime | None,
        metric_type: str | None = "electricity",
    ) -> dict[str, Any]:
        metric = self._require_metric_type(metric_type)
        rows = self._filter_rows(building_id, start_time, end_time)
        series, weather_series = self._series_with_weather(rows, building_id)
        anomalies = self._filter_anomalies(building_id, start_time, end_time)
        values = [item["value"] for item in series]
        summary = self._summarize_values(values)

        if len(values) >= 48:
            head = values[:24]
            tail = values[-24:]
            head_avg = sum(head) / len(head)
            tail_avg = sum(tail) / len(tail)
            change_pct = round(((tail_avg - head_avg) / head_avg * 100) if head_avg else 0.0, 2)
        else:
            change_pct = 0.0

        return {
            "metric_type": metric,
            "metric_label": "电力",
            "unit": "kWh",
            "building_id": building_id or "ALL",
            "series": series,
            "comparison_series": self._comparison_series(rows),
            "weather_series": weather_series,
            "markers": self._trend_markers(anomalies),
            "overlay_available": bool(weather_series),
            "summary": {
                "total_value": summary["total"],
                "avg_value": summary["avg"],
                "peak_value": summary["peak"],
                "volatility_pct": summary["volatility_pct"],
                "window_change_pct": change_pct,
                "temperature_correlation": self._temperature_correlation(series),
            },
        }

    def query_analysis_distribution(
        self,
        building_id: str | None,
        start_time: dt.datetime | None,
        end_time: dt.datetime | None,
        metric_type: str | None = "electricity",
    ) -> dict[str, Any]:
        metric = self._require_metric_type(metric_type)
        rows = self._filter_rows(building_id, start_time, end_time)
        hourly: dict[int, list[float]] = {hour: [] for hour in range(24)}
        weekday_hourly: dict[int, list[float]] = {hour: [] for hour in range(24)}
        weekday_total = 0.0
        weekend_total = 0.0
        day_total = 0.0
        night_total = 0.0
        night_base_values: list[float] = []

        for row in rows:
            value = float(row["electricity_kwh"])
            hourly[row["hour"]].append(value)
            if row["timestamp"].weekday() < 5:
                weekday_total += value
                weekday_hourly[row["hour"]].append(value)
            else:
                weekend_total += value
            if 8 <= row["hour"] <= 20:
                day_total += value
            else:
                night_total += value
            if row["hour"] < 6 or row["hour"] >= 22:
                night_base_values.append(value)

        hourly_profile = [
            {
                "hour": hour,
                "label": f"{hour:02d}:00",
                "avg_value": round(sum(values) / len(values), 4) if values else 0.0,
            }
            for hour, values in hourly.items()
        ]
        total = weekday_total + weekend_total
        active_total = day_total + night_total
        weekday_peak_hours = [
            {
                "hour": hour,
                "label": f"{hour:02d}:00",
                "avg_value": round(sum(values) / len(values), 4),
            }
            for hour, values in weekday_hourly.items()
            if values
        ]
        weekday_peak_hours.sort(key=lambda item: item["avg_value"], reverse=True)
        night_base_avg = round(sum(night_base_values) / len(night_base_values), 4) if night_base_values else 0.0
        summary_avg = round(sum(float(row["electricity_kwh"]) for row in rows) / len(rows), 4) if rows else 0.0

        return {
            "metric_type": metric,
            "metric_label": "电力",
            "unit": "kWh",
            "building_id": building_id or "ALL",
            "hourly_profile": hourly_profile,
            "weekday_weekend_split": [
                {"label": "工作日", "value": round(weekday_total, 4), "ratio_pct": round((weekday_total / total * 100) if total else 0.0, 2)},
                {"label": "周末", "value": round(weekend_total, 4), "ratio_pct": round((weekend_total / total * 100) if total else 0.0, 2)},
            ],
            "day_night_split": [
                {"label": "白天", "value": round(day_total, 4), "ratio_pct": round((day_total / active_total * 100) if active_total else 0.0, 2)},
                {"label": "夜间", "value": round(night_total, 4), "ratio_pct": round((night_total / active_total * 100) if active_total else 0.0, 2)},
            ],
            "weekday_peak_hours": weekday_peak_hours[:3],
            "night_base_load": {
                "avg_value": night_base_avg,
                "ratio_vs_avg_pct": round((night_base_avg / summary_avg * 100) if summary_avg else 0.0, 2),
            },
        }

    def query_analysis_compare(
        self,
        building_id: str | None,
        start_time: dt.datetime | None,
        end_time: dt.datetime | None,
        metric_type: str | None = "electricity",
    ) -> dict[str, Any]:
        metric = self._require_metric_type(metric_type)
        if not building_id:
            return {
                "metric_type": metric,
                "metric_label": "电力",
                "unit": "kWh",
                "building": None,
                "peer_group": None,
                "items": [],
                "peer_ranking": [],
                "message": "请选择单体建筑后查看同类对比。",
            }
        scoped_rows = self._filter_rows(building_id, start_time, end_time)
        target_building_id, _ = self._analysis_target_building(building_id, scoped_rows)
        if not target_building_id:
            return {
                "metric_type": metric,
                "metric_label": "电力",
                "unit": "kWh",
                "building": None,
                "peer_group": None,
                "items": [],
                "peer_ranking": [],
                "message": "当前筛选范围暂无可对比数据。",
            }

        target_meta = self.buildings_meta.get(target_building_id, {})
        peer_category = str(target_meta.get("peer_category", "")).strip()
        if not peer_category:
            return {
                "metric_type": metric,
                "metric_label": "电力",
                "unit": "kWh",
                "building": None,
                "peer_group": None,
                "items": [],
                "peer_ranking": [],
                "message": "当前建筑未纳入代表建筑展示集或未配置对标类别。",
            }

        target_rows = self._filter_rows(target_building_id, start_time, end_time)
        target_values = [float(row["electricity_kwh"]) for row in target_rows]
        target_avg = round(sum(target_values) / len(target_values), 4) if target_values else 0.0
        peer_candidates = self._peer_compare_pool(peer_category, start_time, end_time)
        if target_building_id in peer_candidates:
            peer_candidates = {bid: avg for bid, avg in peer_candidates.items() if bid != target_building_id}
        if not peer_candidates:
            return {
                "metric_type": metric,
                "metric_label": "电力",
                "unit": "kWh",
                "building": {
                    "building_id": target_building_id,
                    "building_name": target_meta.get("display_name", target_building_id),
                    "building_type": target_meta.get("display_category", ""),
                    "avg_value": target_avg,
                },
                "peer_group": {
                    "peer_category": peer_category,
                    "peer_category_label": PEER_CATEGORY_LABELS.get(peer_category, peer_category),
                    "peer_avg_value": 0.0,
                    "peer_count": 0,
                    "vs_peer_pct": 0.0,
                    "gap_pct": 0.0,
                    "peer_percentile": 0.0,
                    "ranking_position": None,
                },
                "items": [{"label": "当前建筑", "value": target_avg}],
                "peer_ranking": [],
                "message": "当前类别在所选时间范围内无可比样本。",
            }

        peer_avg = round(sum(peer_candidates.values()) / len(peer_candidates), 4)
        all_peer_values = list(peer_candidates.values()) + [target_avg]
        peer_total_count = len(all_peer_values)
        vs_peer_pct = round(((target_avg - peer_avg) / peer_avg * 100) if peer_avg else 0.0, 2)
        percentile = round(
            (
                sum(1 for avg in all_peer_values if avg <= target_avg) / peer_total_count * 100
            ) if peer_total_count else 0.0,
            2,
        )
        ranking_position = sum(1 for avg in peer_candidates.values() if avg > target_avg) + 1

        ranking = [
            {
                "building_id": bid,
                "building_name": self.bdq2_metadata.get(bid, {}).get("display_name", showcase_display_name(bid, peer_category)),
                "avg_value": round(avg, 4),
            }
            for bid, avg in sorted(peer_candidates.items(), key=lambda x: x[1], reverse=True)[:8]
        ]

        return {
            "metric_type": metric,
            "metric_label": "电力",
            "unit": "kWh",
            "building": {
                "building_id": target_building_id,
                "building_name": target_meta.get("display_name", target_building_id),
                "building_type": target_meta.get("display_category", ""),
                "avg_value": target_avg,
            },
            "peer_group": {
                "peer_category": peer_category,
                "peer_category_label": PEER_CATEGORY_LABELS.get(peer_category, peer_category),
                "peer_avg_value": peer_avg,
                "peer_count": peer_total_count,
                "vs_peer_pct": vs_peer_pct,
                "gap_pct": vs_peer_pct,
                "peer_percentile": percentile,
                "ranking_position": ranking_position,
            },
            "items": [
                {"label": "当前建筑", "value": target_avg},
                {"label": "同类均值", "value": peer_avg},
            ],
            "peer_ranking": ranking,
            "message": "",
        }

    def query_analysis_insights(
        self,
        building_id: str | None,
        start_time: dt.datetime | None,
        end_time: dt.datetime | None,
        metric_type: str | None = "electricity",
    ) -> dict[str, Any]:
        metric = self._require_metric_type(metric_type)
        rows = self._filter_rows(building_id, start_time, end_time)
        ordered_rows = sorted(rows, key=lambda item: item["timestamp"])
        summary = self.query_analysis_summary(building_id, start_time, end_time, metric)
        trend = self.query_analysis_trend(building_id, start_time, end_time, metric)
        distribution = self.query_analysis_distribution(building_id, start_time, end_time, metric)
        compare = self.query_analysis_compare(building_id, start_time, end_time, metric)
        anomalies = self._filter_anomalies(building_id, start_time, end_time)
        building_meta = self.buildings_meta.get(building_id or "", {})
        weather_stats = self._weather_relation_stats(trend["series"])

        scope_summary = {
            "building_id": building_id or "ALL",
            "building_name": building_meta.get("building_name", "全部建筑"),
            "building_type": building_meta.get("building_type", "portfolio") if building_id else "portfolio",
            "selected_start_time": to_iso(start_time) if start_time else None,
            "selected_end_time": to_iso(end_time) if end_time else None,
            "data_start_time": to_iso(ordered_rows[0]["timestamp"]) if ordered_rows else None,
            "data_end_time": to_iso(ordered_rows[-1]["timestamp"]) if ordered_rows else None,
            "point_count": len(ordered_rows),
            "granularity": self._granularity_label(ordered_rows),
            "anomaly_count": len(anomalies),
            "metric_label": "电力",
            "unit": "kWh",
        }

        trend_findings: list[dict[str, str]] = []
        weather_findings: list[dict[str, str]] = []
        compare_findings: list[dict[str, str]] = []
        saving_opportunities: list[dict[str, Any]] = []

        change_pct = float(trend["summary"].get("window_change_pct", 0.0))
        volatility_pct = float(summary.get("volatility_pct", 0.0))
        temp_corr = float(trend["summary"].get("temperature_correlation", 0.0))
        working_hour_avg = float(summary.get("working_hour_avg", 0.0))
        off_hour_avg = float(summary.get("off_hour_avg", 0.0))
        working_off_gap_pct = ((working_hour_avg - off_hour_avg) / off_hour_avg * 100) if off_hour_avg else 0.0

        if change_pct >= 12:
            trend_findings.append(self._insight_item("近期负荷抬升", f"前后窗口均值上升 {round(change_pct, 2)}%，近期存在明显增载。", "warning"))
        elif change_pct <= -12:
            trend_findings.append(self._insight_item("近期负荷回落", f"前后窗口均值下降 {round(abs(change_pct), 2)}%，近期运行压力有所回落。", "success"))
        else:
            trend_findings.append(self._insight_item("趋势整体平稳", f"窗口变化 {round(change_pct, 2)}%，当前负荷没有明显阶跃变化。", "info"))

        if volatility_pct >= 35:
            trend_findings.append(self._insight_item("波动偏大", f"波动率 {round(volatility_pct, 2)}%，建议重点关注启停频繁时段。", "danger"))
        else:
            trend_findings.append(self._insight_item("稳定性可控", f"波动率 {round(volatility_pct, 2)}%，曲线整体处于可解释区间。", "success"))

        if working_off_gap_pct >= 60:
            trend_findings.append(self._insight_item("工作时段主导负荷", f"工作时段均值较非工作时段高 {round(working_off_gap_pct, 2)}%，节能重点应放在白天主业务时段。", "info"))
        elif off_hour_avg and (off_hour_avg / max(working_hour_avg, 1e-6)) >= 0.7:
            trend_findings.append(self._insight_item("非工作时段负荷偏高", "非工作时段负荷没有明显回落，可能存在待机或常开设备。", "warning"))

        if trend["overlay_available"]:
            if temp_corr >= 0.35:
                weather_findings.append(self._insight_item("温度正相关明显", f"温度相关系数 {round(temp_corr, 2)}，热天气会显著抬升负荷。", "warning"))
            elif temp_corr <= -0.35:
                weather_findings.append(self._insight_item("温度反向影响", f"温度相关系数 {round(temp_corr, 2)}，低温时段负荷抬升更明显。", "info"))
            else:
                weather_findings.append(self._insight_item("天气影响有限", f"温度相关系数 {round(temp_corr, 2)}，当前窗口内负荷更受内部运行策略影响。", "info"))

            if weather_stats["hot_avg_load"] and weather_stats["cold_avg_load"]:
                weather_findings.append(
                    self._insight_item(
                        "冷热天差异",
                        f"热天均值 {weather_stats['hot_avg_load']} kWh，冷天均值 {weather_stats['cold_avg_load']} kWh，差异 {round(weather_stats['hot_cold_gap_pct'], 2)}%。",
                        "info" if abs(weather_stats["hot_cold_gap_pct"]) < 12 else "warning",
                    )
                )

        peer_group = compare.get("peer_group") or {}
        if compare.get("message"):
            compare_findings.append(self._insight_item("同类对比待选定", compare["message"], "info"))
        else:
            compare_findings.append(
                self._insight_item(
                    "同类位置",
                    f"当前建筑在同类中位于前 {round(peer_group.get('peer_percentile', 0.0), 2)} 百分位，排名 {peer_group.get('ranking_position') or '-'} / {peer_group.get('peer_count') or '-'}。",
                    "warning" if float(peer_group.get("gap_pct", 0.0)) > 10 else "info",
                )
            )
            compare_findings.append(
                self._insight_item(
                    "与同类均值差距",
                    f"当前均值较同类均值 {'高' if float(peer_group.get('gap_pct', 0.0)) >= 0 else '低'} {round(abs(float(peer_group.get('gap_pct', 0.0))), 2)}%。",
                    "warning" if abs(float(peer_group.get("gap_pct", 0.0))) >= 10 else "success",
                )
            )

        if distribution["night_base_load"]["avg_value"] > 0:
            saving_opportunities.append(
                self._opportunity_item(
                    "夜间基线负荷优化",
                    f"夜间基线约 {distribution['night_base_load']['avg_value']} kWh，为整体均值的 {distribution['night_base_load']['ratio_vs_avg_pct']}%。可优先排查夜间常开设备。",
                    "high" if distribution["night_base_load"]["ratio_vs_avg_pct"] >= 55 else "medium",
                    distribution["night_base_load"]["avg_value"],
                )
            )

        anomaly_waste = sum(max(float(item["electricity_kwh"]) - float(item["mean_kwh"]), 0.0) for item in anomalies)
        if anomaly_waste > 0:
            saving_opportunities.append(
                self._opportunity_item(
                    "异常浪费回收",
                    f"当前窗口异常额外损失约 {round(anomaly_waste, 2)} kWh，可通过优先处理突增和持续高负荷事件回收损耗。",
                    "high" if len(anomalies) >= 3 else "medium",
                    anomaly_waste,
                )
            )

        if float(peer_group.get("gap_pct", 0.0)) > 8 and summary["point_count"] > 0:
            estimated_gap = max(float(summary["avg_value"]) - float(peer_group.get("peer_avg_value", 0.0)), 0.0) * summary["point_count"]
            saving_opportunities.append(
                self._opportunity_item(
                    "同类差距收敛",
                    f"若将均值收敛到同类水平，当前窗口理论可优化约 {round(estimated_gap, 2)} kWh。",
                    "medium",
                    estimated_gap,
                )
            )

        peak_hours = distribution.get("weekday_peak_hours") or []
        if peak_hours:
            top_hour = peak_hours[0]
            saving_opportunities.append(
                self._opportunity_item(
                    "高峰时段策略优化",
                    f"工作日高峰集中在 {top_hour['label']} 左右，平均负荷 {top_hour['avg_value']} kWh，适合做错峰和设定点优化。",
                    "medium",
                    top_hour["avg_value"],
                )
            )

        anomaly_windows = [
            {
                "anomaly_id": item["anomaly_id"],
                "timestamp": item["timestamp"],
                "anomaly_name": item["anomaly_name"],
                "severity": item["severity"],
                "deviation_pct": item["deviation_pct"],
                "estimated_loss_kwh": item["estimated_loss_kwh"],
                "rule_code": item.get("rule_code", item.get("anomaly_type")),
                "rule_name": item.get("rule_name", item.get("anomaly_name")),
                "rule_summary": item.get("rule_summary", ""),
                "baseline_value": item.get("baseline_value", item.get("mean_kwh")),
                "current_value": item.get("current_value", item.get("electricity_kwh")),
                "threshold_value": item.get("threshold_value", item.get("threshold")),
                "trigger_window": item.get("trigger_window", item["timestamp"] if isinstance(item.get("timestamp"), str) else to_iso(item["timestamp"])),
                "trigger_count": item.get("trigger_count", 1),
                "time_scope_label": item.get("time_scope_label", "异常时段"),
            }
            for item in self._trend_markers(anomalies)[:4]
        ]

        return {
            "metric_type": metric,
            "scope_summary": scope_summary,
            "trend_findings": trend_findings[:4],
            "weather_findings": weather_findings[:3],
            "compare_findings": compare_findings[:3],
            "saving_opportunities": saving_opportunities[:4],
            "anomaly_windows": anomaly_windows,
        }

    def query_trend(self, building_id: str | None, start_time: dt.datetime | None, end_time: dt.datetime | None) -> dict[str, Any]:
        data = self.query_analysis_trend(building_id, start_time, end_time, "electricity")
        return {
            "unit": "kWh",
            "building_id": building_id or "ALL",
            "point_count": len(data["series"]),
            "series": data["series"],
            "summary": {
                "total_kwh": data["summary"]["total_value"],
                "avg_kwh": data["summary"]["avg_value"],
                "peak_kwh": data["summary"]["peak_value"],
            },
        }

    def query_rank(self, month: str | None) -> dict[str, Any]:
        target_month = month or max(r["timestamp"].strftime("%Y-%m") for r in self.rows)
        rank_items = []
        for building_id, rows in self.by_building.items():
            month_rows = [r for r in rows if r["timestamp"].strftime("%Y-%m") == target_month]
            if not month_rows:
                continue
            total = sum(r["electricity_kwh"] for r in month_rows)
            avg = total / len(month_rows)
            rank_items.append(
                {
                    "building_id": building_id,
                    "building_name": month_rows[0]["building_name"],
                    "building_type": month_rows[0]["building_type"],
                    "month": target_month,
                    "total_kwh": round(total, 4),
                    "avg_kwh": round(avg, 4),
                }
            )
        rank_items.sort(key=lambda x: x["avg_kwh"], reverse=True)
        for i, item in enumerate(rank_items, start=1):
            item["rank"] = i
        return {"month": target_month, "unit": "kWh", "items": rank_items}

    def query_anomalies(
        self,
        building_id: str | None,
        start_time: dt.datetime | None,
        end_time: dt.datetime | None,
        anomaly_type: str | None,
        severity: str | None,
        status: str | None,
        page: int = 1,
        page_size: int = 20,
        sort: str = "timestamp_desc",
    ) -> dict[str, Any]:
        filtered = self._filter_anomalies(building_id, start_time, end_time, anomaly_type, severity, status)
        available_type_source = self._filter_anomalies(building_id, start_time, end_time, None, severity, status)
        sorted_rows = self._sort_anomalies(filtered, sort)

        by_type: dict[str, int] = {}
        by_type_code: dict[str, dict[str, Any]] = {}
        by_status: dict[str, int] = {}
        for item in sorted_rows:
            type_label = str(item.get("display_name") or item.get("anomaly_name") or self._anomaly_name(str(item.get("anomaly_type", ""))))
            by_type[type_label] = by_type.get(type_label, 0) + 1
            type_code = str(item.get("anomaly_type", "")).strip()
            if type_code:
                bucket = by_type_code.setdefault(type_code, {"code": type_code, "name": type_label, "count": 0})
                bucket["count"] = int(bucket.get("count", 0)) + 1
            st = self._action_state(int(item["anomaly_id"]))["status"]
            by_status[st] = by_status.get(st, 0) + 1

        available_type_code: dict[str, dict[str, Any]] = {}
        for item in available_type_source:
            type_label = str(item.get("display_name") or item.get("anomaly_name") or self._anomaly_name(str(item.get("anomaly_type", ""))))
            type_code = str(item.get("anomaly_type", "")).strip()
            if not type_code:
                continue
            bucket = available_type_code.setdefault(type_code, {"code": type_code, "name": type_label, "count": 0})
            bucket["count"] = int(bucket.get("count", 0)) + 1

        safe_page = max(1, page)
        safe_page_size = min(max(1, page_size), 200)
        total_count = len(sorted_rows)
        total_pages = max(1, math.ceil(total_count / safe_page_size))
        if safe_page > total_pages:
            safe_page = total_pages

        start_idx = (safe_page - 1) * safe_page_size
        end_idx = start_idx + safe_page_size
        page_rows = sorted_rows[start_idx:end_idx]

        items = []
        for x in page_rows:
            action_state = self._action_state(int(x["anomaly_id"]))
            items.append(
                {
                    **x,
                    "timestamp": to_iso(x["timestamp"]),
                    "anomaly_name": str(x.get("anomaly_name") or self._anomaly_name(x["anomaly_type"])),
                    "display_name": str(x.get("display_name") or self._anomaly_name(x["anomaly_type"])),
                    "status": action_state["status"],
                    "assignee": action_state["assignee"],
                    "last_note": action_state["last_note"],
                    "last_action_at": action_state["last_action_at"],
                }
            )

        return {
            "count": len(items),
            "total_count": total_count,
            "page": safe_page,
            "page_size": safe_page_size,
            "total_pages": total_pages,
            "sort": sort,
            "severity": severity,
            "status": status,
            "by_type": by_type,
            "available_types": sorted(available_type_code.values(), key=lambda item: (-int(item.get("count", 0)), str(item.get("name", "")))),
            "by_status": by_status,
            "items": items,
        }

    def query_metrics_overview(self, building_id: str | None, start_time: dt.datetime | None, end_time: dt.datetime | None) -> dict[str, Any]:
        rows = self._filter_rows(building_id, start_time, end_time)
        anomalies = self._filter_anomalies(building_id, start_time, end_time)
        values = [r["electricity_kwh"] for r in rows]
        total = sum(values)
        avg = total / len(values) if values else 0
        peak = max(values) if values else 0
        peak_excess_pct = ((peak - avg) / avg * 100) if avg else 0
        wasted_kwh = sum(max(a["electricity_kwh"] - a["mean_kwh"], 0) for a in anomalies)
        carbon_reduction_kg = wasted_kwh * CARBON_FACTOR
        return {
            "building_id": building_id or "ALL",
            "unit": "kWh",
            "carbon_factor": CARBON_FACTOR,
            "total_kwh": round(total, 4),
            "avg_kwh": round(avg, 4),
            "peak_kwh": round(peak, 4),
            "anomaly_count": len(anomalies),
            "saving_potential_pct": round(max(peak_excess_pct, 0), 2),
            "wasted_kwh": round(wasted_kwh, 4),
            "carbon_reduction_kg": round(carbon_reduction_kg, 4),
            "method_note": "peak_excess + anomaly_waste",
        }

    def query_saving_potential(self, building_id: str | None, start_time: dt.datetime | None, end_time: dt.datetime | None) -> dict[str, Any]:
        rows = self._filter_rows(building_id, start_time, end_time)
        anomalies = self._filter_anomalies(building_id, start_time, end_time)
        values = [r["electricity_kwh"] for r in rows]
        avg = sum(values) / len(values) if values else 0
        peak = max(values) if values else 0
        peak_excess_pct = ((peak - avg) / avg * 100) if avg else 0
        wasted_kwh = sum(max(a["electricity_kwh"] - a["mean_kwh"], 0) for a in anomalies)
        carbon_reduction_kg = wasted_kwh * CARBON_FACTOR

        by_building: dict[str, list[float]] = {}
        btype: dict[str, str] = {}
        for r in rows:
            by_building.setdefault(r["building_id"], []).append(r["electricity_kwh"])
            btype[r["building_id"]] = r["building_type"]

        peer_info = {"building_id": None, "building_type": None, "vs_type_avg_pct": 0.0, "building_avg_kwh": 0.0, "type_avg_kwh": 0.0}
        if by_building:
            building_avg = {k: sum(v) / len(v) for k, v in by_building.items() if v}
            type_groups: dict[str, list[float]] = {}
            for bid, bav in building_avg.items():
                type_groups.setdefault(btype.get(bid, "unknown"), []).append(bav)
            type_avg = {k: sum(v) / len(v) for k, v in type_groups.items() if v}
            target_bid = building_id if building_id in building_avg else max(building_avg, key=building_avg.get)
            target_type = btype.get(target_bid, "unknown")
            t_avg = type_avg.get(target_type, 0)
            b_avg = building_avg.get(target_bid, 0)
            vs_pct = ((b_avg - t_avg) / t_avg * 100) if t_avg else 0
            peer_info = {
                "building_id": target_bid,
                "building_type": target_type,
                "vs_type_avg_pct": round(vs_pct, 2),
                "building_avg_kwh": round(b_avg, 4),
                "type_avg_kwh": round(t_avg, 4),
            }

        return {
            "building_id": building_id or "ALL",
            "method_peak_excess_pct": round(max(peak_excess_pct, 0), 2),
            "method_anomaly_waste_kwh": round(wasted_kwh, 4),
            "method_anomaly_carbon_kg": round(carbon_reduction_kg, 4),
            "method_peer_compare": peer_info,
        }

    def query_anomaly_detail(self, anomaly_id: int) -> dict[str, Any] | None:
        context = next((item for item in self.anomalies if item["anomaly_id"] == anomaly_id), None)
        if not context:
            return None

        building_rows = self.by_building.get(context["building_id"], [])
        local_start = context["timestamp"] - dt.timedelta(hours=6)
        local_end = context["timestamp"] + dt.timedelta(hours=6)
        neighborhood = [r for r in building_rows if local_start <= r["timestamp"] <= local_end]
        timeline = [{"timestamp": to_iso(r["timestamp"]), "value": round(r["electricity_kwh"], 4)} for r in neighborhood]

        baseline_start = context["timestamp"] - dt.timedelta(hours=24)
        baseline_rows = [r for r in building_rows if baseline_start <= r["timestamp"] < context["timestamp"]]
        baseline_values = [r["electricity_kwh"] for r in baseline_rows]
        baseline_avg = sum(baseline_values) / len(baseline_values) if baseline_values else context["mean_kwh"]
        baseline_max = max(baseline_values) if baseline_values else context["mean_kwh"]
        baseline_min = min(baseline_values) if baseline_values else context["mean_kwh"]

        estimated_loss = max(context["electricity_kwh"] - baseline_avg, 0)
        peer_rows = [r for r in self.rows if r["timestamp"] == context["timestamp"] and r["building_type"] == context["building_type"]]
        peer_avg = (sum(r["electricity_kwh"] for r in peer_rows) / len(peer_rows)) if peer_rows else context["mean_kwh"]
        vs_peer_pct = ((context["electricity_kwh"] - peer_avg) / peer_avg * 100) if peer_avg else 0

        knowledge = self.dict_data.get(context["anomaly_type"], {})
        action_state = self._action_state(anomaly_id)
        note = self.query_anomaly_note(anomaly_id)
        anomaly_name = str(context.get("anomaly_name") or knowledge.get("name", context["anomaly_type"]))

        return {
            "anomaly": {
                **context,
                "timestamp": to_iso(context["timestamp"]),
                "anomaly_name": anomaly_name,
                "display_name": str(context.get("display_name") or anomaly_name),
                "status": action_state["status"],
            },
            "rule_explanation": {
                "rule_code": str(context.get("rule_code", context["anomaly_type"])),
                "rule_name": str(context.get("rule_name", anomaly_name)),
                "rule_summary": str(context.get("rule_summary", "")),
                "current_value": round(float(context.get("current_value", context["electricity_kwh"])), 4),
                "baseline_value": round(float(context.get("baseline_value", context["mean_kwh"])), 4),
                "threshold_value": round(float(context.get("threshold_value", context["threshold"])), 4),
                "deviation_pct": round(float(context.get("deviation_pct", 0.0)), 2),
                "trigger_window": str(context.get("trigger_window", to_iso(context["timestamp"]))),
                "trigger_count": int(context.get("trigger_count", 1) or 1),
                "time_scope_label": str(context.get("time_scope_label", "异常时段")),
            },
            "impact": {
                "estimated_loss_kwh": round(estimated_loss, 4),
                "wasted_kwh": round(estimated_loss, 4),
                "carbon_kg": round(estimated_loss * CARBON_FACTOR, 4),
            },
            "baseline_window": {
                "hours": 24,
                "start_time": to_iso(baseline_start),
                "end_time": to_iso(context["timestamp"]),
                "avg_kwh": round(baseline_avg, 4),
                "max_kwh": round(baseline_max, 4),
                "min_kwh": round(baseline_min, 4),
                "sample_count": len(baseline_rows),
            },
            "peer_compare": {
                "building_type": context["building_type"],
                "peer_avg_kwh": round(peer_avg, 4),
                "vs_peer_pct": round(vs_peer_pct, 2),
                "peer_sample_count": len(peer_rows),
            },
            "timeline": timeline,
            "recommended_actions": knowledge.get("steps", [])[:5],
            "prevention": knowledge.get("prevention", [])[:5],
            "processing_summary": {
                "latest_status": action_state["status"],
                "history_count": action_state["history_count"],
                "latest_note": action_state["last_note"],
                "assignee": action_state["assignee"],
                "last_action_at": action_state["last_action_at"],
            },
            "postmortem_note": note,
        }

    def _search_local_knowledge(self, anomaly_type: str, message: str, limit: int = 3) -> list[dict[str, str]]:
        if not self.knowledge_chunks:
            return []
        keywords = [k.lower() for k in self.dict_data.get(anomaly_type, {}).get("keywords", [])]
        extra_tokens = [t.lower() for t in message.replace("，", " ").replace(",", " ").split() if len(t) >= 2]
        tokens = [t for t in (keywords + extra_tokens) if t]
        if not tokens:
            return []

        scored: list[tuple[int, str, dict[str, Any]]] = []
        for chunk in self.knowledge_chunks:
            text = str(chunk.get("text", "")).lower()
            if not text:
                continue
            score = sum(1 for tk in tokens if tk in text)
            if score > 0:
                scored.append((score, str(chunk.get("chunk_id", "")), chunk))

        scored.sort(key=lambda x: (-x[0], x[1]))
        picked = []
        for _, _, chunk in scored[:limit]:
            picked.append(
                self._normalize_evidence_item(
                    chunk_id=chunk.get("chunk_id", ""),
                    title=chunk.get("title", ""),
                    section=chunk.get("chunk_id", ""),
                    excerpt=chunk.get("text", ""),
                    source_type="local_knowledge",
                )
            )
        return picked

    def _search_knowledge(
        self,
        anomaly_type: str,
        message: str,
        limit: int = 3,
        *,
        query_text: str | None = None,
    ) -> dict[str, Any]:
        search_query = str(query_text or message or anomaly_type).strip()

        # Run RAGFlow retrieval and local keyword search in parallel.
        # RAGFlow gets a 4-second hard deadline at the caller level regardless of its own HTTP timeout.
        # Local search is fast (<100ms) and always finishes in time, so there is zero extra wait on RAGFlow failure.
        with concurrent.futures.ThreadPoolExecutor(max_workers=2) as pool:
            ragflow_future = pool.submit(self._retrieve_ragflow_knowledge, search_query, limit)
            local_future = pool.submit(self._search_local_knowledge, anomaly_type, message, limit)

            try:
                ragflow_result = ragflow_future.result(timeout=4.0)
            except Exception:
                ragflow_result = {
                    "items": [],
                    "knowledge_source": "none",
                    "retrieval_hit_count": 0,
                    "retrieval_error_type": "timeout",
                    "ragflow_session_id": "",
                }

            if ragflow_result["items"]:
                return ragflow_result

            local_items = local_future.result()

        if local_items:
            return {
                "items": local_items,
                "knowledge_source": "local",
                "retrieval_hit_count": len(local_items),
                "retrieval_error_type": ragflow_result.get("retrieval_error_type", ""),
                "ragflow_session_id": "",
            }
        return {
            "items": [],
            "knowledge_source": "none",
            "retrieval_hit_count": 0,
            "retrieval_error_type": ragflow_result.get("retrieval_error_type", "empty_result"),
            "ragflow_session_id": "",
        }

    def _resolve_context(self, payload: dict[str, Any]) -> tuple[dict[str, Any] | None, str | None]:
        anomaly_id = payload.get("anomaly_id")
        building_id = str(payload.get("building_id", "")).strip() or None
        timestamp = parse_time(str(payload.get("timestamp", "")).strip() or None)

        context = None
        if anomaly_id:
            try:
                aid = int(anomaly_id)
            except (TypeError, ValueError):
                aid = None
            if aid is not None:
                context = next((item for item in self.anomalies if item["anomaly_id"] == aid), None)

        if context is None and building_id and timestamp:
            context = next(
                (
                    item
                    for item in self.anomalies
                    if item["building_id"] == building_id and abs((item["timestamp"] - timestamp).total_seconds()) <= 3600
                ),
                None,
            )

        if context is None and building_id:
            context = next((item for item in self.anomalies if item["building_id"] == building_id), None)

        return context, building_id

    def _diagnose_likely_systems(self, building_type: str, anomaly_type: str) -> list[str]:
        systems: list[str] = []
        anomaly_map = {
            "anomaly_spike": ["空调系统", "照明系统", "动力设备"],
            "anomaly_off_hours_load": ["照明系统", "新风系统", "待机设备", "热水设备"],
            "anomaly_sustained_high_load": ["空调主机", "循环水泵", "新风机组", "连续运行设备"],
            "anomaly_workhour_offline": ["配电回路", "控制器", "采集网关", "主机联动"],
            "anomaly_baseload_high": ["照明回路", "热水系统", "待机设备", "新风机组"],
            "anomaly_schedule_shift": ["BMS排程", "空调联动", "照明定时", "新风排程"],
        }
        building_text = str(building_type or "").lower()
        systems.extend(anomaly_map.get(anomaly_type, ["空调系统", "照明系统", "重点用能设备"]))
        if "lodging" in building_text or "residential" in building_text:
            systems.extend(["热水系统", "客房末端设备"])
        elif "office" in building_text or "education" in building_text:
            systems.extend(["办公插座负荷", "教室末端空调"])
        elif "parking" in building_text or "garage" in building_text:
            systems.extend(["通风排烟系统", "照明系统"])
        seen: set[str] = set()
        deduped: list[str] = []
        for item in systems:
            if item and item not in seen:
                deduped.append(item)
                seen.add(item)
        return deduped[:5]

    def _build_diagnose_context(self, payload: dict[str, Any]) -> tuple[dict[str, Any], str, dict[str, Any] | None]:
        context, fallback_building_id = self._resolve_context(payload)
        message = str(payload.get("message", "")).strip()
        anomaly_type = str(payload.get("anomaly_type", "")).strip() or (context["anomaly_type"] if context else self._type_from_keywords(message))
        building_id = (context["building_id"] if context else fallback_building_id) or None
        event_time = context["timestamp"] if context else parse_time(str(payload.get("timestamp", "")).strip() or None)
        building_meta = self.buildings_meta.get(building_id or "", {})
        building_name = (context["building_name"] if context else None) or building_meta.get("building_name") or "当前建筑"
        building_type = (context["building_type"] if context else None) or building_meta.get("building_type") or "portfolio"
        detail = self.query_anomaly_detail(int(context["anomaly_id"])) if context else None

        start_time = parse_time(str(payload.get("start_time", "")).strip() or None)
        end_time = parse_time(str(payload.get("end_time", "")).strip() or None)
        analysis_payload = {
            "building_id": building_id,
            "metric_type": "electricity",
            "start_time": to_iso(start_time) if start_time else None,
            "end_time": to_iso(end_time) if end_time else None,
        }
        analysis_context = self._build_analysis_context(analysis_payload)
        summary = analysis_context.get("summary") or {}
        insights = analysis_context.get("insights") or {}
        distribution = analysis_context.get("distribution") or {}
        compare = analysis_context.get("compare") or {}

        building_rows = self.by_building.get(building_id or "", [])
        before_window: list[dict[str, Any]] = []
        after_window: list[dict[str, Any]] = []
        same_hour_rows: list[dict[str, Any]] = []
        current_temp = None
        temp_band = "normal"
        if event_time and building_rows:
            before_start = event_time - dt.timedelta(hours=24)
            after_end = event_time + dt.timedelta(hours=24)
            before_window = [row for row in building_rows if before_start <= row["timestamp"] < event_time]
            after_window = [row for row in building_rows if event_time < row["timestamp"] <= after_end]
            same_hour_rows = [
                row
                for row in building_rows
                if row["timestamp"] < event_time and row["timestamp"].hour == event_time.hour
            ][-7:]
            site_id = self.building_site_map.get(building_id or "")
            weather_lookup = self.weather_by_site.get(site_id or "", {})
            weather_point = weather_lookup.get(event_time, {})
            current_temp = weather_point.get("temperature_c")
            temps = sorted(
                float(item["temperature_c"])
                for item in weather_lookup.values()
                if item.get("temperature_c") is not None
            )
            if current_temp is not None and temps:
                hot_threshold = temps[min(len(temps) - 1, max(0, int(len(temps) * 0.8)))]
                cold_threshold = temps[min(len(temps) - 1, max(0, int(len(temps) * 0.2)))]
                if current_temp >= hot_threshold:
                    temp_band = "hot"
                elif current_temp <= cold_threshold:
                    temp_band = "cold"

        before_stats = self._summarize_values([float(row["electricity_kwh"]) for row in before_window])
        after_stats = self._summarize_values([float(row["electricity_kwh"]) for row in after_window])
        same_hour_stats = self._summarize_values([float(row["electricity_kwh"]) for row in same_hour_rows])
        peer_group = compare.get("peer_group") or {}
        trend_findings = insights.get("trend_findings") or []
        weather_findings = insights.get("weather_findings") or []
        compare_findings = insights.get("compare_findings") or []
        night_base = distribution.get("night_base_load") or {}
        phenomenon_tags = []
        if context:
            phenomenon_tags.append(str(detail["anomaly"]["anomaly_name"]) if detail else anomaly_type)
        if before_stats.get("avg"):
            deviation_vs_before = ((float(context["electricity_kwh"]) - before_stats["avg"]) / before_stats["avg"] * 100) if context and before_stats["avg"] else 0.0
            if abs(deviation_vs_before) >= 25:
                phenomenon_tags.append("负荷明显偏离近24小时基线")
        if float(night_base.get("ratio_vs_avg_pct", 0) or 0) >= 65:
            phenomenon_tags.append("夜间基线偏高")
        if any("非工作时段" in f"{item.get('title', '')}{item.get('detail', '')}" for item in trend_findings if isinstance(item, dict)):
            phenomenon_tags.append("非工作时段负荷不降")
        if any("温度" in f"{item.get('title', '')}{item.get('detail', '')}" for item in weather_findings if isinstance(item, dict)):
            phenomenon_tags.append("天气联动明显")
        if any("同类" in f"{item.get('title', '')}{item.get('detail', '')}" for item in compare_findings if isinstance(item, dict)):
            phenomenon_tags.append("同类对照存在差异")
        seen_tags: set[str] = set()
        phenomenon_tags = [item for item in phenomenon_tags if item and not (item in seen_tags or seen_tags.add(item))]

        data_evidence: list[dict[str, Any]] = []
        if context and detail:
            data_evidence.append(
                {
                    "title": "异常点与24h基线",
                    "detail": f"异常时刻负荷 {round(float(context['electricity_kwh']), 2)} kWh，24h基线均值 {detail['baseline_window']['avg_kwh']} kWh，偏差 {round(float(context['deviation_pct']), 2)}%。",
                    "source_type": "data_signal",
                }
            )
        if same_hour_stats.get("avg"):
            data_evidence.append(
                {
                    "title": "同小时历史基线",
                    "detail": f"近7个同小时均值 {same_hour_stats['avg']} kWh，可用于排除单纯时段性波动。",
                    "source_type": "data_signal",
                }
            )
        if before_stats.get("avg") or after_stats.get("avg"):
            data_evidence.append(
                {
                    "title": "异常前后窗口",
                    "detail": f"异常前24h均值 {before_stats.get('avg', 0)} kWh，异常后24h均值 {after_stats.get('avg', 0)} kWh。",
                    "source_type": "data_signal",
                }
            )
        if peer_group:
            data_evidence.append(
                {
                    "title": "同类建筑对照",
                    "detail": f"同类样本 {peer_group.get('peer_count', 0)} 栋，偏离同类均值 {peer_group.get('gap_pct', 0)}%，同类百分位 {peer_group.get('peer_percentile', 0)}。",
                    "source_type": "data_compare",
                }
            )
        if current_temp is not None:
            data_evidence.append(
                {
                    "title": "天气条件",
                    "detail": f"异常时段温度约 {round(float(current_temp), 1)}°C，当前判定为 {('高温区间' if temp_band == 'hot' else '低温区间' if temp_band == 'cold' else '常规温度区间')}。",
                    "source_type": "data_weather",
                }
            )
        if detail:
            processing = detail.get("processing_summary") or {}
            data_evidence.append(
                {
                    "title": "处理闭环状态",
                    "detail": f"当前状态 {processing.get('latest_status', '-') }，历史处理 {processing.get('history_count', 0)} 次，处理人 {processing.get('assignee') or '未指派'}。",
                    "source_type": "ops_status",
                }
            )

        return (
            {
                "anomaly_id": context["anomaly_id"] if context else None,
                "building_id": building_id,
                "building_name": building_name,
                "building_type": building_type,
                "timestamp": to_iso(event_time) if event_time else None,
                "anomaly_type": anomaly_type,
                "anomaly_name": (detail or {}).get("anomaly", {}).get("anomaly_name") if detail else "",
                "status": (detail or {}).get("processing_summary", {}).get("latest_status") if detail else None,
                "deviation_pct": round(float(context["deviation_pct"]), 2) if context else 0.0,
                "value_kwh": round(float(context["electricity_kwh"]), 2) if context else 0.0,
                "analysis_scope": {
                    "start_time": analysis_context.get("start_time"),
                    "end_time": analysis_context.get("end_time"),
                    "metric_type": analysis_context.get("metric_type"),
                },
                "window_context": {
                    "baseline_24h_avg_kwh": (detail or {}).get("baseline_window", {}).get("avg_kwh"),
                    "same_hour_avg_kwh": same_hour_stats.get("avg", 0.0),
                    "before_24h_avg_kwh": before_stats.get("avg", 0.0),
                    "after_24h_avg_kwh": after_stats.get("avg", 0.0),
                    "point_count": (analysis_context.get("summary") or {}).get("point_count", 0),
                },
                "analysis_background": {
                    "trend_findings": trend_findings[:3],
                    "weather_findings": weather_findings[:3],
                    "compare_findings": compare_findings[:3],
                    "night_base_load": night_base,
                },
                "peer_context": peer_group,
                "weather_context": {
                    "temperature_c": current_temp,
                    "temperature_band": temp_band,
                    "temperature_correlation": (analysis_context.get("trend") or {}).get("summary", {}).get("temperature_correlation"),
                },
                "operations_context": {
                    "processing_summary": (detail or {}).get("processing_summary"),
                    "postmortem_note": (detail or {}).get("postmortem_note"),
                },
                "phenomenon_tags": phenomenon_tags[:6],
                "likely_systems": self._diagnose_likely_systems(building_type, anomaly_type),
                "data_evidence": data_evidence,
            },
            anomaly_type,
            context,
        )

    def _diagnose_by_template(self, payload: dict[str, Any]) -> dict[str, Any]:
        message = str(payload.get("message", "")).strip()
        diagnose_context, chosen_type, raw_context = self._build_diagnose_context(payload)

        knowledge = self.dict_data.get(chosen_type, self.dict_data.get("anomaly_spike", {}))
        anomaly_name = knowledge.get("name", chosen_type)
        causes = self._clean_text_list(knowledge.get("possible_causes", []), max_items=4)
        steps = self._clean_text_list(knowledge.get("steps", []), max_items=4)
        prevention = self._clean_text_list(knowledge.get("prevention", []), max_items=4)
        retrieval = self._search_knowledge(
            chosen_type,
            message,
            limit=3,
            query_text=self._build_diagnose_knowledge_query(anomaly_name, message, diagnose_context),
        )
        evidence = retrieval["items"]

        enriched_causes = self._build_diagnosis_causes(diagnose_context, raw_context, causes)
        enriched_steps = self._build_diagnosis_steps(diagnose_context, raw_context, steps)
        enriched_prevention = self._build_diagnosis_prevention(diagnose_context, prevention)
        recommended_actions = self._build_diagnosis_actions(diagnose_context, chosen_type, enriched_steps, steps[:3] if steps else prevention[:3])

        if raw_context:
            conclusion = (
                f"{diagnose_context['building_name']} 在 {diagnose_context['timestamp']} 出现{anomaly_name}，"
                f"异常点负荷 {diagnose_context['value_kwh']} kWh，较24小时基线 "
                f"{diagnose_context['window_context'].get('baseline_24h_avg_kwh') or raw_context['mean_kwh']} kWh 偏离 {diagnose_context['deviation_pct']}%。"
                f"结合同小时基线、建筑类型和知识证据，这更像是 { '、'.join(diagnose_context.get('likely_systems', [])[:2]) or '关键用能系统' } 在该时段运行异常，而不是普通波动。"
            )
        else:
            conclusion = (
                f"根据你提供的问题，系统判定为 {anomaly_name} 场景。"
                "建议先执行基础排查步骤，再结合实时曲线和知识依据确认是否恢复。"
            )

        if raw_context and evidence:
            confidence = 0.85
        elif raw_context:
            confidence = 0.72
        elif evidence:
            confidence = 0.58
        else:
            confidence = 0.45

        risk_level = "high" if confidence >= 0.8 else "medium" if confidence >= 0.6 else "low"
        return {
            "diagnosis": {
                "anomaly_type": chosen_type,
                "anomaly_name": anomaly_name,
                "conclusion": conclusion,
                "causes": enriched_causes,
                "possible_causes": enriched_causes,
                "steps": enriched_steps,
                "prevention": enriched_prevention,
                "recommended_actions": recommended_actions,
                "data_evidence": diagnose_context.get("data_evidence", []),
                "evidence": evidence,
                "confidence": round(confidence, 2),
                "risk_level": risk_level,
                "knowledge_source": retrieval["knowledge_source"],
                "retrieval_hit_count": retrieval["retrieval_hit_count"],
                "retrieval_error_type": retrieval["retrieval_error_type"],
                "ragflow_session_id": retrieval.get("ragflow_session_id", ""),
            },
            "context": diagnose_context,
        }

    def diagnose_stream_events(
        self, payload: dict[str, Any]
    ) -> Generator[tuple[str, dict[str, Any]], None, None]:
        """Yield (event_name, data_dict) tuples for SSE streaming diagnosis."""
        template_result = self._diagnose_by_template(payload)
        diag = template_result["diagnosis"]
        context = template_result["context"]
        trace_id = uuid.uuid4().hex
        requested_provider = str(payload.get("provider", "auto")).strip().lower() or "auto"

        yield "template", diag

        api_key = os.getenv("OPENAI_API_KEY", "").strip()
        if not api_key:
            yield "done", {
                **diag,
                "provider": "template_provider",
                "requested_provider": requested_provider,
                "fallback_used": True,
                "degrade_message": self._friendly_degrade_message("LLM provider not configured"),
                "trace_id": trace_id,
                "latency_ms": 0,
            }
            return

        base_url = os.getenv("OPENAI_BASE_URL", "https://api.deepseek.com").strip() or "https://api.deepseek.com"
        model = os.getenv("OPENAI_MODEL", "deepseek-chat").strip() or "deepseek-chat"
        base_timeout_sec = float(os.getenv("OPENAI_TIMEOUT_SEC", "12"))
        timeout_sec = float(os.getenv("OPENAI_DIAGNOSE_TIMEOUT_SEC", str(max(base_timeout_sec, 45.0))))

        evidence_text = "\n".join(
            f"- {x.get('title', '')}: {x.get('excerpt', '')}"
            for x in diag.get("evidence", [])[:3]
            if isinstance(x, dict)
        ).strip()

        system_prompt = (
            "你是建筑能源运维诊断助手。请输出可直接展示在运维工作台上的中文诊断结论。"
            "要求比普通数据复述更深入，必须说明异常意味着什么、优先怀疑哪些系统、为什么要这样判断。"
            "结论要引用建筑名称、异常时间、偏差值、基线或同小时对比，不要输出JSON、标题或列表。"
        )
        user_prompt = (
            f"异常类型: {diag.get('anomaly_name', '')}\n"
            f"建筑: {context.get('building_name', '')}，"
            f"时间: {context.get('timestamp', '')}，"
            f"负荷 {context.get('value_kwh', '')} kWh，"
            f"偏离 {context.get('deviation_pct', '')}%\n"
            f"模板结论: {diag.get('conclusion', '')}\n"
            f"知识证据:\n{evidence_text or '无'}\n"
            "请给出更专业的诊断结论，引用具体数值。"
        )

        start = time.perf_counter()
        try:
            llm_provider = self.providers["llm"]
            for token in llm_provider._call_chat_completion_stream(
                base_url=base_url,
                api_key=api_key,
                model=model,
                timeout_sec=timeout_sec,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
            ):
                yield "token", {"text": token}
            latency_ms = int((time.perf_counter() - start) * 1000)
            yield "done", {
                **diag,
                "conclusion": "",
                "provider": "llm_provider",
                "requested_provider": requested_provider,
                "fallback_used": False,
                "trace_id": trace_id,
                "latency_ms": latency_ms,
            }
        except Exception as exc:
            yield "done", {
                **diag,
                "provider": "template_provider",
                "requested_provider": requested_provider,
                "fallback_used": True,
                "degrade_message": self._friendly_degrade_message(str(exc)),
                "trace_id": trace_id,
                "latency_ms": int((time.perf_counter() - start) * 1000),
            }

    def _build_analysis_context(
        self,
        payload: dict[str, Any],
    ) -> dict[str, Any]:
        metric_type = str(payload.get("metric_type", "electricity")).strip().lower() or "electricity"
        building_id = str(payload.get("building_id", "")).strip() or None
        start_time = parse_time(str(payload.get("start_time", "")).strip() or None)
        end_time = parse_time(str(payload.get("end_time", "")).strip() or None)

        summary = self.query_analysis_summary(building_id, start_time, end_time, metric_type)
        trend = self.query_analysis_trend(building_id, start_time, end_time, metric_type)
        distribution = self.query_analysis_distribution(building_id, start_time, end_time, metric_type)
        compare = self.query_analysis_compare(building_id, start_time, end_time, metric_type)
        payload_insights = payload.get("insights")
        insights = payload_insights if isinstance(payload_insights, dict) else self.query_analysis_insights(building_id, start_time, end_time, metric_type)
        building = compare.get("building") or {}

        return {
            "building_id": building_id or building.get("building_id"),
            "building_name": building.get("building_name") or (self.buildings_meta.get(building_id or "", {}) or {}).get("building_name"),
            "building_type": building.get("building_type") or (self.buildings_meta.get(building_id or "", {}) or {}).get("building_type"),
            "metric_type": metric_type,
            "start_time": to_iso(start_time) if start_time else None,
            "end_time": to_iso(end_time) if end_time else None,
            "summary": summary,
            "trend": trend,
            "distribution": distribution,
            "compare": compare,
            "insights": insights,
        }

    def _truncate_text(self, value: Any, limit: int = 160) -> str:
        text = str(value or "").strip()
        if len(text) <= limit:
            return text
        return f"{text[:limit].rstrip()}..."

    def _clean_text_list(self, items: Any, max_items: int = 5) -> list[str]:
        if not isinstance(items, list):
            return []
        cleaned: list[str] = []
        seen: set[str] = set()
        for item in items:
            text = self._truncate_text(item, 180)
            if not text or text in seen:
                continue
            cleaned.append(text)
            seen.add(text)
            if len(cleaned) >= max_items:
                break
        return cleaned

    def _merge_text_lists(self, primary: Any, fallback: Any, *, max_items: int = 5, min_items: int = 1) -> list[str]:
        merged = self._clean_text_list(primary, max_items=max_items)
        if len(merged) < min_items:
            for item in self._clean_text_list(fallback, max_items=max_items):
                if item not in merged:
                    merged.append(item)
                if len(merged) >= max_items:
                    break
        return merged[:max_items]

    def _build_diagnosis_causes(
        self,
        diagnose_context: dict[str, Any],
        raw_context: dict[str, Any] | None,
        base_causes: list[str],
    ) -> list[str]:
        causes = list(base_causes or [])
        window_context = diagnose_context.get("window_context") or {}
        weather_context = diagnose_context.get("weather_context") or {}
        peer_context = diagnose_context.get("peer_context") or {}
        likely_systems = diagnose_context.get("likely_systems") or []
        if raw_context and window_context.get("baseline_24h_avg_kwh") is not None:
            causes.append(
                f"异常点负荷 {diagnose_context.get('value_kwh', 0)} kWh，相对24小时基线 {window_context.get('baseline_24h_avg_kwh')} kWh 偏高 {diagnose_context.get('deviation_pct', 0)}%，说明不是普通时段波动。"
            )
        if window_context.get("same_hour_avg_kwh"):
            causes.append(f"近7个同小时均值仅 {window_context.get('same_hour_avg_kwh')} kWh，当前时点明显偏离常规同小时负荷。")
        if weather_context.get("temperature_band") == "hot":
            causes.append("异常发生在高温区间，空调主机、新风机组或末端设备可能因负荷上升而共同抬高总电量。")
        elif weather_context.get("temperature_band") == "cold":
            causes.append("异常发生在低温区间，采暖、伴热或防冻运行策略可能推高对应时段能耗。")
        if float(peer_context.get("gap_pct", 0) or 0) >= 10:
            causes.append(f"当前建筑较同类均值偏高 {round(float(peer_context.get('gap_pct', 0) or 0), 2)}%，问题更可能来自本建筑运行策略或局部设备异常。")
        if likely_systems:
            causes.append(f"结合建筑类型与异常类型，建议优先怀疑 { '、'.join(likely_systems[:3]) } 的运行状态或联动策略。")
        return self._merge_text_lists(causes, base_causes, max_items=4, min_items=3)

    def _build_diagnosis_steps(
        self,
        diagnose_context: dict[str, Any],
        raw_context: dict[str, Any] | None,
        base_steps: list[str],
    ) -> list[str]:
        steps = list(base_steps or [])
        likely_systems = diagnose_context.get("likely_systems") or []
        ts = diagnose_context.get("timestamp") or "当前异常时段"
        steps.insert(0, f"先核对 {ts} 前后 1 小时的总表、分项表和分路数据，确认异常不是采集抖动或单点坏值。")
        if likely_systems:
            steps.append(f"按 { '、'.join(likely_systems[:2]) } 的顺序检查设备启停、设定点、手自动状态和联动信号。")
        if raw_context:
            steps.append("对照近7个同小时曲线与前后24小时窗口，确认异常是瞬时触发还是持续运行问题。")
        steps.append("同步查看 BMS 排程、临时加班记录和现场设备运行台账，判断是否存在计划外启停或联动遗漏。")
        return self._merge_text_lists(steps, base_steps, max_items=4, min_items=4)

    def _build_diagnosis_actions(
        self,
        diagnose_context: dict[str, Any],
        chosen_type: str,
        step_items: list[str],
        base_actions: list[str],
    ) -> list[str]:
        actions = list(base_actions or [])
        if chosen_type in {"anomaly_spike", "anomaly_sustained_high_load", "anomaly_off_hours_load"}:
            actions.append("先对异常时段相关空调、新风、照明或动力回路执行限时核减，避免异常负荷继续放大。")
        if chosen_type == "anomaly_workhour_offline":
            actions.append("先恢复关键主机、控制器和通信链路，再确认现场末端是否已恢复正常运行。")
        if chosen_type in {"anomaly_baseload_high", "anomaly_schedule_shift"}:
            actions.append("立即复核夜间/启停排程，取消不必要的常开策略，并保留调参前后曲线截图用于复盘。")
        if step_items:
            actions.append(f"优先执行：{step_items[0]}")
        return self._merge_text_lists(actions, base_actions, max_items=3, min_items=3)

    def _build_diagnosis_prevention(self, diagnose_context: dict[str, Any], base_prevention: list[str]) -> list[str]:
        prevention = list(base_prevention or [])
        likely_systems = diagnose_context.get("likely_systems") or []
        prevention.append("把异常时段的排程、设定点和现场操作记录纳入每周复盘，避免同类问题重复发生。")
        prevention.append("针对夜间基线、异常峰值和工作时段停运建立自动阈值告警，并与分项电表趋势联动校验。")
        if likely_systems:
            prevention.append(f"对 { '、'.join(likely_systems[:2]) } 建立月度巡检与季节切换前专项检查清单。")
        return self._merge_text_lists(prevention, base_prevention, max_items=4, min_items=3)

    def _build_analysis_findings(self, insights: dict[str, Any], summary: dict[str, Any]) -> list[str]:
        findings: list[str] = []
        for item in (insights.get("trend_findings") or [])[:3]:
            findings.append(f"{item['title']}：{item['detail']}")
        for item in (insights.get("weather_findings") or [])[:2]:
            findings.append(f"{item['title']}：{item['detail']}")
        for item in (insights.get("compare_findings") or [])[:2]:
            findings.append(f"{item['title']}：{item['detail']}")
        if summary.get("peak_value") and summary.get("avg_value"):
            findings.append(f"峰值 {summary['peak_value']} kWh，高于均值 {summary['avg_value']} kWh，说明当前窗口存在明显的峰段管理空间。")
        return self._merge_text_lists(findings, [], max_items=5, min_items=3)

    def _build_analysis_possible_causes(self, insights: dict[str, Any]) -> list[str]:
        possible_causes: list[str] = []
        for item in insights.get("weather_findings") or []:
            text = str(item.get("detail", ""))
            if "热天气" in text or "温度正相关" in text:
                possible_causes.append("负荷与外气温存在正相关，空调、新风或末端策略可能在高温时段放大能耗。")
            elif "低温" in text or "反向影响" in text:
                possible_causes.append("低温时段相关负荷上升，可能与采暖、伴热或防冻运行有关。")
        for item in insights.get("trend_findings") or []:
            text = f"{item.get('title', '')}{item.get('detail', '')}"
            if "非工作时段" in text:
                possible_causes.append("非工作时段负荷回落不充分，说明夜间待机、常开设备或排程收敛不足。")
            if "波动" in text or "启停" in text:
                possible_causes.append("设备启停节奏偏密，可能造成短周期波动和无效能耗叠加。")
        for item in insights.get("compare_findings") or []:
            text = f"{item.get('title', '')}{item.get('detail', '')}"
            if "同类" in text and "高" in text:
                possible_causes.append("与同类建筑差距偏大，问题更可能来自本建筑控制策略、运行时长或局部设备效率。")
        possible_causes.append("建议把异常窗口、夜间基线和同类差距结合起来看，优先排查最能解释这三类信号的系统。")
        return self._merge_text_lists(possible_causes, [], max_items=4, min_items=3)

    def _build_analysis_energy_saving(self, insights: dict[str, Any]) -> list[str]:
        energy_saving: list[str] = []
        for item in (insights.get("saving_opportunities") or [])[:4]:
            estimated_value = round(float(item.get("estimated_kwh", item.get("estimated_loss_kwh", 0)) or 0), 2)
            energy_saving.append(f"{item['title']}：{item['detail']} 可优先作为节能专项，当前窗口影响估算 {estimated_value} kWh。")
        if not energy_saving:
            energy_saving.append("当前窗口未形成高置信度节能机会，建议继续跟踪高峰和夜间基线后再组织专项优化。")
        return self._merge_text_lists(energy_saving, [], max_items=4, min_items=3)

    def _build_analysis_operations(self, insights: dict[str, Any]) -> list[str]:
        operations: list[str] = []
        for item in (insights.get("anomaly_windows") or [])[:3]:
            operations.append(f"先复核 {item['timestamp']} 的 {item['anomaly_name']}，偏差 {item['deviation_pct']}%，影响约 {item['estimated_loss_kwh']} kWh，并核查对应设备工况与排程。")
        operations.append("把高峰时段排程、夜间基线和异常窗口放在同一张时间轴上核对，优先确认哪些负荷是可控的。")
        operations.append("对已经识别出的异常窗口补充分项电表、BMS 操作记录和现场值班记录，形成可闭环的问题清单。")
        return self._merge_text_lists(operations, [], max_items=4, min_items=3)

    def _select_relevant_saving_opportunities(
        self,
        opportunities: list[dict[str, Any]],
        message: str,
        focus: str | None = None,
        limit: int = 3,
    ) -> list[dict[str, Any]]:
        items = list(opportunities or [])
        if not items:
            return []
        focus_text = str(focus or "").strip()
        if focus_text:
            focus_map = {
                "night_baseload": "夜间基线负荷优化",
                "anomaly_waste": "异常浪费回收",
                "peer_gap": "同类差距收敛",
                "peak_strategy": "高峰时段策略优化",
            }
            target_title = focus_map.get(focus_text, focus_text)
            focused = [item for item in items if target_title in str(item.get("title", ""))]
            if focused:
                return focused[:limit]
        question = str(message or "").strip().lower()
        if not question:
            return items[:limit]

        keyword_groups = [
            ["夜间", "基线", "待机", "常开", "night"],
            ["异常", "浪费", "突增", "高负荷", "回收", "anomaly"],
            ["同类", "对标", "差距", "peer", "benchmark"],
            ["峰时", "高峰", "时段", "错峰", "排程", "schedule", "peak"],
        ]

        ranked: list[tuple[int, float, dict[str, Any]]] = []
        for item in items:
            title = str(item.get("title", ""))
            detail = str(item.get("detail", ""))
            corpus = f"{title} {detail}".lower()
            score = 0
            for group in keyword_groups:
                hits = sum(1 for keyword in group if keyword and keyword.lower() in question and keyword.lower() in corpus)
                if hits:
                    score += hits * 3
            parts = [part for part in re.split(r"[\s,，、:：()（）]+", corpus) if len(part) >= 2]
            score += sum(1 for part in parts[:10] if part in question)
            estimated_value = float(item.get("estimated_kwh", item.get("estimated_loss_kwh", 0)) or 0)
            ranked.append((score, estimated_value, item))

        ranked.sort(key=lambda x: (-x[0], -x[1]))
        matched = [item for score, _, item in ranked if score > 0]
        if not matched:
            return items[:limit]
        best_score = max(score for score, _, _ in ranked)
        best_items = [item for score, _, item in ranked if score == best_score and score > 0]
        return best_items[:limit]

    def _build_analysis_prompt_context(
        self,
        payload: dict[str, Any],
        context: dict[str, Any],
        analysis_seed: dict[str, Any],
    ) -> dict[str, Any]:
        summary = context.get("summary") or {}
        insights = context.get("insights") or {}
        compare = context.get("compare") or {}
        trend_snapshot = payload.get("trend_snapshot") if isinstance(payload.get("trend_snapshot"), dict) else {}
        distribution_snapshot = payload.get("distribution_snapshot") if isinstance(payload.get("distribution_snapshot"), dict) else {}
        compare_snapshot = payload.get("compare_snapshot") if isinstance(payload.get("compare_snapshot"), dict) else {}

        def compact_findings(items: list[Any], limit: int = 4) -> list[dict[str, str]]:
            compacted = []
            for item in (items or [])[:limit]:
                if not isinstance(item, dict):
                    continue
                compacted.append(
                    {
                        "title": self._truncate_text(item.get("title", ""), 80),
                        "detail": self._truncate_text(item.get("detail", ""), 180),
                        "severity": str(item.get("severity", "")).strip() or "info",
                    }
                )
            return compacted

        def compact_opportunities(items: list[Any], limit: int = 4) -> list[dict[str, Any]]:
            compacted = []
            for item in (items or [])[:limit]:
                if not isinstance(item, dict):
                    continue
                compacted.append(
                    {
                        "title": self._truncate_text(item.get("title", ""), 80),
                        "detail": self._truncate_text(item.get("detail", ""), 180),
                        "priority": str(item.get("priority", "")).strip() or "info",
                        "estimated_loss_kwh": item.get("estimated_loss_kwh", 0),
                    }
                )
            return compacted

        def compact_windows(items: list[Any], limit: int = 5) -> list[dict[str, Any]]:
            compacted = []
            for item in (items or [])[:limit]:
                if not isinstance(item, dict):
                    continue
                compacted.append(
                    {
                        "timestamp": item.get("timestamp"),
                        "anomaly_name": self._truncate_text(item.get("anomaly_name", ""), 60),
                        "severity": item.get("severity"),
                        "deviation_pct": item.get("deviation_pct"),
                        "estimated_loss_kwh": item.get("estimated_loss_kwh"),
                    }
                )
            return compacted

        compact_evidence = []
        for idx, item in enumerate((analysis_seed.get("evidence") or [])[:3], start=1):
            if isinstance(item, dict):
                compact_evidence.append(
                    {
                        "chunk_id": item.get("chunk_id", f"seed-{idx}"),
                        "title": self._truncate_text(item.get("title", ""), 80),
                        "section": self._truncate_text(item.get("section", ""), 80),
                        "excerpt": self._truncate_text(item.get("excerpt", ""), 180),
                    }
                )
            else:
                compact_evidence.append(
                    {
                        "chunk_id": f"seed-{idx}",
                        "title": "seed",
                        "section": "",
                        "excerpt": self._truncate_text(item, 180),
                    }
                )

        return {
            "building": {
                "building_id": context.get("building_id"),
                "building_name": context.get("building_name"),
                "building_type": context.get("building_type"),
                "metric_type": context.get("metric_type"),
                "start_time": context.get("start_time"),
                "end_time": context.get("end_time"),
            },
            "summary": {
                "metric_label": summary.get("metric_label"),
                "unit": summary.get("unit"),
                "point_count": summary.get("point_count"),
                "total_value": summary.get("total_value"),
                "avg_value": summary.get("avg_value"),
                "peak_value": summary.get("peak_value"),
                "volatility_pct": summary.get("volatility_pct"),
                "anomaly_count": summary.get("anomaly_count"),
                "working_hour_avg": summary.get("working_hour_avg"),
                "off_hour_avg": summary.get("off_hour_avg"),
            },
            "trend_snapshot": {
                "window_change_pct": trend_snapshot.get("window_change_pct", context.get("trend", {}).get("summary", {}).get("window_change_pct")),
                "temperature_correlation": trend_snapshot.get("temperature_correlation", context.get("trend", {}).get("summary", {}).get("temperature_correlation")),
            },
            "distribution_snapshot": {
                "weekday_peak_hours": distribution_snapshot.get("weekday_peak_hours") or (context.get("distribution", {}).get("weekday_peak_hours") or [])[:4],
                "night_base_load": distribution_snapshot.get("night_base_load") or context.get("distribution", {}).get("night_base_load") or {},
            },
            "compare_snapshot": compare_snapshot or compare.get("peer_group") or {},
            "insight_snapshot": {
                "scope_summary": insights.get("scope_summary") or {},
                "trend_findings": compact_findings(insights.get("trend_findings") or []),
                "weather_findings": compact_findings(insights.get("weather_findings") or []),
                "compare_findings": compact_findings(insights.get("compare_findings") or []),
                "saving_opportunities": compact_opportunities(insights.get("saving_opportunities") or []),
                "anomaly_windows": compact_windows(insights.get("anomaly_windows") or []),
            },
            "analysis_seed": {
                "summary": self._truncate_text(analysis_seed.get("summary", ""), 220),
                "findings": [self._truncate_text(item, 160) for item in (analysis_seed.get("findings") or [])[:5]],
                "possible_causes": [self._truncate_text(item, 160) for item in (analysis_seed.get("possible_causes") or [])[:4]],
                "energy_saving_suggestions": [self._truncate_text(item, 160) for item in (analysis_seed.get("energy_saving_suggestions") or [])[:4]],
                "operations_suggestions": [self._truncate_text(item, 160) for item in (analysis_seed.get("operations_suggestions") or [])[:4]],
                "evidence": compact_evidence,
            },
        }

    def _compact_analysis_prompt(self, ctx: dict[str, Any], max_chars: int = 6000) -> str:
        """Return a JSON string of prompt_context trimmed to max_chars.

        Priority order (drop later items first when over budget):
          1. building + summary (always kept)
          2. trend_snapshot + distribution_snapshot
          3. insight_snapshot.trend_findings (top 3)
          4. insight_snapshot.weather_findings (top 2)
          5. insight_snapshot.anomaly_windows (top 3) → drop if still over
          6. insight_snapshot.compare_findings → drop if still over
          7. analysis_seed.evidence excerpts → truncate to 60 chars each
        """
        import copy
        c = copy.deepcopy(ctx)

        def _try(obj: dict[str, Any]) -> str:
            return json.dumps(obj, ensure_ascii=False)

        if len(_try(c)) <= max_chars:
            return _try(c)

        # Trim anomaly_windows
        snap = c.get("insight_snapshot", {})
        snap["anomaly_windows"] = snap.get("anomaly_windows", [])[:3]
        if len(_try(c)) <= max_chars:
            return _try(c)

        # Drop compare_findings
        snap["compare_findings"] = []
        if len(_try(c)) <= max_chars:
            return _try(c)

        # Trim weather findings
        snap["weather_findings"] = snap.get("weather_findings", [])[:2]
        if len(_try(c)) <= max_chars:
            return _try(c)

        # Trim trend findings
        snap["trend_findings"] = snap.get("trend_findings", [])[:2]
        if len(_try(c)) <= max_chars:
            return _try(c)

        # Truncate evidence excerpts
        for ev in c.get("analysis_seed", {}).get("evidence", []):
            if isinstance(ev, dict):
                ev["excerpt"] = str(ev.get("excerpt", ""))[:60]
        if len(_try(c)) <= max_chars:
            return _try(c)

        # Last resort: drop saving_opportunities
        snap["saving_opportunities"] = snap.get("saving_opportunities", [])[:2]
        return _try(c)

    def _llm_template_fallback_enabled(self) -> bool:
        raw = str(os.getenv("LLM_ENABLE_TEMPLATE_FALLBACK", "0")).strip().lower()
        return raw in {"1", "true", "yes", "on"}

    def _raise_llm_failure(self, preferred: str, error_message: str | None) -> None:
        if preferred == "template":
            return
        message = str(error_message or "").strip() or "LLM request failed"
        raise RuntimeError(message)

    def _analyze_by_template(self, payload: dict[str, Any]) -> dict[str, Any]:
        context = self._build_analysis_context(payload)
        summary = context["summary"]
        insights = context["insights"]
        building_name = context.get("building_name") or "当前建筑"
        metric_label = "电力"
        findings = self._build_analysis_findings(insights, summary)
        possible_causes = self._build_analysis_possible_causes(insights)
        energy_saving = self._build_analysis_energy_saving(insights)
        operations = self._build_analysis_operations(insights)

        message = str(payload.get("message", "")).strip() or f"{building_name} {metric_label} 分析"
        selected_saving_opportunities = self._select_relevant_saving_opportunities(
            insights.get("saving_opportunities") or [],
            message,
            payload.get("saving_focus"),
            limit=3,
        )
        retrieval = self._search_knowledge(
            "anomaly_sustained_high_load",
            message,
            limit=3,
            query_text=self._build_analysis_knowledge_query(payload, context, insights),
        )
        evidence = retrieval["items"]

        scope = insights.get("scope_summary") or {}
        summary_text = (
            f"{building_name} 当前{metric_label}分析覆盖 {scope.get('data_start_time') or '-'} 至 {scope.get('data_end_time') or '-'}，"
            f"共 {scope.get('point_count', 0)} 个数据点；均值 {summary['avg_value']} kWh，"
            f"峰值 {summary['peak_value']} kWh，波动率 {summary['volatility_pct']}%。"
            f"当前结论已结合趋势、同类差距、异常窗口和节能机会，可直接用于汇报与运维排查。"
        )

        return {
            "analysis": {
                "summary": summary_text,
                "findings": findings,
                "possible_causes": possible_causes,
                "energy_saving_suggestions": energy_saving,
                "operations_suggestions": operations,
                "saving_opportunities": selected_saving_opportunities,
                "evidence": evidence,
                "knowledge_source": retrieval["knowledge_source"],
                "retrieval_hit_count": retrieval["retrieval_hit_count"],
                "retrieval_error_type": retrieval["retrieval_error_type"],
                "ragflow_session_id": retrieval.get("ragflow_session_id", ""),
            },
            "context": context,
        }

    def analyze(self, payload: dict[str, Any]) -> dict[str, Any]:
        preferred = str(payload.get("provider", "auto")).strip().lower()
        if preferred not in {"template", "llm", "auto"}:
            preferred = "auto"

        start = time.perf_counter()
        fallback_used = False
        error_message = None
        trace_id = uuid.uuid4().hex
        template_result = self._analyze_by_template(payload)
        result = template_result
        provider_name = "template_provider"
        fallback_enabled = self._llm_template_fallback_enabled()

        try:
            if preferred in {"llm", "auto"}:
                llm_provider = self.providers["llm"]
                if payload.get("simulate_llm_failure"):
                    raise RuntimeError("Simulated llm failure")
                api_key = os.getenv("OPENAI_API_KEY", "").strip()
                if not api_key:
                    raise RuntimeError("LLM provider not configured")

                base_url = os.getenv("OPENAI_BASE_URL", "https://api.deepseek.com").strip() or "https://api.deepseek.com"
                model = os.getenv("OPENAI_MODEL", "deepseek-chat").strip() or "deepseek-chat"
                base_timeout_sec = float(os.getenv("OPENAI_TIMEOUT_SEC", "12"))
                timeout_sec = float(os.getenv("OPENAI_ANALYZE_TIMEOUT_SEC", str(max(base_timeout_sec, 45.0))))
                max_retries = int(os.getenv("OPENAI_MAX_RETRIES", "2"))
                context = template_result["context"]
                analysis_seed = template_result["analysis"]
                user_question = str(payload.get("message", "")).strip() or "请围绕当前筛选范围输出一份可直接用于汇报的分析结论。"
                user_question = self._truncate_text(user_question, 180)
                system_prompt = (
                    "你是建筑能源分析助手，面向楼宇能源管理和运维答辩场景。"
                    "你必须只输出一个JSON对象，不要输出任何解释、前后缀或Markdown。"
                    "JSON字段必须包含：summary, findings, possible_causes, energy_saving_suggestions, operations_suggestions, evidence。"
                    "findings/possible_causes/energy_saving_suggestions/operations_suggestions/evidence 必须是数组，每个数组至少包含1条内容，不允许空数组。"
                    "findings 最多5条，possible_causes 最多4条，energy_saving_suggestions 最多3条，operations_suggestions 最多3条，evidence 最多3条。"
                    "请使用中文；不要输出空泛套话；summary必须引用建筑名称和时间范围；结论必须引用趋势、同类对比、天气联动或异常窗口中的具体数值。"
                    "建议动作尽量落到具体时段、运行策略、夜间基线或异常事件。请尽量简洁，避免冗长重复。"
                )
                prompt_context = self._build_analysis_prompt_context(payload, context, analysis_seed)
                prompt_context_str = json.dumps(prompt_context, ensure_ascii=False)
                if len(prompt_context_str) > 2800:
                    prompt_context_str = self._compact_analysis_prompt(prompt_context, max_chars=2800)
                user_prompt = (
                    f"当前分析上下文: {prompt_context_str}\n"
                    f"用户补充问题: {user_question}\n"
                    "请提炼最重要的趋势、原因、节能动作和运维动作。\n"
                    "如果证据不足，请明确说明证据不足，但仍要给出最稳妥的建议。\n"
                    "请输出严格JSON。"
                )
                last_err: Exception | None = None
                llm_obj: dict[str, Any] | None = None
                for attempt in range(max_retries + 1):
                    try:
                        response = llm_provider._call_chat_completion(
                            base_url=base_url,
                            api_key=api_key,
                            model=model,
                            timeout_sec=timeout_sec,
                            max_tokens=1100,
                            messages=[
                                {"role": "system", "content": system_prompt},
                                {"role": "user", "content": user_prompt},
                            ],
                        )
                        content = response.get("choices", [{}])[0].get("message", {}).get("content", "")
                        if isinstance(content, list):
                            content = "".join(
                                str(part.get("text", "")) if isinstance(part, dict) else str(part)
                                for part in content
                            )
                        llm_obj = llm_provider._extract_json_object(str(content))
                        break
                    except HTTPError as exc:
                        retryable = exc.code == 429 or 500 <= exc.code < 600
                        last_err = RuntimeError(f"llm http status {exc.code}")
                        if retryable and attempt < max_retries:
                            time.sleep(0.6 * (attempt + 1))
                            continue
                        break
                    except (URLError, TimeoutError, socket.timeout, ConnectionResetError) as exc:
                        last_err = RuntimeError(f"llm network error: {type(exc).__name__}")
                        if attempt < max_retries:
                            time.sleep(0.6 * (attempt + 1))
                            continue
                        break
                    except Exception as exc:
                        last_err = RuntimeError(str(exc) if str(exc) else f"llm parse error: {type(exc).__name__}")
                        if attempt < max_retries:
                            time.sleep(0.4 * (attempt + 1))
                            continue
                        break
                if llm_obj is None:
                    raise RuntimeError(str(last_err or "llm unknown error"))
                template_analysis = template_result["analysis"]
                result = {
                    "analysis": {
                        "summary": str(llm_obj.get("summary", "")).strip() or template_analysis["summary"],
                        "findings": self._merge_text_lists(llm_provider._coerce_list_of_str(llm_obj.get("findings")), template_analysis["findings"], max_items=5, min_items=3),
                        "possible_causes": self._merge_text_lists(llm_provider._coerce_list_of_str(llm_obj.get("possible_causes")), template_analysis["possible_causes"], max_items=4, min_items=3),
                        "energy_saving_suggestions": self._merge_text_lists(llm_provider._coerce_list_of_str(llm_obj.get("energy_saving_suggestions")), template_analysis["energy_saving_suggestions"], max_items=4, min_items=3),
                        "operations_suggestions": self._merge_text_lists(llm_provider._coerce_list_of_str(llm_obj.get("operations_suggestions")), template_analysis["operations_suggestions"], max_items=4, min_items=3),
                        "saving_opportunities": template_analysis.get("saving_opportunities", []),
                        "evidence": template_analysis["evidence"],
                        "knowledge_source": template_analysis.get("knowledge_source", "none"),
                        "retrieval_hit_count": template_analysis.get("retrieval_hit_count", 0),
                        "retrieval_error_type": template_analysis.get("retrieval_error_type", ""),
                    },
                    "context": context,
                }
                provider_name = llm_provider.name
        except Exception as exc:
            error_message = str(exc)
            if fallback_enabled:
                fallback_used = True
                result = template_result
            else:
                self._append_ai_event(
                    {
                        "timestamp": to_iso(dt.datetime.now()),
                        "trace_id": trace_id,
                        "requested_provider": preferred,
                        "provider": provider_name,
                        "fallback_used": False,
                        "latency_ms": int((time.perf_counter() - start) * 1000),
                        "error_type": "llm_error",
                        "building_id": template_result.get("context", {}).get("building_id"),
                        "event_type": "analysis",
                        "field_complete": False,
                        "result_risk_level": "",
                        "knowledge_source": str(template_result.get("analysis", {}).get("knowledge_source", "")),
                        "retrieval_hit_count": int(template_result.get("analysis", {}).get("retrieval_hit_count", 0) or 0),
                        "retrieval_error_type": str(template_result.get("analysis", {}).get("retrieval_error_type", "")),
                    }
                )
                self._raise_llm_failure(preferred, error_message)

        latency_ms = int((time.perf_counter() - start) * 1000)
        result["analysis"]["provider"] = provider_name
        result["analysis"]["requested_provider"] = preferred
        result["analysis"]["latency_ms"] = latency_ms
        result["analysis"]["fallback_used"] = fallback_used
        result["analysis"]["trace_id"] = trace_id
        if error_message:
            result["analysis"]["degrade_message"] = self._friendly_degrade_message(error_message)

        error_type = ""
        if fallback_used:
            msg = (error_message or "").lower()
            if "http status 429" in msg:
                error_type = "rate_limit"
            elif "maximum context length" in msg or "reduce the length of the messages" in msg:
                error_type = "context_too_long"
            elif "timed out" in msg:
                error_type = "timeout"
            elif "http status" in msg:
                error_type = "http_error"
            elif "network error" in msg:
                error_type = "network_error"
            elif "parse error" in msg:
                error_type = "parse_error"
            elif "not configured" in msg:
                error_type = "not_configured"
            else:
                error_type = "unknown_error"

        field_complete = self._required_analysis_fields_complete(result.get("analysis", {}))
        self._append_ai_event(
            {
                "timestamp": to_iso(dt.datetime.now()),
                "trace_id": trace_id,
                "requested_provider": preferred,
                "provider": provider_name,
                "fallback_used": fallback_used,
                "latency_ms": latency_ms,
                "error_type": error_type,
                "building_id": result.get("context", {}).get("building_id"),
                "event_type": "analysis",
                "field_complete": field_complete,
                "result_risk_level": "",
                "knowledge_source": str(result.get("analysis", {}).get("knowledge_source", "")),
                "retrieval_hit_count": int(result.get("analysis", {}).get("retrieval_hit_count", 0) or 0),
                "retrieval_error_type": str(result.get("analysis", {}).get("retrieval_error_type", "")),
            }
        )
        return result

    def diagnose(self, payload: dict[str, Any]) -> dict[str, Any]:
        preferred = str(payload.get("provider", "auto")).strip().lower()
        if preferred not in {"template", "llm", "auto"}:
            preferred = "auto"

        start = time.perf_counter()
        fallback_used = False
        error_message = None
        trace_id = uuid.uuid4().hex
        fallback_enabled = self._llm_template_fallback_enabled()

        try:
            if preferred in {"llm", "auto"}:
                result = self.providers["llm"].diagnose(self, payload)
                provider_name = self.providers["llm"].name
            else:
                result = self.providers["template"].diagnose(self, payload)
                provider_name = self.providers["template"].name
        except Exception as exc:
            error_message = str(exc)
            if fallback_enabled:
                fallback_used = True
                result = self.providers["template"].diagnose(self, payload)
                provider_name = self.providers["template"].name
            else:
                self._append_ai_event(
                    {
                        "timestamp": to_iso(dt.datetime.now()),
                        "trace_id": trace_id,
                        "requested_provider": preferred,
                        "provider": "llm_provider",
                        "fallback_used": False,
                        "latency_ms": int((time.perf_counter() - start) * 1000),
                        "error_type": "llm_error",
                        "anomaly_id": None,
                        "has_message": bool(str(payload.get("message", "")).strip()),
                        "event_type": "diagnose",
                        "field_complete": False,
                        "result_risk_level": "",
                        "knowledge_source": "",
                        "retrieval_hit_count": 0,
                        "retrieval_error_type": "",
                    }
                )
                self._raise_llm_failure(preferred, error_message)

        latency_ms = int((time.perf_counter() - start) * 1000)
        result["diagnosis"]["provider"] = provider_name
        result["diagnosis"]["requested_provider"] = preferred
        result["diagnosis"]["latency_ms"] = latency_ms
        result["diagnosis"]["fallback_used"] = fallback_used
        result["diagnosis"]["trace_id"] = trace_id
        if error_message:
            result["diagnosis"]["degrade_message"] = self._friendly_degrade_message(error_message)

        error_type = ""
        if fallback_used:
            msg = (error_message or "").lower()
            if "http status 429" in msg:
                error_type = "rate_limit"
            elif "http status" in msg:
                error_type = "http_error"
            elif "network error" in msg:
                error_type = "network_error"
            elif "parse error" in msg:
                error_type = "parse_error"
            elif "not configured" in msg:
                error_type = "not_configured"
            else:
                error_type = "unknown_error"
        field_complete = self._required_diag_fields_complete(result.get("diagnosis", {}))
        self._append_ai_event(
            {
                "timestamp": to_iso(dt.datetime.now()),
                "trace_id": trace_id,
                "requested_provider": preferred,
                "provider": provider_name,
                "fallback_used": fallback_used,
                "latency_ms": latency_ms,
                "error_type": error_type,
                "anomaly_id": result.get("context", {}).get("anomaly_id"),
                "has_message": bool(str(payload.get("message", "")).strip()),
                "event_type": "diagnose",
                "field_complete": field_complete,
                "result_risk_level": str(result.get("diagnosis", {}).get("risk_level", "")),
                "knowledge_source": str(result.get("diagnosis", {}).get("knowledge_source", "")),
                "retrieval_hit_count": int(result.get("diagnosis", {}).get("retrieval_hit_count", 0) or 0),
                "retrieval_error_type": str(result.get("diagnosis", {}).get("retrieval_error_type", "")),
            }
        )
        return result

    def _type_from_keywords(self, text: str) -> str:
        lowered = text.lower()
        for anomaly_type, details in self.dict_data.items():
            keywords = details.get("keywords", [])
            if any(k.lower() in lowered for k in keywords):
                return anomaly_type
        return "anomaly_spike"

    def _friendly_degrade_message(self, error_message: str | None) -> str:
        raw = str(error_message or "").strip()
        lowered = raw.lower()
        if "not configured" in lowered:
            return "当前环境未配置 DeepSeek API Key，本次使用模板兜底。"
        if "maximum context length" in lowered or "reduce the length of the messages" in lowered:
            return "发送给 DeepSeek 的分析上下文过长，系统已切换到模板兜底。"
        if "timed out" in lowered:
            return "DeepSeek 分析响应超时，系统已切换到模板兜底。"
        if "network error" in lowered:
            return "DeepSeek 网络请求失败，本次使用模板兜底。"
        if "http status 429" in lowered:
            return "DeepSeek 当前触发限流，本次使用模板兜底。"
        if "http status" in lowered:
            return "DeepSeek 服务响应异常，本次使用模板兜底。"
        if "parse error" in lowered:
            return "DeepSeek 返回格式异常，本次使用模板兜底。"
        if "simulated llm failure" in lowered:
            return "当前为模拟失败场景，本次使用模板兜底。"
        return "在线模型暂时不可用，系统已切换到模板兜底。"


class FileImportSource(EnergyRepository):
    storage_backend = "file"


class MySQLRepository(EnergyRepository):
    storage_backend = "mysql"

    def __init__(
        self,
        demo_data_file: Path,
        normalized_data_file: Path,
        metadata_file: Path,
        weather_file: Path,
        dict_file: Path,
        knowledge_file: Path,
        action_log_file: Path,
        ai_call_log_file: Path,
        note_log_file: Path,
        regression_summary_file: Path,
        mysql_client: MySQLClient | None = None,
    ) -> None:
        self.mysql = mysql_client or MySQLClient.from_env()
        super().__init__(
            demo_data_file,
            normalized_data_file,
            metadata_file,
            weather_file,
            dict_file,
            knowledge_file,
            action_log_file,
            ai_call_log_file,
            note_log_file,
            regression_summary_file,
        )

    def _mysql_table_count(self, table_name: str) -> int:
        try:
            value = self.mysql.query_scalar(f"SELECT COUNT(*) FROM {table_name}")
            return int(str(value or "0").strip() or "0")
        except Exception:
            return 0

    def _ensure_runtime_storage(self) -> None:
        health = self.mysql.health()
        if not health["configured"]:
            raise RuntimeError("MySQL backend selected but MYSQL_HOST/MYSQL_DATABASE/MYSQL_USER not configured")
        if not health["available"]:
            raise RuntimeError("MySQL backend selected but mysql client not found")
        self.mysql.ensure_schema()

    def _load_bdg2_metadata(self) -> dict[str, dict[str, Any]]:
        if self._mysql_table_count("buildings") <= 0:
            return {}
        rows = self.mysql.query_json_rows(
            """
            SELECT JSON_OBJECT(
                'building_id', building_id,
                'site_id', site_id,
                'primaryspaceusage', primaryspaceusage,
                'sub_primaryspaceusage', sub_primaryspaceusage,
                'peer_category', peer_category,
                'display_category', display_category,
                'display_name', display_name
            )
            FROM buildings
            ORDER BY building_id
            """
        )
        metadata: dict[str, dict[str, Any]] = {}
        for raw in rows:
            building_id = str(raw.get("building_id", "")).strip()
            if not building_id:
                continue
            metadata[building_id] = {
                "building_id": building_id,
                "site_id": str(raw.get("site_id", "")).strip(),
                "primaryspaceusage": str(raw.get("primaryspaceusage", "")).strip(),
                "sub_primaryspaceusage": str(raw.get("sub_primaryspaceusage", "")).strip(),
                "peer_category": str(raw.get("peer_category", "")).strip() or None,
                "display_category": str(raw.get("display_category", "")).strip(),
                "display_name": str(raw.get("display_name", "")).strip() or showcase_display_name(building_id),
            }
        return metadata

    def _load_rows(self) -> list[dict[str, Any]]:
        if self._mysql_table_count("energy_timeseries") <= 0:
            return []
        rows = self.mysql.query_json_rows(
            """
            SELECT JSON_OBJECT(
                'record_id', record_id,
                'building_id', building_id,
                'building_name', building_name,
                'building_type', building_type,
                'timestamp', DATE_FORMAT(ts, '%Y-%m-%d %H:%i:%s'),
                'hour', hour_of_day,
                'electricity_kwh', value,
                'source', source
            )
            FROM energy_timeseries
            WHERE metric_type = 'electricity'
            ORDER BY building_id, ts
            """,
            timeout_sec=120,
        )
        normalized_rows: list[dict[str, Any]] = []
        for raw in rows:
            ts = parse_time(str(raw.get("timestamp", "")))
            if not ts:
                continue
            try:
                record_id = int(raw.get("record_id") or len(normalized_rows) + 1)
            except (TypeError, ValueError):
                record_id = len(normalized_rows) + 1
            normalized_rows.append(
                {
                    "record_id": record_id,
                    "building_id": str(raw.get("building_id", "")).strip(),
                    "building_name": str(raw.get("building_name", "")).strip(),
                    "building_type": str(raw.get("building_type", "")).strip(),
                    "timestamp": ts,
                    "hour": int(raw.get("hour") or ts.hour),
                    "electricity_kwh": float(raw.get("electricity_kwh") or 0.0),
                    "source": str(raw.get("source", "mysql")).strip() or "mysql",
                }
            )
        return normalized_rows

    def _load_raw_electricity_headers(self) -> set[str]:
        if self._mysql_table_count("energy_timeseries") <= 0:
            return set()
        rows = self.mysql.query_rows(
            """
            SELECT DISTINCT building_id
            FROM energy_timeseries
            WHERE metric_type = 'electricity'
            """
        )
        return {str(row.get("building_id", "")).strip() for row in rows if str(row.get("building_id", "")).strip()}

    def _peer_compare_pool(
        self,
        peer_category: str,
        start_time: dt.datetime | None,
        end_time: dt.datetime | None,
    ) -> dict[str, float]:
        cache_key = (
            peer_category,
            to_iso(start_time) if start_time else None,
            to_iso(end_time) if end_time else None,
        )
        cached = self.compare_pool_cache.get(cache_key)
        if cached is not None:
            return dict(cached)

        candidate_ids = list(self.peer_category_to_buildings.get(peer_category, []))
        if not candidate_ids:
            self.compare_pool_cache[cache_key] = {}
            return {}

        placeholders = ", ".join(["%s"] * len(candidate_ids))
        params: list[Any] = list(candidate_ids)
        sql = f"""
        SELECT building_id, SUM(total_value) / NULLIF(SUM(sample_count), 0) AS avg_value
        FROM peer_energy_daily
        WHERE metric_type = 'electricity' AND building_id IN ({placeholders})
        """
        if start_time:
            sql += " AND day >= %s"
            params.append(start_time.date().isoformat())
        if end_time:
            sql += " AND day <= %s"
            params.append(end_time.date().isoformat())
        sql += " GROUP BY building_id"
        rows = self.mysql.query_rows(sql, params=params, timeout_sec=120)
        result = {
            str(row.get("building_id", "")).strip(): float(row.get("avg_value") or 0.0)
            for row in rows
            if str(row.get("building_id", "")).strip()
        }
        self.compare_pool_cache[cache_key] = dict(result)
        return result

    def _load_weather_by_site(self) -> dict[str, dict[dt.datetime, dict[str, float]]]:
        if self._mysql_table_count("weather_timeseries") <= 0:
            return {}
        rows = self.mysql.query_json_rows(
            """
            SELECT JSON_OBJECT(
                'site_id', site_id,
                'timestamp', DATE_FORMAT(ts, '%Y-%m-%d %H:%i:%s'),
                'temperature_c', temperature_c,
                'wind_speed', wind_speed
            )
            FROM weather_timeseries
            ORDER BY site_id, ts
            """,
            timeout_sec=120,
        )
        weather_by_site: dict[str, dict[dt.datetime, dict[str, float]]] = {}
        for raw in rows:
            site_id = str(raw.get("site_id", "")).strip()
            timestamp = parse_time(str(raw.get("timestamp", "")))
            if not site_id or not timestamp:
                continue
            weather_by_site.setdefault(site_id, {})[timestamp] = {
                "temperature_c": round(float(raw.get("temperature_c") or 0.0), 2),
                "wind_speed": round(float(raw.get("wind_speed") or 0.0), 2),
            }
        return weather_by_site

    def _load_actions(self) -> None:
        if self._mysql_table_count("anomaly_actions") <= 0:
            self.action_events = []
            self._rebuild_action_index()
            return
        rows = self.mysql.query_json_rows(
            """
            SELECT JSON_OBJECT(
                'anomaly_id', anomaly_id,
                'action', action_name,
                'status_before', status_before,
                'status', status_after,
                'assignee', assignee,
                'note', note,
                'created_at', DATE_FORMAT(created_at, '%Y-%m-%d %H:%i:%s')
            )
            FROM anomaly_actions
            ORDER BY created_at, id
            """
        )
        self.action_events = []
        for event in rows:
            try:
                event["anomaly_id"] = int(event.get("anomaly_id"))
            except (TypeError, ValueError):
                continue
            event["created_at"] = str(event.get("created_at", ""))
            self.action_events.append(event)
        self._rebuild_action_index()

    def _append_action_event(self, event: dict[str, Any]) -> None:
        sql = f"""
        INSERT INTO anomaly_actions (
            anomaly_id, action_name, status_before, status_after, assignee, note, created_at
        ) VALUES (
            {int(event['anomaly_id'])},
            {sql_literal(event.get('action', ''))},
            {sql_literal(event.get('status_before', ''))},
            {sql_literal(event.get('status', ''))},
            {sql_literal(event.get('assignee', ''))},
            {sql_literal(event.get('note', ''))},
            {sql_literal(event.get('created_at', ''))}
        )
        """
        self.mysql.execute(sql)
        self.action_events.append(event)
        self._rebuild_action_index()

    def _load_notes(self) -> None:
        if self._mysql_table_count("anomaly_notes") <= 0:
            self.note_events = []
            self._rebuild_note_index()
            return
        rows = self.mysql.query_json_rows(
            """
            SELECT JSON_OBJECT(
                'anomaly_id', anomaly_id,
                'cause_confirmed', cause_confirmed,
                'action_taken', action_taken,
                'result_summary', result_summary,
                'recurrence_risk', recurrence_risk,
                'reviewer', reviewer,
                'updated_at', DATE_FORMAT(updated_at, '%Y-%m-%d %H:%i:%s')
            )
            FROM anomaly_notes
            ORDER BY updated_at, id
            """
        )
        self.note_events = []
        for event in rows:
            try:
                event["anomaly_id"] = int(event.get("anomaly_id"))
            except (TypeError, ValueError):
                continue
            event["updated_at"] = str(event.get("updated_at", ""))
            self.note_events.append(event)
        self._rebuild_note_index()

    def _append_note_event(self, event: dict[str, Any]) -> None:
        sql = f"""
        INSERT INTO anomaly_notes (
            anomaly_id, cause_confirmed, action_taken, result_summary, recurrence_risk, reviewer, updated_at
        ) VALUES (
            {int(event['anomaly_id'])},
            {sql_literal(event.get('cause_confirmed', ''))},
            {sql_literal(event.get('action_taken', ''))},
            {sql_literal(event.get('result_summary', ''))},
            {sql_literal(event.get('recurrence_risk', 'medium'))},
            {sql_literal(event.get('reviewer', ''))},
            {sql_literal(event.get('updated_at', ''))}
        )
        """
        self.mysql.execute(sql)
        self.note_events.append(event)
        self._rebuild_note_index()

    def _load_ai_events(self) -> None:
        if self._mysql_table_count("ai_calls") <= 0:
            self.ai_events = []
            return
        rows = self.mysql.query_json_rows(
            """
            SELECT JSON_OBJECT(
                'timestamp', DATE_FORMAT(event_time, '%Y-%m-%d %H:%i:%s'),
                'trace_id', trace_id,
                'requested_provider', requested_provider,
                'provider', provider,
                'building_id', building_id,
                'anomaly_id', anomaly_id,
                'has_message', has_message,
                'event_type', scene,
                'fallback_used', fallback_used,
                'latency_ms', latency_ms,
                'error_type', error_type,
                'field_complete', field_complete,
                'result_risk_level', result_risk_level,
                'knowledge_source', knowledge_source,
                'retrieval_hit_count', retrieval_hit_count,
                'retrieval_error_type', retrieval_error_type,
                'feedback_label', feedback_label
            )
            FROM ai_calls
            ORDER BY event_time, trace_id
            """
        )
        self.ai_events = list(rows)

    def _append_ai_event(self, event: dict[str, Any]) -> None:
        timestamp = str(event.get("timestamp", "")).strip() or to_iso(dt.datetime.now())
        trace_id = str(event.get("trace_id", "")).strip() or uuid.uuid4().hex
        success = 0 if str(event.get("error_type", "")).strip() else 1
        sql = f"""
        INSERT INTO ai_calls (
            trace_id, event_time, requested_provider, provider, scene, building_id, anomaly_id,
            has_message, result_risk_level, knowledge_source, retrieval_hit_count, retrieval_error_type,
            fallback_used, field_complete, latency_ms, success, error_type, feedback_label, created_at, updated_at
        ) VALUES (
            {sql_literal(trace_id)},
            {sql_literal(timestamp)},
            {sql_literal(event.get('requested_provider', ''))},
            {sql_literal(event.get('provider', ''))},
            {sql_literal(event.get('event_type', ''))},
            {sql_literal(event.get('building_id'))},
            {sql_literal(event.get('anomaly_id'))},
            {1 if bool(event.get('has_message', False)) else 0},
            {sql_literal(event.get('result_risk_level', ''))},
            {sql_literal(event.get('knowledge_source', ''))},
            {int(event.get('retrieval_hit_count', 0) or 0)},
            {sql_literal(event.get('retrieval_error_type', ''))},
            {1 if bool(event.get('fallback_used', False)) else 0},
            {1 if bool(event.get('field_complete', False)) else 0},
            {int(event.get('latency_ms', 0) or 0)},
            {success},
            {sql_literal(event.get('error_type', ''))},
            {sql_literal(event.get('feedback_label', ''))},
            {sql_literal(timestamp)},
            {sql_literal(timestamp)}
        )
        ON DUPLICATE KEY UPDATE
            requested_provider = VALUES(requested_provider),
            provider = VALUES(provider),
            scene = VALUES(scene),
            building_id = VALUES(building_id),
            anomaly_id = VALUES(anomaly_id),
            has_message = VALUES(has_message),
            result_risk_level = VALUES(result_risk_level),
            knowledge_source = VALUES(knowledge_source),
            retrieval_hit_count = VALUES(retrieval_hit_count),
            retrieval_error_type = VALUES(retrieval_error_type),
            fallback_used = VALUES(fallback_used),
            field_complete = VALUES(field_complete),
            latency_ms = VALUES(latency_ms),
            success = VALUES(success),
            error_type = VALUES(error_type),
            feedback_label = VALUES(feedback_label),
            updated_at = VALUES(updated_at)
        """
        self.mysql.execute(sql)
        for idx, existing in enumerate(self.ai_events):
            if str(existing.get("trace_id", "")) == trace_id:
                merged = dict(existing)
                merged.update(event)
                merged["trace_id"] = trace_id
                merged["timestamp"] = timestamp
                self.ai_events[idx] = merged
                break
        else:
            merged = dict(event)
            merged["trace_id"] = trace_id
            merged["timestamp"] = timestamp
            self.ai_events.append(merged)

    def save_ai_feedback(self, payload: dict[str, Any]) -> dict[str, Any]:
        trace_id = str(payload.get("trace_id", "")).strip()
        label = str(payload.get("label", "")).strip().lower()
        if not trace_id:
            raise ValueError("trace_id required")
        if label not in {"useful", "not_useful"}:
            raise ValueError("label must be useful/not_useful")
        sql = f"""
        UPDATE ai_calls
        SET feedback_label = {sql_literal(label)}, updated_at = {sql_literal(to_iso(dt.datetime.now()))}
        WHERE trace_id = {sql_literal(trace_id)}
        """
        self.mysql.execute(sql)
        row_count = self.mysql.query_scalar(
            f"SELECT COUNT(*) FROM ai_calls WHERE trace_id = {sql_literal(trace_id)}"
        )
        if int(str(row_count or "0") or "0") <= 0:
            raise LookupError("trace_id not found")
        for ev in self.ai_events:
            if str(ev.get("trace_id", "")) == trace_id:
                ev["feedback_label"] = label
                break
        return {"trace_id": trace_id, "label": label}

    def _load_regression_summary_from_mysql(self) -> dict[str, Any] | None:
        if self._mysql_table_count("system_snapshots") <= 0:
            return None
        rows = self.mysql.query_json_rows(
            """
            SELECT JSON_OBJECT(
                'payload_json', payload_json
            )
            FROM system_snapshots
            WHERE snapshot_key = 'regression_summary'
            LIMIT 1
            """
        )
        if not rows:
            return None
        payload_raw = str(rows[0].get("payload_json", "")).strip()
        if not payload_raw:
            return None
        try:
            return json.loads(payload_raw)
        except json.JSONDecodeError:
            return None

    def query_system_health(self) -> dict[str, Any]:
        data = super().query_system_health()
        mysql_health = self.mysql.health()
        regression = self._load_regression_summary_from_mysql()
        if regression:
            data["recent_regression"] = regression
        data["storage"] = {
            "backend": self.storage_backend,
            "mysql": mysql_health,
        }
        if data["status"] == "ok" and not mysql_health["connected"]:
            data["status"] = "degraded"
        return data


FileRepository = FileImportSource


def create_repository() -> EnergyRepository:
    backend = (os.getenv("STORAGE_BACKEND", STORAGE_BACKEND_DEFAULT).strip().lower() or STORAGE_BACKEND_DEFAULT)
    repo_cls: type[EnergyRepository] = MySQLRepository if backend != "file" else FileImportSource
    return repo_cls(
        DEMO_DATA_FILE,
        NORMALIZED_DATA_FILE,
        METADATA_FILE,
        WEATHER_FILE,
        DICT_FILE,
        KNOWLEDGE_FILE,
        ACTION_LOG_FILE,
        AI_CALL_LOG_FILE,
        NOTE_LOG_FILE,
        REGRESSION_SUMMARY_FILE,
    )


REPO = create_repository()


class Handler(BaseHTTPRequestHandler):
    server_version = "A8EnergyServer/0.4"

    def _set_json_headers(self, status: int = HTTPStatus.OK) -> None:
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET,POST,OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def _set_file_headers(self, status: int, content_type: str) -> None:
        self.send_response(status)
        self.send_header("Content-Type", content_type)
        self.send_header("Cache-Control", "no-store, no-cache, must-revalidate")
        self.send_header("Pragma", "no-cache")
        self.send_header("Expires", "0")
        self.send_header("Access-Control-Allow-Origin", "*")
        self.end_headers()

    def _set_sse_headers(self) -> None:
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", "text/event-stream; charset=utf-8")
        self.send_header("Cache-Control", "no-cache")
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET,POST,OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def _send_sse(self, event: str, data: Any) -> bool:
        """Write one SSE event; returns False if the connection was broken."""
        line = f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"
        try:
            self.wfile.write(line.encode("utf-8"))
            self.wfile.flush()
            return True
        except (BrokenPipeError, ConnectionAbortedError, ConnectionResetError):
            return False

    def _json(self, data: Any, status: int = HTTPStatus.OK) -> None:
        self._set_json_headers(status)
        payload = json.dumps(data, ensure_ascii=False).encode("utf-8")
        try:
            self.wfile.write(payload)
        except (BrokenPipeError, ConnectionAbortedError, ConnectionResetError):
            return

    def _csv(self, text: str, filename: str) -> None:
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", "text/csv; charset=utf-8")
        self.send_header("Content-Disposition", f'attachment; filename="{filename}"')
        self.send_header("Access-Control-Allow-Origin", "*")
        self.end_headers()
        self.wfile.write(text.encode("utf-8-sig"))

    def _file_bytes(self, content: bytes, filename: str, content_type: str) -> None:
        encoded_name = quote(filename, safe="")
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{encoded_name}")
        self.send_header("Content-Length", str(len(content)))
        self.send_header("Cache-Control", "no-store, no-cache, must-revalidate")
        self.send_header("Pragma", "no-cache")
        self.send_header("Expires", "0")
        self.send_header("Access-Control-Allow-Origin", "*")
        self.end_headers()
        try:
            self.wfile.write(content)
        except (BrokenPipeError, ConnectionAbortedError, ConnectionResetError):
            return

    def do_OPTIONS(self) -> None:
        self._set_json_headers(HTTPStatus.NO_CONTENT)

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path.startswith("/api/"):
            self._handle_api_get(parsed)
            return
        self._serve_static(parsed.path)

    def do_POST(self) -> None:
        parsed = urlparse(self.path)
        length = int(self.headers.get("Content-Length", "0"))
        body = self.rfile.read(length) if length else b"{}"
        try:
            payload = json.loads(body.decode("utf-8"))
        except json.JSONDecodeError:
            self._json({"code": 400, "message": "invalid json", "data": None}, HTTPStatus.BAD_REQUEST)
            return

        if parsed.path == "/api/ai/stream":
            self._set_sse_headers()
            try:
                for event_name, event_data in REPO.diagnose_stream_events(payload):
                    if not self._send_sse(event_name, event_data):
                        break
            except Exception as exc:
                self._send_sse("error", {"message": str(exc)})
            return

        if parsed.path == "/api/ai/diagnose":
            try:
                result = REPO.diagnose(payload)
            except ValueError as exc:
                self._json({"code": 400, "message": str(exc), "data": None}, HTTPStatus.BAD_REQUEST)
                return
            except Exception as exc:
                self._json({"code": 502, "message": str(exc), "data": None}, HTTPStatus.BAD_GATEWAY)
                return
            self._json({"code": 0, "message": "ok", "data": result})
            return

        if parsed.path == "/api/ai/analyze":
            try:
                result = REPO.analyze(payload)
            except ValueError as exc:
                self._json({"code": 400, "message": str(exc), "data": None}, HTTPStatus.BAD_REQUEST)
                return
            except Exception as exc:
                self._json({"code": 502, "message": str(exc), "data": None}, HTTPStatus.BAD_GATEWAY)
                return
            self._json({"code": 0, "message": "ok", "data": result})
            return

        if parsed.path == "/api/assistant/report/export":
            try:
                content, filename, content_type = REPO.export_assistant_report(payload)
            except ValueError as exc:
                self._json({"code": 400, "message": str(exc), "data": None}, HTTPStatus.BAD_REQUEST)
                return
            except Exception as exc:
                self._json({"code": 502, "message": str(exc), "data": None}, HTTPStatus.BAD_GATEWAY)
                return
            self._file_bytes(content, filename, content_type)
            return

        if parsed.path == "/api/ragflow/chat/stream":
            self._set_sse_headers()
            try:
                for event_name, event_data in REPO.ragflow_chat_stream_events(payload):
                    if not self._send_sse(event_name, event_data):
                        break
            except Exception as exc:
                self._send_sse("error", {"message": str(exc)})
            return

        if parsed.path == "/api/ragflow/chat":
            try:
                result = REPO.ask_ragflow_chat(payload)
            except ValueError as exc:
                self._json({"code": 400, "message": str(exc), "data": None}, HTTPStatus.BAD_REQUEST)
                return
            except Exception as exc:
                self._json({"code": 502, "message": str(exc), "data": None}, HTTPStatus.BAD_GATEWAY)
                return
            self._json({"code": 0, "message": "ok", "data": result})
            return

        if parsed.path == "/api/ragflow/reference/document":
            try:
                result = REPO.ask_ragflow_reference_document(payload)
            except ValueError as exc:
                self._json({"code": 400, "message": str(exc), "data": None}, HTTPStatus.BAD_REQUEST)
                return
            except LookupError as exc:
                self._json({"code": 404, "message": str(exc), "data": None}, HTTPStatus.NOT_FOUND)
                return
            except Exception as exc:
                self._json({"code": 502, "message": str(exc), "data": None}, HTTPStatus.BAD_GATEWAY)
                return
            self._json({"code": 0, "message": "ok", "data": result})
            return

        if parsed.path == "/api/anomaly/action":
            try:
                result = REPO.apply_anomaly_action(payload)
            except ValueError as exc:
                self._json({"code": 400, "message": str(exc), "data": None}, HTTPStatus.BAD_REQUEST)
                return
            except LookupError as exc:
                self._json({"code": 404, "message": str(exc), "data": None}, HTTPStatus.NOT_FOUND)
                return
            self._json({"code": 0, "message": "ok", "data": result})
            return

        if parsed.path == "/api/anomaly/note":
            try:
                result = REPO.upsert_anomaly_note(payload)
            except ValueError as exc:
                self._json({"code": 400, "message": str(exc), "data": None}, HTTPStatus.BAD_REQUEST)
                return
            except LookupError as exc:
                self._json({"code": 404, "message": str(exc), "data": None}, HTTPStatus.NOT_FOUND)
                return
            self._json({"code": 0, "message": "ok", "data": result})
            return

        if parsed.path == "/api/ai/evaluate":
            hours_raw = payload.get("hours", 24)
            try:
                hours = int(hours_raw)
            except (TypeError, ValueError):
                hours = 24
            self._json({"code": 0, "message": "ok", "data": REPO.query_ai_evaluate(hours)})
            return

        if parsed.path == "/api/ai/feedback":
            try:
                result = REPO.save_ai_feedback(payload)
            except ValueError as exc:
                self._json({"code": 400, "message": str(exc), "data": None}, HTTPStatus.BAD_REQUEST)
                return
            except LookupError as exc:
                self._json({"code": 404, "message": str(exc), "data": None}, HTTPStatus.NOT_FOUND)
                return
            self._json({"code": 0, "message": "ok", "data": result})
            return

        self._json({"code": 404, "message": "not found", "data": None}, HTTPStatus.NOT_FOUND)

    def _handle_api_get(self, parsed: Any) -> None:
        params = parse_qs(parsed.query)
        building_id = params.get("building_id", [None])[0]
        start_time = parse_time(params.get("start_time", [None])[0])
        end_time = parse_time(params.get("end_time", [None])[0])
        metric_type = params.get("metric_type", ["electricity"])[0]

        if parsed.path == "/api/buildings":
            self._json({"code": 0, "message": "ok", "data": REPO.query_buildings()})
            return

        if parsed.path == "/api/energy/trend":
            self._json({"code": 0, "message": "ok", "data": REPO.query_trend(building_id, start_time, end_time)})
            return

        if parsed.path == "/api/energy/rank":
            month = params.get("month", [None])[0]
            self._json({"code": 0, "message": "ok", "data": REPO.query_rank(month)})
            return

        if parsed.path == "/api/analysis/summary":
            try:
                data = REPO.query_analysis_summary(building_id, start_time, end_time, metric_type)
            except ValueError as exc:
                self._json({"code": 400, "message": str(exc), "data": None}, HTTPStatus.BAD_REQUEST)
                return
            self._json({"code": 0, "message": "ok", "data": data})
            return

        if parsed.path == "/api/analysis/trend":
            try:
                data = REPO.query_analysis_trend(building_id, start_time, end_time, metric_type)
            except ValueError as exc:
                self._json({"code": 400, "message": str(exc), "data": None}, HTTPStatus.BAD_REQUEST)
                return
            self._json({"code": 0, "message": "ok", "data": data})
            return

        if parsed.path == "/api/analysis/distribution":
            try:
                data = REPO.query_analysis_distribution(building_id, start_time, end_time, metric_type)
            except ValueError as exc:
                self._json({"code": 400, "message": str(exc), "data": None}, HTTPStatus.BAD_REQUEST)
                return
            self._json({"code": 0, "message": "ok", "data": data})
            return

        if parsed.path == "/api/analysis/insights":
            try:
                data = REPO.query_analysis_insights(building_id, start_time, end_time, metric_type)
            except ValueError as exc:
                self._json({"code": 400, "message": str(exc), "data": None}, HTTPStatus.BAD_REQUEST)
                return
            self._json({"code": 0, "message": "ok", "data": data})
            return

        if parsed.path == "/api/analysis/compare":
            try:
                data = REPO.query_analysis_compare(building_id, start_time, end_time, metric_type)
            except ValueError as exc:
                self._json({"code": 400, "message": str(exc), "data": None}, HTTPStatus.BAD_REQUEST)
                return
            self._json({"code": 0, "message": "ok", "data": data})
            return

        if parsed.path == "/api/anomaly/list":
            anomaly_type = params.get("anomaly_type", [None])[0]
            severity = params.get("severity", [None])[0]
            status = params.get("status", [None])[0]
            if status and status not in STATUS_VALUES:
                self._json({"code": 400, "message": "invalid status", "data": None}, HTTPStatus.BAD_REQUEST)
                return
            sort = params.get("sort", ["timestamp_desc"])[0]
            try:
                page = int(params.get("page", ["1"])[0])
            except ValueError:
                page = 1
            try:
                page_size = int(params.get("page_size", ["20"])[0])
            except ValueError:
                page_size = 20

            data = REPO.query_anomalies(building_id, start_time, end_time, anomaly_type, severity, status, page, page_size, sort)
            self._json({"code": 0, "message": "ok", "data": data})
            return

        if parsed.path == "/api/anomaly/history":
            anomaly_id = params.get("anomaly_id", [None])[0]
            if not anomaly_id:
                self._json({"code": 400, "message": "anomaly_id required", "data": None}, HTTPStatus.BAD_REQUEST)
                return
            try:
                data = REPO.query_anomaly_history(int(anomaly_id))
            except ValueError:
                self._json({"code": 400, "message": "invalid anomaly_id", "data": None}, HTTPStatus.BAD_REQUEST)
                return
            if not data:
                self._json({"code": 404, "message": "anomaly not found", "data": None}, HTTPStatus.NOT_FOUND)
                return
            self._json({"code": 0, "message": "ok", "data": data})
            return

        if parsed.path == "/api/anomaly/export":
            anomaly_type = params.get("anomaly_type", [None])[0]
            severity = params.get("severity", [None])[0]
            status = params.get("status", [None])[0]
            sort = params.get("sort", ["timestamp_desc"])[0]
            content = REPO.export_anomalies_csv(building_id, start_time, end_time, anomaly_type, severity, status, sort)
            ts = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
            self._csv(content, f"a8_anomaly_export_{ts}.csv")
            return

        if parsed.path == "/api/ai/stats":
            try:
                hours = int(params.get("hours", ["24"])[0])
            except ValueError:
                hours = 24
            self._json({"code": 0, "message": "ok", "data": REPO.query_ai_stats(hours)})
            return

        if parsed.path == "/api/system/health":
            self._json({"code": 0, "message": "ok", "data": REPO.query_system_health()})
            return

        if parsed.path == "/api/metrics/overview":
            self._json({"code": 0, "message": "ok", "data": REPO.query_metrics_overview(building_id, start_time, end_time)})
            return

        if parsed.path == "/api/metrics/saving-potential":
            self._json({"code": 0, "message": "ok", "data": REPO.query_saving_potential(building_id, start_time, end_time)})
            return

        if parsed.path == "/api/anomaly/detail":
            anomaly_id = params.get("anomaly_id", [None])[0]
            if not anomaly_id:
                self._json({"code": 400, "message": "anomaly_id required", "data": None}, HTTPStatus.BAD_REQUEST)
                return
            try:
                data = REPO.query_anomaly_detail(int(anomaly_id))
            except ValueError:
                self._json({"code": 400, "message": "invalid anomaly_id", "data": None}, HTTPStatus.BAD_REQUEST)
                return
            if not data:
                self._json({"code": 404, "message": "anomaly not found", "data": None}, HTTPStatus.NOT_FOUND)
                return
            self._json({"code": 0, "message": "ok", "data": data})
            return

        self._json({"code": 404, "message": "not found", "data": None}, HTTPStatus.NOT_FOUND)

    def _serve_static(self, path: str) -> None:
        clean_path = path or "/"
        target = FRONTEND_DIR / "index.html" if clean_path in ("/", "/index.html") else FRONTEND_DIR / clean_path.lstrip("/")
        if not target.exists() or not target.is_file():
            self._json({"code": 404, "message": "not found", "data": None}, HTTPStatus.NOT_FOUND)
            return

        content_type = "text/plain; charset=utf-8"
        if target.suffix == ".html":
            content_type = "text/html; charset=utf-8"
        elif target.suffix == ".css":
            content_type = "text/css; charset=utf-8"
        elif target.suffix == ".js":
            content_type = "application/javascript; charset=utf-8"

        self._set_file_headers(HTTPStatus.OK, content_type)
        self.wfile.write(target.read_bytes())


def run(host: str = "127.0.0.1", port: int = 8000) -> None:
    server = ThreadingHTTPServer((host, port), Handler)
    print(f"A8 server running at http://{host}:{port}")
    print("Press Ctrl+C to stop.")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("Shutting down...")
    finally:
        server.server_close()


if __name__ == "__main__":
    run()
